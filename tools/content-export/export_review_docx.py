#!/usr/bin/env python3
"""Client-review .docx renderer for the content plugin (house format).

Renders generated markdown pages into a single Word document built for client
review and Google Docs import.

House format spec:
  * "Page N" opens every page, styled Heading 1 so it lands in the Google Docs
    outline / document tabs. Larger, and the only element in a different font.
  * Each page starts on a new page (page break before every "Page N" after the
    first).
  * Top-section sequence: Page N -> URL -> blank -> Primary Keyword ->
    Secondary Keywords -> blank -> Meta Title -> Meta Description.
  * Body headings render as literal "H1:" / "H2:" / "H3:" tags. The tag is grey,
    the heading text is black and bold, both at body font size.
  * FAQ questions (a line that is entirely bold) render as H3.
  * Bullets render as real Word bullets. Links render standard blue + underlined
    and remain clickable.
  * All content black; only heading tags are grey and only links are blue.

Usage:
    python3 export_review_docx.py <client_content_dir> <out.docx> [entry_id ...]

If no entry ids are given, every entry with content/<id>/generated.md is used,
service hubs first, then the rest alphabetically.
"""
from __future__ import annotations

import json
import re
import sys
from pathlib import Path

import docx
from docx.enum.text import WD_BREAK
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

BODY_FONT = "Calibri"
BODY_SIZE = Pt(11)
TITLE_FONT = "Georgia"          # the only different font, per spec
TITLE_SIZE = Pt(20)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x80, 0x80, 0x80)
LINK_BLUE = "0563C1"

# [text](url)  and  **bold**
_INLINE = re.compile(r"(\[[^\]]+\]\([^)]+\)|\*\*[^*]+\*\*)")
_LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_BOLD_ONLY = re.compile(r"^\*\*(.+?)\*\*:?$")


def add_hyperlink(paragraph, url: str, text: str) -> None:
    """Insert a real, clickable hyperlink run (blue + underlined)."""
    r_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    colour = OxmlElement("w:color")
    colour.set(qn("w:val"), LINK_BLUE)
    rpr.append(colour)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), BODY_FONT)
    fonts.set(qn("w:hAnsi"), BODY_FONT)
    rpr.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(int(BODY_SIZE.pt * 2)))
    rpr.append(size)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


def style_run(run, *, bold=False, colour=BLACK, font=BODY_FONT, size=BODY_SIZE):
    run.bold = bold
    run.font.color.rgb = colour
    run.font.name = font
    run.font.size = size
    return run


def write_inline(paragraph, text: str, *, bold_all: bool = False) -> None:
    """Write text into a paragraph, honouring **bold** and [links](url)."""
    for piece in _INLINE.split(text):
        if not piece:
            continue
        link = _LINK.fullmatch(piece)
        if link:
            add_hyperlink(paragraph, link.group(2), link.group(1))
        elif piece.startswith("**") and piece.endswith("**"):
            style_run(paragraph.add_run(piece[2:-2]), bold=True)
        else:
            style_run(paragraph.add_run(piece), bold=bold_all)


def add_heading_line(doc, tag: str, text: str):
    """Body-sized paragraph: grey 'Hn:' tag, then bold black heading text."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    style_run(p.add_run(f"{tag}: "), bold=True, colour=GREY)
    write_inline(p, text, bold_all=True)
    for run in p.runs:
        run.bold = True
    return p


def add_meta_line(doc, label: str, value: str, *, link: bool = False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    style_run(p.add_run(f"{label}: "), bold=True)
    if link and value.startswith("http"):
        add_hyperlink(p, value, value)
    else:
        write_inline(p, value)
    return p


def add_blank(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    style_run(p.add_run(""))
    return p


def render_body(doc, markdown: str) -> None:
    lines = markdown.splitlines()
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            add_heading_line(doc, "H3", line[4:].strip())
        elif line.startswith("## "):
            add_heading_line(doc, "H2", line[3:].strip())
        elif line.startswith("# "):
            add_heading_line(doc, "H1", line[2:].strip())
        elif _BOLD_ONLY.match(line.strip()):
            # a line that is entirely bold == an FAQ question -> H3
            add_heading_line(doc, "H3", _BOLD_ONLY.match(line.strip()).group(1).strip())
        elif re.match(r"^\s*[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            write_inline(p, re.sub(r"^\s*[-*]\s+", "", line))
        elif re.match(r"^\s*\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(2)
            write_inline(p, re.sub(r"^\s*\d+\.\s+", "", line))
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            write_inline(p, line)


def order_entries(content_dir: Path, ids: list[str]) -> list[str]:
    hubs = [e for e in ids if not e.startswith("cleaners-")]
    spokes = sorted(e for e in ids if e.startswith("cleaners-"))
    return sorted(hubs) + spokes


def main(argv: list[str]) -> int:
    content_dir = Path(argv[1]).resolve()
    out_path = Path(argv[2]).resolve()
    ids = argv[3:] or [p.parent.name for p in content_dir.glob("content/*/generated.md")]
    ids = order_entries(content_dir, ids)

    doc = docx.Document()
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal.font.color.rgb = BLACK

    for idx, eid in enumerate(ids, start=1):
        entry = json.loads((content_dir / "entries" / f"{eid}.json").read_text())
        attrs = entry.get("attributes", {}) or {}
        body = (content_dir / "content" / eid / "generated.md").read_text()

        # "Page N" -> Heading 1 so Google Docs lists it in the outline / tabs.
        head = doc.add_heading("", level=1)
        if idx > 1:
            head.paragraph_format.page_break_before = True
        style_run(head.add_run(f"Page {idx}"), bold=True, colour=BLACK,
                  font=TITLE_FONT, size=TITLE_SIZE)
        head.paragraph_format.space_after = Pt(10)

        add_meta_line(doc, "URL", entry.get("url", ""), link=True)
        add_blank(doc)
        add_meta_line(doc, "Primary Keyword", entry.get("primary_keyword", ""))
        secondaries = entry.get("secondary_keywords", []) or []
        add_meta_line(doc, "Secondary Keywords", ", ".join(secondaries))
        add_blank(doc)
        add_meta_line(doc, "Meta Title", attrs.get("meta_title", ""))
        add_meta_line(doc, "Meta Description", attrs.get("meta_description", ""))
        add_blank(doc)

        render_body(doc, body)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"wrote {out_path}  ({len(ids)} pages)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
