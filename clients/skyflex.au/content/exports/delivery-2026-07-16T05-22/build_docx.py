import re, json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.shared import OxmlElement, qn

CLIENT = Path("/home/invoi/fahad_projects/clients/skyflex.au/content")
GREY = RGBColor(0x88, 0x88, 0x88)
LINKC = "1758C4"

ORDER = [
    ("homepage", "Homepage", "add-blocks", "pergolas melbourne", "https://skyflex.au/"),
    ("louvred-pergolas-sydney", "Louvred Pergolas Sydney (service-location)", "add-blocks", "pergolas sydney", "https://skyflex.au/louvred-pergolas-sydney/"),
    ("delta-pro-retractable-roof", "Delta Pro Retractable Roof", "rewrite-existing", "retractable roof system melbourne", "https://skyflex.au/product/delta-pro-retractable-roof/"),
    ("delta-commercial-folding-arm", "Delta Commercial Folding Arm Awning", "rewrite-existing", "retractable awning melbourne", "https://skyflex.au/product/delta-commercial-folding-arm/"),
    ("skyflex-4k-android-smart-outdoor-tv", "Skyflex 4K Outdoor TV", "rewrite-existing", "waterproof tv australia", "https://skyflex.au/product/skyflex-4k-android-smart-outdoor-tv/"),
    ("skyflex-bbq-pods", "Skyflex BBQ Pods", "rewrite-existing", "bbq pods melbourne", "https://skyflex.au/product/skyflex-bbq-pods/"),
    ("smart-toilets", "Smart Toilets (category, new page)", "new-page", "smart toilets melbourne", "https://skyflex.au/smart-toilets/"),
]
EXISTING_JSON = {"homepage": "content/homepage-existing.json",
                 "louvred-pergolas-sydney": "content/louvred-pergolas-sydney-existing.json"}

def add_hyperlink(paragraph, url, text, bold=False):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    h = OxmlElement('w:hyperlink'); h.set(qn('r:id'), r_id)
    r = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color'); c.set(qn('w:val'), LINKC); rPr.append(c)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    if bold:
        rPr.append(OxmlElement('w:b'))
    r.append(rPr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text; r.append(t)
    h.append(r); paragraph._p.append(h)

def add_inline(par, text, base_bold=False, highlight=None):
    """Render text with **bold** and [text](url) links into runs on par."""
    pos = 0
    segs = []
    for m in re.finditer(r'\[([^\]]+)\]\((https?://[^)]+)\)', text):
        if m.start() > pos: segs.append(('t', text[pos:m.start()]))
        segs.append(('a', m.group(1), m.group(2)))
        pos = m.end()
    if pos < len(text): segs.append(('t', text[pos:]))
    for s in segs:
        if s[0] == 'a':
            add_hyperlink(par, s[2], s[1], bold=base_bold)
        else:
            for p in re.split(r'(\*\*.+?\*\*)', s[1]):
                if not p: continue
                if p.startswith('**') and p.endswith('**'):
                    run = par.add_run(p[2:-2]); run.bold = True
                else:
                    run = par.add_run(p); run.bold = base_bold
                if highlight: run.font.highlight_color = highlight

def heading_para(doc, level, text, highlight=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    tag = p.add_run(f"H{level} ")
    tag.bold = True; tag.font.color.rgb = GREY
    if highlight: tag.font.highlight_color = highlight
    add_inline(p, text, base_bold=True, highlight=highlight)
    return p

def body_para(doc, text, highlight=None, bullet=False, callout=False):
    style = 'List Bullet' if bullet else None
    p = doc.add_paragraph(style=style)
    if callout:
        p.paragraph_format.space_before = Pt(4)
    add_inline(p, text, base_bold=False, highlight=highlight)
    return p

FAQ_Q = re.compile(r'^\*\*(.+?\?)\*\*\s*(.*)$')
def is_cta(bold_text):
    return ('call us now' in bold_text.lower()) or re.search(r'0[0-9] ?[0-9]{4} ?[0-9]{4}', bold_text)

def render_markdown(doc, md, highlight=None):
    lines = md.splitlines()
    i = 0
    # skip title/meta handled by caller; find body start
    while i < len(lines) and (lines[i].startswith('title:') or lines[i].startswith('meta_description:') or not lines[i].strip()):
        i += 1
    while i < len(lines):
        ln = lines[i].rstrip()
        if not ln.strip():
            i += 1; continue
        m = re.match(r'^(#{1,4})\s+(.*)', ln)
        if m:
            heading_para(doc, len(m.group(1)), m.group(2), highlight); i += 1; continue
        if ln.lstrip().startswith('- '):
            body_para(doc, ln.lstrip()[2:], highlight, bullet=True); i += 1; continue
        # FAQ question (bold ending with ?) -> H3 + optional inline answer
        fm = FAQ_Q.match(ln.strip())
        if fm and not is_cta(fm.group(1)):
            heading_para(doc, 3, fm.group(1), highlight)
            if fm.group(2).strip():
                body_para(doc, fm.group(2).strip(), highlight)
            i += 1; continue
        # CTA bold line -> bold callout (not a heading)
        if ln.strip().startswith('**') and is_cta(ln.strip()):
            body_para(doc, ln.strip(), highlight, callout=True); i += 1; continue
        body_para(doc, ln.strip(), highlight); i += 1

# ---- build ----
doc = Document()
normal = doc.styles['Normal']
normal.font.name = 'Calibri'; normal.font.size = Pt(11)
# make List Bullet same size
try: doc.styles['List Bullet'].font.size = Pt(11)
except Exception: pass

# cover
t = doc.add_paragraph(); r = t.add_run("Skyflex — Content Deliverable"); r.bold = True; r.font.size = Pt(11)
doc.add_paragraph("7 pages. Headings are tagged H1/H2/H3 and bold, at the same size as body text, so the developer can apply the correct heading level. FAQ questions are H3. Existing page content that is kept is highlighted yellow. Internal links are live hyperlinks. Generated 2026-07-16.")
sep = doc.add_paragraph(); sep.add_run("Legend: ").bold = True
sep.add_run("plain = new content to add.  ")
hr = sep.add_run("yellow = existing content on the live page, keep as-is."); hr.font.highlight_color = WD_COLOR_INDEX.YELLOW

for idx, (slug, name, mode, kw, url) in enumerate(ORDER):
    doc.add_page_break()
    ph = doc.add_paragraph(); pr = ph.add_run(f"PAGE {idx+1}: {name}"); pr.bold = True; pr.font.size = Pt(11)
    meta = doc.add_paragraph()
    meta.add_run("URL: ").bold = True; add_hyperlink(meta, url, url)
    meta.add_run(f"    Mode: {mode}    Primary keyword: ").bold = False
    kwr = meta.add_run(kw); kwr.bold = True
    # SEO title + meta description
    md = (CLIENT / f"content/{slug}/generated.md").read_text()
    lines = md.splitlines()
    title = next((l.split(':',1)[1].strip() for l in lines[:3] if l.startswith('title:') or l.lower().startswith('title_tag:')), "")
    desc = next((l.split(':',1)[1].strip() for l in lines[:3] if l.startswith('meta_description:') or l.lower().startswith('meta description:')), "")
    sp = doc.add_paragraph(); sp.add_run("SEO title: ").bold = True; sp.add_run(title)
    dp = doc.add_paragraph(); dp.add_run("Meta description: ").bold = True; dp.add_run(desc)
    doc.add_paragraph()
    if mode == "add-blocks":
        lab = doc.add_paragraph(); lr = lab.add_run("NEW CONTENT (add these sections; they lead the page):"); lr.bold = True
        render_markdown(doc, md)
        # existing content, yellow, in position (after new sections)
        ex = json.loads((CLIENT / EXISTING_JSON[slug]).read_text())
        doc.add_paragraph()
        el = doc.add_paragraph(); er = el.add_run("EXISTING PAGE CONTENT BELOW — keep as-is (highlighted yellow):"); er.bold = True; er.font.highlight_color = WD_COLOR_INDEX.YELLOW
        n2 = doc.add_paragraph(); n2r = n2.add_run("Note: the new sections above supply the page H1. Demote the existing H1 below to H2 so the page has a single H1. On the Sydney page also fix the 'designs and installs' intro, the skyflex.com.au link and the Melbourne phone; rename any 'Why Choose SkyFlex' heading; fill or remove empty U6/U7 grid tiles."); n2r.italic = True
        st = doc.add_paragraph(); st.add_run("Current SEO title: ").bold = True
        stt = st.add_run(ex.get("title","")); stt.font.highlight_color = WD_COLOR_INDEX.YELLOW
        st.add_run("  (replace with the new SEO title above)")
        for h in ex.get("headings", []):
            tg = h.get("tag","H2"); tx = h.get("text","").strip()
            if not tx or 'cart' in tx.lower(): continue
            lvl = tg[1] if tg.startswith('H') and tg[1:].isdigit() else '2'
            heading_para(doc, lvl, tx, highlight=WD_COLOR_INDEX.YELLOW)
        paras = [p for p in ex.get("paragraphs", []) if len(p) > 50 and 'reCAPTCHA' not in p and '$0.00' not in p][:5]
        if paras:
            note = doc.add_paragraph(); note.add_run("Sample of existing body copy (kept):").bold = True
            for p in paras: body_para(doc, p, highlight=WD_COLOR_INDEX.YELLOW)
    else:
        if mode == "rewrite-existing":
            note = doc.add_paragraph(); nr = note.add_run("Rewrite: the content below replaces the current page body."); nr.italic = True
        render_markdown(doc, md)

outdir = CLIENT / "exports" / "delivery-2026-07-16T05-22"
out = outdir / "skyflex-content-deliverable.docx"
doc.save(str(out))
print("WROTE", out)
