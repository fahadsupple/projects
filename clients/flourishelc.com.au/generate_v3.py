from docx import Document
from docx.shared import Pt

doc = Document()

# ─────────────────────────────────────────────
# CONTENT VARIANTS
# ─────────────────────────────────────────────

# 4 versions of the Programs H2 section
programs = {
    'A': {
        'heading': 'Programs That Cover Every Part of Childhood',
        'body': (
            "STEM Learning, Creative Arts, Language, Mathematics, Fine Motor Skills, Music and Movement, "
            "Cultural Understanding, Social Justice, Independence, Exercise, Health, and Social and Emotional "
            "development are all woven into the daily rhythm at Flourish ELC. These are not separate subjects "
            "delivered in isolation: they come to life through play, projects, conversations, and discovery "
            "that children find genuinely exciting.\n\n"
            "Academic foundations matter. Language and Mathematics give children the skills they need when "
            "formal schooling begins. Fine Motor Skills build the physical dexterity that underpins writing, "
            "drawing, and creative tasks. STEM Learning sparks the curiosity that keeps children engaged with "
            "ideas long after they leave our centre. Creative Arts and Music and Movement give children "
            "expressive outlets that build confidence and a love of learning that stays with them for life.\n\n"
            "Five nutritious meals are prepared and served each weekday: breakfast, morning tea, lunch, "
            "afternoon tea, and a late snack. School readiness is woven through everything we do, so children "
            "move into their next chapter with the skills, the confidence, and the enthusiasm they need."
        )
    },
    'B': {
        'heading': 'Learning That Develops the Whole Child',
        'body': (
            "Play is how children learn best, and at Flourish ELC, learning through play is not a compromise: "
            "it is the method. STEM Learning, Creative Arts, Language, Mathematics, Fine Motor Skills, Music "
            "and Movement, Cultural Understanding, Social Justice, Independence, Exercise, Health, and Social "
            "and Emotional development are all delivered through experiences that children find joyful and "
            "meaningful.\n\n"
            "Social and Emotional development sits at the heart of our curriculum because children who feel "
            "secure and connected learn better and grow more confidently. Independence is nurtured from an "
            "early age, with educators standing alongside children as they build judgment and self-belief. "
            "Cultural Understanding and Social Justice are woven into everyday experiences because we want "
            "children to grow into people who are curious about the world and kind within it.\n\n"
            "Five nutritious meals are served each weekday: breakfast, morning tea, lunch, afternoon tea, and "
            "a late snack. Our school readiness programs make the transition to kindergarten a confident and "
            "joyful one rather than an uncertain one."
        )
    },
    'C': {
        'heading': 'A Curriculum Built for Curious, Capable Children',
        'body': (
            "The Flourish ELC curriculum is broad, intentional, and full of life. Every day, children engage "
            "with Creative Arts, STEM Learning, Language, Mathematics, Fine Motor Skills, Music and Movement, "
            "Cultural Understanding, Social Justice, Independence, Exercise, Health, and Social and Emotional "
            "development through experiences that feel like adventure rather than instruction.\n\n"
            "Mathematics and Language build the academic foundations children carry forward into school. STEM "
            "Learning cultivates the inventive thinking that serves children across every discipline. Creative "
            "Arts and Music and Movement give children ways to express themselves that go far beyond words. "
            "Cultural Understanding and Social Justice ensure that children develop empathy and a broad view "
            "of the world from the very beginning. Exercise and Health are part of every day, so children "
            "move and grow physically alongside everything else they are discovering.\n\n"
            "Meals are nutritious and plentiful: breakfast, morning tea, lunch, afternoon tea, and a late "
            "snack every weekday. Our school readiness programs are embedded throughout so that children are "
            "genuinely prepared when kindergarten arrives."
        )
    },
    'D': {
        'heading': 'Rich Programs Across Every Learning Area',
        'body': (
            "At Flourish ELC, no dimension of a child's development is left to chance. Our daily programs "
            "cover Creative Arts, STEM Learning, Language, Mathematics, Fine Motor Skills, Music and Movement, "
            "Cultural Understanding, Social Justice, Independence, Exercise, Health, and Social and Emotional "
            "development, not as a checklist, but as a living curriculum that responds to what children are "
            "curious about and what they need next.\n\n"
            "Fine Motor Skills develop through hands-on creative activities that make the foundational work "
            "feel like play. Language and Mathematics are integrated naturally into the daily flow of life at "
            "the centre. Exercise and Health are embedded in outdoor play and movement experiences every day. "
            "Independence is nurtured gently, because children who can manage themselves feel capable and "
            "proud, and that confidence carries forward into everything they do.\n\n"
            "Five nutritious meals are provided each weekday: breakfast, morning tea, lunch, afternoon tea, "
            "and a late snack. School readiness programs ensure that every child is ready to embrace "
            "kindergarten with genuine excitement and a strong foundation."
        )
    }
}

# 4 versions of the Environment H2 section
environment = {
    'A': {
        'heading': 'Purpose-Built, NQF-Aligned, and Always Welcoming',
        'body': (
            "Our facility is purpose-built for children aged 6 weeks to school age. Every room, every outdoor "
            "space, and every learning corner has been thoughtfully designed with children's safety, wellbeing, "
            "and development in mind: not repurposed from something else, but built from the ground up for the "
            "work of early learning. We meet and exceed the National Quality Framework (NQF), which means the "
            "quality of our educators, our programs, and our environments is assessed against Australia's "
            "highest national standards for early education and care.\n\n"
            "Our educators are qualified professionals who bring both formal training and genuine warmth to "
            "every interaction. They take the time to know each child as an individual: what makes them "
            "curious, what makes them feel safe, and what they are ready to explore next. That personal "
            "attention makes a real difference to children and families alike.\n\n"
            "We are open Monday to Friday, 7:00am to 6:00pm. Our parent communication app delivers real-time "
            "updates, photos, and observations throughout the day so families always feel genuinely connected "
            "to their child's world, even on the busiest days."
        )
    },
    'B': {
        'heading': 'Qualified Educators in a Space Designed for Children',
        'body': (
            "Flourish ELC is purpose-built. Every space, from indoor learning rooms to outdoor play areas, "
            "has been designed for the specific needs of young children at different developmental stages. The "
            "environment itself is a teaching tool, and we have built ours with care and intention. Our "
            "qualified educators meet and exceed the National Quality Framework (NQF), and their dedication "
            "shows in the way they engage with children: attentively, warmly, and with genuine professional "
            "skill.\n\n"
            "The NQF is not just a compliance standard at Flourish ELC: it is a reflection of our values. "
            "Every element of the centre, the spaces, the programs, the educator qualifications, and the "
            "relationships we build with families, has been assessed against Australia's national quality "
            "standards. We hold ourselves to those standards because they matter to us, not just because "
            "they are required.\n\n"
            "We are open 7:00am to 6:00pm Monday to Friday, and our parent communication app keeps families "
            "genuinely connected throughout the day with real-time updates, photos, and notes from our "
            "educators."
        )
    },
    'C': {
        'heading': 'A Safe, Warm Environment That Meets the NQF',
        'body': (
            "The Flourish ELC environment has been purpose-built for children from 6 weeks to school age. "
            "Safe, bright, and welcoming, it is a place where children feel at home from the moment they "
            "arrive. Meeting and exceeding the National Quality Framework (NQF) is not a target we reach for: "
            "it is the standard we live every single day. Our educators hold recognised qualifications and "
            "bring a passion for early childhood to every interaction.\n\n"
            "Children spend their days in spaces designed to inspire movement, creativity, and connection. "
            "Outdoor areas give children room to run, explore, and experience nature as part of their "
            "learning. Indoor environments are warm and purposeful, with activities that reflect the interests "
            "and developmental needs of the children who use them. The result is a place that children love "
            "to come to and families feel proud to be part of.\n\n"
            "We are open Monday to Friday, 7:00am to 6:00pm. Our parent communication app sends real-time "
            "updates and photos throughout the day, giving families a genuine window into what their child "
            "has been experiencing and achieving."
        )
    },
    'D': {
        'heading': 'An Environment Built Around Children, Open Every Weekday',
        'body': (
            "Everything about the Flourish ELC facility has been built with children in mind. Our purpose-built "
            "centre meets and exceeds the National Quality Framework (NQF), giving families confidence that "
            "the care, the people, and the programs their children experience have been assessed against "
            "Australia's highest national standards for early education. That standard is one we are genuinely "
            "proud of and work hard to uphold every day.\n\n"
            "Our educators are at the heart of everything. They are qualified, warm, and genuinely committed "
            "to the children in their care. They observe carefully, listen closely, and respond thoughtfully, "
            "creating a daily experience that is both nurturing and stimulating. That quality of attention "
            "is what families notice and what children thrive on.\n\n"
            "We are open 7:00am to 6:00pm Monday to Friday, offering families a flexible window for drop-off "
            "and pick-up. Our parent communication app delivers real-time photos and updates throughout the "
            "day, keeping families connected even when life gets busy."
        )
    }
}

# 4 versions of the Why Choose Us H2 content (heading is generated per page using keyword)
why_choose = {
    'A': (
        "Flourish ELC is guided by the Early Years Learning Framework (EYLF): Belonging, Being and Becoming. "
        "That framework is not just words on a policy document: it is the living philosophy behind how our "
        "educators plan, observe, and respond to each child every single day. Families choose us because they "
        "can feel the difference it makes, from the way children are greeted each morning to the way "
        "milestones are noticed and celebrated.\n\n"
        "Five nutritious meals are served each weekday: breakfast, morning tea, lunch, afternoon tea, and a "
        "late snack. Dietary requirements and allergies are always carefully managed. Our centre is open "
        "Monday to Friday, 7:00am to 6:00pm, giving families a generous and flexible window. Our parent "
        "communication app delivers real-time updates and photos throughout the day so you always know what "
        "your child has been exploring, creating, and enjoying.\n\n"
        "Flourish ELC is a community, not just a centre. Families feel genuinely welcomed here, and children "
        "feel genuinely known. That sense of belonging is at the core of everything we offer, and it is what "
        "families consistently tell us makes Flourish ELC the right choice for their child."
    ),
    'B': (
        "Choosing early learning is one of the most significant decisions a family makes, and we do not take "
        "that trust lightly. Flourish ELC is guided by the Early Years Learning Framework (EYLF): Belonging, "
        "Being and Becoming, which means our programs are built around your child's identity, wellbeing, and "
        "sense of belonging, not just their academic preparation.\n\n"
        "Our educators hold recognised qualifications and meet and exceed the National Quality Framework (NQF). "
        "They bring both professional expertise and genuine warmth to every session. Our school readiness "
        "programs ensure children transition to kindergarten with the skills, the confidence, and the love "
        "of learning already well established. Nothing is left to chance in that preparation.\n\n"
        "We are open Monday to Friday, 7:00am to 6:00pm, with five nutritious daily meals provided and a "
        "parent communication app that keeps you connected in real time. Everything families need is here, "
        "including the peace of mind that their child is thriving."
    ),
    'C': (
        "Flourish ELC stands apart because every decision we make, from how we design our spaces to how we "
        "hire and support our educators, is guided by what is genuinely best for children. Our purpose-built "
        "facility meets and exceeds the National Quality Framework (NQF). Our programs are shaped by the "
        "Early Years Learning Framework (EYLF): Belonging, Being and Becoming. And our educators bring both "
        "formal qualifications and genuine heart to every day.\n\n"
        "Families who choose Flourish ELC find a warm, inclusive community where children are celebrated for "
        "who they are and gently supported in who they are becoming. School readiness is woven through "
        "everything we do, so the transition to kindergarten is confident and smooth rather than uncertain "
        "or rushed.\n\n"
        "Five nutritious meals are served each weekday: breakfast, morning tea, lunch, afternoon tea, and a "
        "late snack. We are open Monday to Friday, 7:00am to 6:00pm. Our parent communication app keeps you "
        "genuinely connected to your child's day throughout the hours you are apart."
    ),
    'D': (
        "Flourish ELC is a purpose-built early learning centre guided by the Early Years Learning Framework "
        "(EYLF): Belonging, Being and Becoming. Our qualified educators meet and exceed the National Quality "
        "Framework (NQF), and our programs are designed around the whole child rather than a narrow focus "
        "on academic preparation alone.\n\n"
        "The details matter to us. Five nutritious meals are prepared fresh every weekday: breakfast, morning "
        "tea, lunch, afternoon tea, and a late snack. Dietary requirements and allergies are always managed "
        "carefully. Our centre is open 7:00am to 6:00pm Monday to Friday, and our parent communication app "
        "delivers real-time photos and updates so you never feel far from your child's world.\n\n"
        "What you will find at Flourish ELC is a centre where children love to be, where educators genuinely "
        "know each child, and where families feel like true partners in their child's learning and growth. "
        "That is what we offer, and it is what sets us apart."
    )
}

# 4 FAQ sets per page type
faqs_childcare = {
    'A': [
        ("What age groups does Flourish ELC cater to?",
         "We welcome children from 6 weeks through to school age, with dedicated programs and spaces for each developmental stage."),
        ("What meals are provided each day?",
         "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. All dietary requirements and allergies are carefully managed."),
        ("What are your opening hours?",
         "Our centre is open Monday to Friday, 7:00am to 6:00pm."),
        ("How does the parent communication app work?",
         "The app delivers real-time updates, photos, and observations throughout the day so families always feel genuinely connected to their child's world."),
        ("What is the Early Years Learning Framework?",
         "The EYLF (Belonging, Being and Becoming) is the national framework guiding how our educators plan and deliver learning experiences, always placing your child's identity, wellbeing, and engagement at the centre of practice."),
    ],
    'B': [
        ("How old does my child need to be to enrol at Flourish ELC?",
         "We welcome children from 6 weeks of age through to school age. Each age group has its own dedicated program and learning environment."),
        ("What meals does Flourish ELC serve?",
         "Five nutritious meals daily: breakfast, morning tea, lunch, afternoon tea, and a late snack. Dietary needs and allergies are always accommodated."),
        ("What are the centre's opening hours?",
         "We are open Monday to Friday from 7:00am to 6:00pm, giving families a generous window for drop-off and pick-up."),
        ("How do I stay connected to my child's day?",
         "Our parent communication app sends real-time photos and updates throughout the day so you are always close to what your child is experiencing."),
        ("What does the NQF mean for my child?",
         "The National Quality Framework sets and assesses quality standards for early education in Australia. Meeting and exceeding the NQF means our programs, educators, and environments are held to the highest national benchmarks."),
    ],
    'C': [
        ("Does Flourish ELC cater for very young babies?",
         "Yes, we welcome children from 6 weeks of age through to school age, with age-appropriate programs and dedicated spaces for each stage."),
        ("Are meals included?",
         "Five nutritious meals are provided every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. All dietary requirements and allergies are managed carefully."),
        ("When is the centre open?",
         "Monday to Friday, 7:00am to 6:00pm."),
        ("How will I know what my child has been doing each day?",
         "Our parent communication app delivers real-time updates, photos, and notes from educators so you always have a genuine window into your child's day."),
        ("What school readiness programs does Flourish ELC offer?",
         "Our school readiness programs develop the academic, social, and emotional skills children need to transition confidently into kindergarten and feel genuinely ready from day one."),
    ],
    'D': [
        ("What age can children start at Flourish ELC?",
         "From 6 weeks of age. We cater for children all the way through to school age, with programs and spaces tailored to each developmental stage."),
        ("What food is provided at Flourish ELC?",
         "Breakfast, morning tea, lunch, afternoon tea, and a late snack: five nutritious meals every weekday. Dietary requirements and allergies are always carefully catered for."),
        ("What are your hours of operation?",
         "We are open 7:00am to 6:00pm, Monday to Friday."),
        ("How does Flourish ELC keep families informed during the day?",
         "Our parent communication app sends real-time photos, updates, and observations so families always feel connected and informed about their child's day."),
        ("How does the EYLF guide what happens at Flourish ELC?",
         "The Early Years Learning Framework shapes how our educators plan experiences, observe children, and support their development, always with belonging, wellbeing, and engagement at the heart of everything."),
    ]
}

faqs_preschool = {
    'A': [
        ("What age is preschool at Flourish ELC?",
         "Our preschool programs are designed for children aged 3 to 5. Our full centre welcomes children from 6 weeks to school age."),
        ("What meals are included in the preschool program?",
         "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. All dietary requirements and allergies are carefully accommodated."),
        ("What are the preschool opening hours?",
         "Monday to Friday, 7:00am to 6:00pm."),
        ("How do educators share updates with families?",
         "Our parent communication app delivers real-time photos and observations throughout the day so you always know what your child has been exploring and learning."),
        ("How does Flourish ELC prepare preschoolers for kindergarten?",
         "Our school readiness programs develop literacy, numeracy, independence, and social and emotional skills progressively so children transition to kindergarten with confidence and genuine readiness."),
    ],
    'B': [
        ("What age does the preschool program cater for?",
         "Preschool programs are designed for children aged 3 to 5. Our centre welcomes all children from 6 weeks through to school age."),
        ("What food does Flourish ELC provide for preschoolers?",
         "Five nutritious meals each weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. Dietary needs and allergies are always managed."),
        ("What are the opening hours for preschool?",
         "Our centre is open Monday to Friday, 7:00am to 6:00pm."),
        ("How will I know what my preschooler has been up to during the day?",
         "Our parent communication app sends real-time updates, photos, and observations so you always feel connected to your child's preschool day."),
        ("What does school readiness look like at Flourish ELC?",
         "We build literacy, numeracy, independence, and social and emotional skills throughout the preschool years so children step into kindergarten ready, confident, and excited."),
    ],
    'C': [
        ("What is the age range for preschool at Flourish ELC?",
         "3 to 5 years. Our full centre caters for children from 6 weeks right through to school age."),
        ("Are meals included in the preschool program?",
         "Yes, five nutritious meals are served every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. All allergies and dietary requirements are managed."),
        ("How long is the preschool day?",
         "Our centre is open Monday to Friday, 7:00am to 6:00pm."),
        ("How does Flourish ELC keep preschool families in the loop?",
         "Our parent communication app sends real-time photos and updates throughout the day, keeping you genuinely connected to your child's preschool experience."),
        ("How does the EYLF guide the preschool program?",
         "The Early Years Learning Framework shapes how our educators plan experiences and support each child's unique journey, placing belonging, wellbeing, and curiosity at the heart of every preschool day."),
    ],
    'D': [
        ("What age is the preschool program for?",
         "Children aged 3 to 5. Our centre also welcomes children from 6 weeks to school age across all programs."),
        ("What meals are served during preschool?",
         "Breakfast, morning tea, lunch, afternoon tea, and a late snack: five nutritious meals prepared fresh every weekday."),
        ("What are the preschool hours at Flourish ELC?",
         "Monday to Friday, 7:00am to 6:00pm."),
        ("How are families kept updated during the preschool day?",
         "Our parent communication app delivers real-time photos and observations so you are always aware of what your child has been learning and experiencing."),
        ("What school readiness support does Flourish ELC offer in preschool?",
         "School readiness is embedded in our preschool programs from the outset, building literacy, numeracy, social skills, and independence so children are genuinely prepared for kindergarten."),
    ]
}

faqs_kinder = {
    'A': [
        ("What age is kindergarten at Flourish ELC?",
         "Our kindergarten program is designed for children aged 4 to 5, the year before they begin primary school."),
        ("What meals are included during kindergarten?",
         "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack."),
        ("What are the kindergarten opening hours?",
         "Monday to Friday, 7:00am to 6:00pm."),
        ("How does the parent app work during the kindergarten year?",
         "It delivers real-time updates and photos throughout the day so families always feel connected to their child's kindergarten experience."),
        ("How does Flourish ELC prepare kindergarteners for primary school?",
         "Our school readiness programs develop literacy, numeracy, independence, and social and emotional skills throughout the kindergarten year so children begin primary school genuinely ready and excited."),
    ],
    'B': [
        ("What age is the kindergarten program?",
         "Designed for children aged 4 to 5 in the year before primary school begins."),
        ("What food is provided during kindergarten?",
         "Five nutritious meals each weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack."),
        ("What are your kindergarten hours?",
         "Our centre is open Monday to Friday, 7:00am to 6:00pm."),
        ("How will I know what my kindergartener has been doing during the day?",
         "Our parent communication app sends real-time photos and updates so you always have a clear picture of your child's kindergarten day."),
        ("What does school readiness mean at Flourish ELC?",
         "We build literacy, numeracy, independence, and social and emotional skills throughout the kindergarten year so children walk into primary school confident, capable, and genuinely excited."),
    ],
    'C': [
        ("At what age do children join kindergarten at Flourish ELC?",
         "Kindergarten is for children aged 4 to 5, the year before primary school. Our full centre welcomes children from 6 weeks to school age."),
        ("Are meals included in kindergarten?",
         "Yes, five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. Dietary needs are always managed."),
        ("What are the hours for kindergarten?",
         "Monday to Friday, 7:00am to 6:00pm."),
        ("How does Flourish ELC keep kindergarten families connected?",
         "Our parent communication app delivers real-time updates and photos throughout the kindergarten day so families always know what their child has been up to."),
        ("How does the EYLF shape the kindergarten program?",
         "The Early Years Learning Framework guides how our educators plan and deliver kindergarten experiences, always placing each child's belonging, wellbeing, and engagement at the heart of practice."),
    ],
    'D': [
        ("What is the kindergarten age group at Flourish ELC?",
         "4 to 5 years, in the year before primary school begins."),
        ("What meals are served in kindergarten?",
         "Breakfast, morning tea, lunch, afternoon tea, and a late snack: five nutritious meals prepared fresh every weekday."),
        ("When is the kindergarten program available?",
         "Monday to Friday, 7:00am to 6:00pm."),
        ("How are kindergarten families updated throughout the day?",
         "Our parent communication app sends real-time photos and observations, keeping you close to your child's kindergarten world even on busy days."),
        ("What makes Flourish ELC kindergarten different?",
         "Our EYLF-guided programs, qualified educators, purpose-built spaces, and embedded school readiness focus combine to give every child the best possible start to their formal education journey."),
    ]
}

# ─────────────────────────────────────────────
# PAGE DATA
# ─────────────────────────────────────────────
# var: 0=Programs→Env→Why, 1=Env→Why→Programs, 2=Why→Programs→Env, 3=Programs→Why→Env

pages = [
    # ── NSW Childcare ──
    {
        "num": 1, "type": "childcare",
        "url": "https://flourishelc.com.au/daycare-childcare-oakville-nsw/",
        "primary": "childcare oakville", "secondary": "daycare oakville",
        "meta_title": "Childcare Oakville | Flourish Early Learning Centre",
        "meta_desc": "Warm, purpose-built childcare in Oakville with qualified educators, five daily meals, and programs guided by the EYLF. Book your tour today.",
        "h1": "Childcare Oakville Families Love: Flourish Early Learning Centre",
        "intro": (
            "Flourish Early Learning Centre is a purpose-built childcare in Oakville designed around children "
            "thriving in every sense of the word. Guided by the Early Years Learning Framework (EYLF): "
            "Belonging, Being and Becoming, our qualified educators bring genuine passion to every single day. "
            "Oakville families choose us because we go beyond care: we create a place where children light up, "
            "grow, and flourish."
        ),
        "why_heading": "Why Choose Flourish ELC for Childcare in Oakville",
        "prog": 'A', "env": 'A', "why": 'A', "faq": 'childcare_A', "var": 0,
    },
    {
        "num": 2, "type": "childcare",
        "url": "https://flourishelc.com.au/daycare-childcare-box-hill-nsw/",
        "primary": "childcare near box hill", "secondary": "daycare near box hill",
        "meta_title": "Childcare Near Box Hill NSW | Flourish Early Learning Centre",
        "meta_desc": "Searching for childcare near Box Hill? Flourish ELC is a purpose-built centre with qualified educators, five daily meals, and EYLF-guided programs. Book your tour today.",
        "h1": "Childcare Near Box Hill: Where Every Child Belongs, Grows and Flourishes",
        "intro": (
            "Flourish Early Learning Centre is a warm, purpose-built early learning centre conveniently located "
            "for families in and around Box Hill. Guided by the Early Years Learning Framework (EYLF): "
            "Belonging, Being and Becoming, our qualified, passionate educators bring that philosophy to life "
            "every single day. If you are looking for childcare near Box Hill that feels like a second home "
            "for your child, you have found it."
        ),
        "why_heading": "Why Choose Flourish ELC for Childcare Near Box Hill",
        "prog": 'B', "env": 'B', "why": 'B', "faq": 'childcare_B', "var": 1,
    },
    {
        "num": 3, "type": "childcare",
        "url": "https://flourishelc.com.au/daycare-childcare-vineyard-nsw/",
        "primary": "childcare near vineyard", "secondary": "daycare near vineyard",
        "meta_title": "Childcare Near Vineyard NSW | Flourish Early Learning Centre",
        "meta_desc": "Looking for childcare near Vineyard? Flourish ELC is purpose-built, guided by the EYLF, and offers five daily meals and qualified educators. Book your tour today.",
        "h1": "Childcare Near Vineyard That Puts Children First",
        "intro": (
            "Flourish Early Learning Centre is a purpose-built early learning centre serving families in and "
            "around Vineyard. Built around the Early Years Learning Framework (EYLF): Belonging, Being and "
            "Becoming, our centre is a place where children are celebrated for who they are, supported in "
            "what they are becoming, and given every opportunity to belong. Our qualified, passionate educators "
            "make Flourish ELC a truly special place for Vineyard families."
        ),
        "why_heading": "Why Choose Flourish ELC for Childcare Near Vineyard",
        "prog": 'C', "env": 'C', "why": 'C', "faq": 'childcare_C', "var": 2,
    },
    {
        "num": 4, "type": "childcare",
        "url": "https://flourishelc.com.au/daycare-childcare-gables-nsw/",
        "primary": "childcare near gables", "secondary": "daycare near gables",
        "meta_title": "Childcare Near Gables NSW | Flourish Early Learning Centre",
        "meta_desc": "Warm, purpose-built childcare near Gables with EYLF-guided programs, qualified educators, and five nutritious daily meals. Book your family tour today.",
        "h1": "Childcare Near Gables: A Place Where Children Truly Flourish",
        "intro": (
            "Flourish Early Learning Centre is a purpose-built childcare centre welcoming families in and "
            "around Gables. Everything here, from the way our educators engage with children to the way our "
            "spaces are designed, is shaped by the Early Years Learning Framework (EYLF): Belonging, Being "
            "and Becoming. Children at Flourish ELC do not just attend a centre: they belong to a warm, "
            "vibrant community that celebrates every step of their journey."
        ),
        "why_heading": "Why Choose Flourish ELC for Childcare Near Gables",
        "prog": 'D', "env": 'D', "why": 'D', "faq": 'childcare_D', "var": 3,
    },
    {
        "num": 5, "type": "childcare",
        "url": "https://flourishelc.com.au/daycare-childcare-grantham-farm-nsw/",
        "primary": "childcare near grantham farm", "secondary": "daycare near grantham farm",
        "meta_title": "Childcare Near Grantham Farm | Flourish Early Learning Centre",
        "meta_desc": "Purpose-built childcare near Grantham Farm with EYLF-guided learning, qualified educators, five daily meals, and a real-time parent app. Book your tour today.",
        "h1": "Childcare Near Grantham Farm: Warm, Qualified, and Built for Growing Families",
        "intro": (
            "Families near Grantham Farm looking for childcare that goes beyond the basics will find exactly "
            "that at Flourish Early Learning Centre. Our purpose-built centre is guided by the Early Years "
            "Learning Framework (EYLF): Belonging, Being and Becoming, and staffed by qualified educators who "
            "genuinely love what they do. At Flourish ELC, children are known, celebrated, and given every "
            "opportunity to thrive."
        ),
        "why_heading": "Why Choose Flourish ELC for Childcare Near Grantham Farm",
        "prog": 'A', "env": 'B', "why": 'A', "faq": 'childcare_A', "var": 0,
    },
    {
        "num": 6, "type": "childcare",
        "url": "https://flourishelc.com.au/daycare-childcare-riverstone-nsw/",
        "primary": "childcare near riverstone", "secondary": "daycare near riverstone",
        "meta_title": "Childcare Near Riverstone NSW | Flourish Early Learning Centre",
        "meta_desc": "Warm, purpose-built childcare near Riverstone with EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour.",
        "h1": "Childcare Near Riverstone: Belonging, Growing, and Flourishing Together",
        "intro": (
            "Flourish Early Learning Centre offers families near Riverstone a warm, purpose-built early "
            "learning environment where children are known by name, celebrated for who they are, and gently "
            "guided toward all they are becoming. Our programs are shaped by the Early Years Learning "
            "Framework (EYLF): Belonging, Being and Becoming, and delivered by qualified educators who bring "
            "both skill and heart to every day."
        ),
        "why_heading": "Why Choose Flourish ELC for Childcare Near Riverstone",
        "prog": 'B', "env": 'C', "why": 'B', "faq": 'childcare_B', "var": 1,
    },
    {
        "num": 7, "type": "childcare",
        "url": "https://flourishelc.com.au/daycare-childcare-mulgrave-nsw/",
        "primary": "childcare near mulgrave", "secondary": "daycare near mulgrave",
        "meta_title": "Childcare Near Mulgrave NSW | Flourish Early Learning Centre",
        "meta_desc": "Looking for childcare near Mulgrave? Flourish ELC is purpose-built with EYLF-guided programs, five nutritious daily meals, and qualified educators. Book your tour.",
        "h1": "Childcare Near Mulgrave That Families Trust and Children Love",
        "intro": (
            "Flourish Early Learning Centre is a purpose-built early learning centre welcoming families near "
            "Mulgrave into a warm, vibrant community. Guided by the Early Years Learning Framework (EYLF): "
            "Belonging, Being and Becoming, our centre is a place where children feel secure, excited, and "
            "free to explore who they are. Our qualified, passionate educators are the heart of everything "
            "we do."
        ),
        "why_heading": "Why Choose Flourish ELC for Childcare Near Mulgrave",
        "prog": 'C', "env": 'D', "why": 'C', "faq": 'childcare_C', "var": 2,
    },
    # ── NSW Preschool ──
    {
        "num": 8, "type": "preschool",
        "url": "https://flourishelc.com.au/preschool-oakville-nsw/",
        "primary": "preschool oakville", "secondary": "",
        "meta_title": "Preschool Oakville NSW | Flourish Early Learning Centre",
        "meta_desc": "Looking for a preschool in Oakville? Flourish ELC offers EYLF-guided programs, qualified educators, five daily meals, and a warm, purpose-built environment. Enquire today.",
        "h1": "Preschool Oakville Families Choose: Flourish Early Learning Centre",
        "intro": (
            "Flourish Early Learning Centre is a warm, purpose-built preschool in Oakville where children "
            "step into a world designed entirely for them. Our preschool programs are guided by the Early "
            "Years Learning Framework (EYLF): Belonging, Being and Becoming, and delivered by qualified "
            "educators who genuinely care about every child's growth. Families in Oakville choose Flourish "
            "ELC because it is the kind of place that children look forward to, and parents feel proud of."
        ),
        "why_heading": "Why Choose Flourish ELC for Preschool in Oakville",
        "prog": 'D', "env": 'A', "why": 'D', "faq": 'preschool_A', "var": 3,
    },
    {
        "num": 9, "type": "preschool",
        "url": "https://flourishelc.com.au/preschool-box-hill-nsw/",
        "primary": "preschool near box hill", "secondary": "",
        "meta_title": "Preschool Near Box Hill NSW | Flourish Early Learning Centre",
        "meta_desc": "Searching for a preschool near Box Hill? Flourish ELC is purpose-built with qualified educators, EYLF-guided programs, and five daily meals. Book your tour today.",
        "h1": "Preschool Near Box Hill: Where Children Grow Into Themselves",
        "intro": (
            "Flourish Early Learning Centre is a purpose-built preschool near Box Hill where the Early Years "
            "Learning Framework (EYLF): Belonging, Being and Becoming, is not just a document on a wall but "
            "the living philosophy behind every experience. Our qualified, passionate educators know each "
            "child by name, by interest, and by strength. Families near Box Hill choose Flourish ELC because "
            "their children flourish here in every sense of the word."
        ),
        "why_heading": "Why Choose Flourish ELC for Preschool Near Box Hill",
        "prog": 'A', "env": 'B', "why": 'A', "faq": 'preschool_B', "var": 0,
    },
    {
        "num": 10, "type": "preschool",
        "url": "https://flourishelc.com.au/preschool-vineyard-nsw/",
        "primary": "preschool near vineyard", "secondary": "",
        "meta_title": "Preschool Near Vineyard NSW | Flourish Early Learning Centre",
        "meta_desc": "Looking for a preschool near Vineyard? Flourish ELC is purpose-built with EYLF-guided programs, qualified educators, and five daily nutritious meals. Enquire today.",
        "h1": "Preschool Near Vineyard: A Warm Place Where Learning Comes Alive",
        "intro": (
            "Flourish Early Learning Centre is a purpose-built preschool near Vineyard where the Early Years "
            "Learning Framework (EYLF): Belonging, Being and Becoming, shapes every experience, every "
            "relationship, and every room. Our qualified educators bring knowledge and genuine warmth to the "
            "preschool years, creating an environment where children discover their strengths, build their "
            "confidence, and develop a love of learning that will stay with them for life."
        ),
        "why_heading": "Why Choose Flourish ELC for Preschool Near Vineyard",
        "prog": 'B', "env": 'C', "why": 'B', "faq": 'preschool_C', "var": 1,
    },
    {
        "num": 11, "type": "preschool",
        "url": "https://flourishelc.com.au/preschool-gables-nsw/",
        "primary": "preschool near gables", "secondary": "",
        "meta_title": "Preschool Near Gables NSW | Flourish Early Learning Centre",
        "meta_desc": "Looking for a preschool near Gables? Flourish ELC is purpose-built with EYLF-guided learning, qualified educators, and five daily meals. Book your tour today.",
        "h1": "Preschool Near Gables: Where Little Learners Truly Belong",
        "intro": (
            "Flourish Early Learning Centre is a warm, purpose-built preschool near Gables where the Early "
            "Years Learning Framework (EYLF): Belonging, Being and Becoming, is more than a framework: it is "
            "the spirit of everything we do. Our qualified, passionate educators create an environment where "
            "preschoolers feel safe, celebrated, and inspired to explore the world around them. Families near "
            "Gables choose Flourish ELC because their children thrive here, and they can see it every day."
        ),
        "why_heading": "Why Choose Flourish ELC for Preschool Near Gables",
        "prog": 'C', "env": 'D', "why": 'C', "faq": 'preschool_D', "var": 2,
    },
    {
        "num": 12, "type": "preschool",
        "url": "https://flourishelc.com.au/preschool-grantham-farm-nsw/",
        "primary": "preschool near grantham farm", "secondary": "",
        "meta_title": "Preschool Near Grantham Farm | Flourish Early Learning Centre",
        "meta_desc": "Looking for a preschool near Grantham Farm? Flourish ELC is purpose-built with EYLF-guided programs, five daily meals, and qualified educators. Book your tour.",
        "h1": "Preschool Near Grantham Farm: Where Every Child Gets the Best Start",
        "intro": (
            "Flourish Early Learning Centre is a purpose-built preschool near Grantham Farm where the Early "
            "Years Learning Framework (EYLF): Belonging, Being and Becoming, guides every experience. Our "
            "qualified educators are passionate about early childhood and bring that passion to life in "
            "programs that are joyful, intentional, and genuinely child-centred. Families near Grantham Farm "
            "trust Flourish ELC because it is the kind of preschool that makes a real difference."
        ),
        "why_heading": "Why Choose Flourish ELC for Preschool Near Grantham Farm",
        "prog": 'D', "env": 'A', "why": 'D', "faq": 'preschool_A', "var": 3,
    },
    {
        "num": 13, "type": "preschool",
        "url": "https://flourishelc.com.au/preschool-riverstone-nsw/",
        "primary": "preschool near riverstone", "secondary": "",
        "meta_title": "Preschool Near Riverstone NSW | Flourish Early Learning Centre",
        "meta_desc": "Looking for a preschool near Riverstone? Flourish ELC is purpose-built with EYLF-guided programs, five daily meals, and passionate qualified educators. Book today.",
        "h1": "Preschool Near Riverstone That Families Are Proud to Choose",
        "intro": (
            "Flourish Early Learning Centre is a purpose-built preschool near Riverstone guided by the Early "
            "Years Learning Framework (EYLF): Belonging, Being and Becoming. Our qualified educators bring "
            "both professional expertise and genuine warmth to preschool, creating an environment where "
            "children feel at home, feel inspired, and feel proud of who they are. Families near Riverstone "
            "choose Flourish ELC because they can see, and feel, the difference a truly great preschool makes."
        ),
        "why_heading": "Why Choose Flourish ELC for Preschool Near Riverstone",
        "prog": 'A', "env": 'B', "why": 'A', "faq": 'preschool_B', "var": 0,
    },
    {
        "num": 14, "type": "preschool",
        "url": "https://flourishelc.com.au/preschool-mulgrave-nsw/",
        "primary": "preschool near mulgrave", "secondary": "",
        "meta_title": "Preschool Near Mulgrave NSW | Flourish Early Learning Centre",
        "meta_desc": "Looking for a preschool near Mulgrave? Flourish ELC offers EYLF-guided programs, five daily meals, qualified educators, and a warm purpose-built environment. Enquire today.",
        "h1": "Preschool Near Mulgrave: Joyful Learning, Qualified Educators, Happy Children",
        "intro": (
            "Flourish Early Learning Centre is a warm, purpose-built preschool near Mulgrave where children "
            "are welcomed into an environment shaped by the Early Years Learning Framework (EYLF): Belonging, "
            "Being and Becoming. Our qualified educators are passionate about the preschool years and create "
            "daily experiences that are rich, joyful, and genuinely meaningful. Families near Mulgrave trust "
            "Flourish ELC because children love being here, and the learning shows."
        ),
        "why_heading": "Why Choose Flourish ELC for Preschool Near Mulgrave",
        "prog": 'B', "env": 'C', "why": 'B', "faq": 'preschool_C', "var": 1,
    },
    # ── WA Childcare ──
    {
        "num": 15, "type": "childcare",
        "url": "https://flourishelc.com.au/daycare-childcare-bull-creek-wa/",
        "primary": "childcare bull creek", "secondary": "daycare bull creek",
        "meta_title": "Childcare Bull Creek WA | Flourish Early Learning Centre",
        "meta_desc": "Warm, purpose-built childcare in Bull Creek with EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour today.",
        "h1": "Childcare in Bull Creek: Warm, Qualified, and Built Around Your Child",
        "intro": (
            "Flourish Early Learning Centre offers families in Bull Creek a purpose-built early learning "
            "environment where children are guided by the best, and so are their educators. Our programs are "
            "shaped by the Early Years Learning Framework (EYLF): Belonging, Being and Becoming, and delivered "
            "by qualified, passionate educators who genuinely love what they do. Bull Creek families choose "
            "Flourish ELC because it is a place where children thrive, and parents feel completely at ease."
        ),
        "why_heading": "Why Choose Flourish ELC for Childcare in Bull Creek",
        "prog": 'C', "env": 'D', "why": 'C', "faq": 'childcare_D', "var": 2,
    },
    {
        "num": 16, "type": "childcare",
        "url": "https://flourishelc.com.au/daycare-childcare-leeming-wa/",
        "primary": "childcare near leeming", "secondary": "daycare near leeming",
        "meta_title": "Childcare Near Leeming WA | Flourish Early Learning Centre",
        "meta_desc": "Looking for childcare near Leeming? Flourish ELC is purpose-built with EYLF-guided programs, five daily meals, qualified educators, and a parent communication app. Enquire today.",
        "h1": "Childcare Near Leeming: Where Children Belong, Grow, and Flourish",
        "intro": (
            "Flourish Early Learning Centre is a purpose-built early learning centre near Leeming where the "
            "Early Years Learning Framework (EYLF): Belonging, Being and Becoming, is the foundation of every "
            "program, every interaction, and every space. Our qualified, passionate educators create an "
            "environment where children feel genuinely at home, genuinely seen, and genuinely free to learn "
            "at their own wonderful pace. Families near Leeming trust Flourish ELC because the warmth and "
            "quality are real."
        ),
        "why_heading": "Why Choose Flourish ELC for Childcare Near Leeming",
        "prog": 'D', "env": 'A', "why": 'D', "faq": 'childcare_A', "var": 3,
    },
    {
        "num": 17, "type": "childcare",
        "url": "https://flourishelc.com.au/daycare-childcare-bateman-wa/",
        "primary": "childcare near bateman", "secondary": "daycare near bateman",
        "meta_title": "Childcare Near Bateman WA | Flourish Early Learning Centre",
        "meta_desc": "Looking for childcare near Bateman? Flourish ELC is purpose-built with EYLF-guided learning, five daily meals, qualified educators, and a real-time parent app. Book a tour.",
        "h1": "Childcare Near Bateman: Qualified Educators, Warm Community, Happy Children",
        "intro": (
            "Flourish Early Learning Centre is a purpose-built early learning centre near Bateman where "
            "families find exactly what they have been looking for: a warm, qualified, and genuinely "
            "child-centred environment. Guided by the Early Years Learning Framework (EYLF): Belonging, "
            "Being and Becoming, our programs and our people are dedicated to making each child's early years "
            "as rich and joyful as possible. Families near Bateman choose Flourish ELC because the quality "
            "is evident from the moment you walk through the door."
        ),
        "why_heading": "Why Choose Flourish ELC for Childcare Near Bateman",
        "prog": 'A', "env": 'B', "why": 'A', "faq": 'childcare_B', "var": 0,
    },
    {
        "num": 18, "type": "childcare",
        "url": "https://flourishelc.com.au/daycare-childcare-rossmoyne-wa/",
        "primary": "childcare near rossmoyne", "secondary": "daycare near rossmoyne",
        "meta_title": "Childcare Near Rossmoyne WA | Flourish Early Learning Centre",
        "meta_desc": "Looking for childcare near Rossmoyne? Flourish ELC is purpose-built with EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book today.",
        "h1": "Childcare Near Rossmoyne: Where Learning Feels Like Home",
        "intro": (
            "Flourish Early Learning Centre is a warm, purpose-built early learning centre near Rossmoyne "
            "where children are welcomed into an environment guided by the Early Years Learning Framework "
            "(EYLF): Belonging, Being and Becoming. Our qualified, passionate educators are the heart of "
            "Flourish ELC, and it shows. Families near Rossmoyne choose us because their children genuinely "
            "love coming, and the development they witness speaks for itself."
        ),
        "why_heading": "Why Choose Flourish ELC for Childcare Near Rossmoyne",
        "prog": 'B', "env": 'C', "why": 'B', "faq": 'childcare_C', "var": 1,
    },
    {
        "num": 19, "type": "childcare",
        "url": "https://flourishelc.com.au/daycare-childcare-shelley-wa/",
        "primary": "childcare near shelley", "secondary": "daycare near shelley",
        "meta_title": "Childcare Near Shelley WA | Flourish Early Learning Centre",
        "meta_desc": "Looking for childcare near Shelley? Flourish ELC is purpose-built with EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour.",
        "h1": "Childcare Near Shelley: A Place Where Children Thrive",
        "intro": (
            "Flourish Early Learning Centre is a purpose-built early learning centre near Shelley where "
            "families find something genuinely special: a warm community, qualified educators, and programs "
            "shaped by the Early Years Learning Framework (EYLF): Belonging, Being and Becoming. At Flourish "
            "ELC, children are known, celebrated, and supported in becoming the best version of themselves. "
            "Families near Shelley trust us because their children love it here, and the learning is real."
        ),
        "why_heading": "Why Choose Flourish ELC for Childcare Near Shelley",
        "prog": 'C', "env": 'D', "why": 'C', "faq": 'childcare_D', "var": 2,
    },
    {
        "num": 20, "type": "childcare",
        "url": "https://flourishelc.com.au/daycare-childcare-willetton-wa/",
        "primary": "childcare near willetton", "secondary": "daycare near willetton",
        "meta_title": "Childcare Near Willetton WA | Flourish Early Learning Centre",
        "meta_desc": "Looking for childcare near Willetton? Flourish ELC is purpose-built with EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book today.",
        "h1": "Childcare Near Willetton: Warmth, Quality, and Children Who Flourish",
        "intro": (
            "Flourish Early Learning Centre is a purpose-built childcare centre near Willetton where "
            "everything is designed around children having the very best start in life. Guided by the Early "
            "Years Learning Framework (EYLF): Belonging, Being and Becoming, our programs are thoughtful, "
            "joyful, and genuinely focused on each child. Our qualified, passionate educators are what "
            "families near Willetton talk about most, because great educators make everything else possible."
        ),
        "why_heading": "Why Choose Flourish ELC for Childcare Near Willetton",
        "prog": 'D', "env": 'A', "why": 'D', "faq": 'childcare_A', "var": 3,
    },
    {
        "num": 21, "type": "childcare",
        "url": "https://flourishelc.com.au/daycare-childcare-brentwood-wa/",
        "primary": "childcare near brentwood", "secondary": "daycare near brentwood",
        "meta_title": "Childcare Near Brentwood WA | Flourish Early Learning Centre",
        "meta_desc": "Looking for childcare near Brentwood? Flourish ELC is purpose-built with EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book today.",
        "h1": "Childcare Near Brentwood: Where Children Are Known, Loved, and Inspired",
        "intro": (
            "Flourish Early Learning Centre is a purpose-built early learning centre near Brentwood where "
            "families find something they often say is hard to describe but impossible to miss: the feeling "
            "that their child truly belongs here. Our programs are guided by the Early Years Learning "
            "Framework (EYLF): Belonging, Being and Becoming, and our qualified, passionate educators bring "
            "that philosophy to life with warmth and expertise every single day."
        ),
        "why_heading": "Why Choose Flourish ELC for Childcare Near Brentwood",
        "prog": 'A', "env": 'B', "why": 'A', "faq": 'childcare_B', "var": 0,
    },
    # ── WA Kindergarten ──
    {
        "num": 22, "type": "kinder",
        "url": "https://flourishelc.com.au/kindergarten-bull-creek-wa/",
        "primary": "kindergarten bull creek", "secondary": "",
        "meta_title": "Kindergarten Bull Creek WA | Flourish Early Learning Centre",
        "meta_desc": "Looking for kindergarten in Bull Creek? Flourish ELC offers EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour today.",
        "h1": "Kindergarten in Bull Creek: A Joyful, Qualified Start to School Life",
        "intro": (
            "Flourish Early Learning Centre offers a warm, EYLF-guided kindergarten program in Bull Creek "
            "where children are prepared for school in the best possible way: through joyful, intentional, "
            "play-based learning. Our Early Years Learning Framework, Belonging, Being and Becoming, is not "
            "just a concept here: it is the lived experience of every child in our care. Our qualified "
            "educators know each child well and create experiences that light them up."
        ),
        "why_heading": "Why Choose Flourish ELC for Kindergarten in Bull Creek",
        "prog": 'B', "env": 'C', "why": 'B', "faq": 'kinder_A', "var": 1,
    },
    {
        "num": 23, "type": "kinder",
        "url": "https://flourishelc.com.au/kindergarten-leeming-wa/",
        "primary": "kindergarten near leeming", "secondary": "",
        "meta_title": "Kindergarten Near Leeming WA | Flourish Early Learning Centre",
        "meta_desc": "Looking for kindergarten near Leeming? Flourish ELC offers EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour today.",
        "h1": "Kindergarten Near Leeming: Where Children Grow Ready, Willing, and Able",
        "intro": (
            "Flourish Early Learning Centre offers a warm, play-based kindergarten near Leeming guided by "
            "the Early Years Learning Framework (EYLF): Belonging, Being and Becoming. Our kindergarten year "
            "is one of the most important a child will have, and we take that responsibility joyfully and "
            "seriously. Our qualified educators design kindergarten experiences that build confident, curious, "
            "capable children who walk into primary school genuinely ready for everything that comes next."
        ),
        "why_heading": "Why Choose Flourish ELC for Kindergarten Near Leeming",
        "prog": 'C', "env": 'D', "why": 'C', "faq": 'kinder_B', "var": 2,
    },
    {
        "num": 24, "type": "kinder",
        "url": "https://flourishelc.com.au/kindergarten-bateman-wa/",
        "primary": "kindergarten near bateman", "secondary": "",
        "meta_title": "Kindergarten Near Bateman WA | Flourish Early Learning Centre",
        "meta_desc": "Looking for kindergarten near Bateman? Flourish ELC offers EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour today.",
        "h1": "Kindergarten Near Bateman: Play-Based, Purposeful, and Genuinely Wonderful",
        "intro": (
            "Flourish Early Learning Centre offers a warm, EYLF-guided kindergarten program near Bateman "
            "where the year before primary school is spent in the very best way: learning through play, "
            "connecting with others, and growing into a confident, capable child. The Early Years Learning "
            "Framework, Belonging, Being and Becoming, shapes every experience in our kindergarten, and our "
            "qualified, passionate educators bring it to life with skill and joy."
        ),
        "why_heading": "Why Choose Flourish ELC for Kindergarten Near Bateman",
        "prog": 'D', "env": 'A', "why": 'D', "faq": 'kinder_C', "var": 3,
    },
    {
        "num": 25, "type": "kinder",
        "url": "https://flourishelc.com.au/kindergarten-rossmoyne-wa/",
        "primary": "kindergarten near rossmoyne", "secondary": "",
        "meta_title": "Kindergarten Near Rossmoyne WA | Flourish Early Learning Centre",
        "meta_desc": "Looking for kindergarten near Rossmoyne? Flourish ELC offers EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour today.",
        "h1": "Kindergarten Near Rossmoyne: A Joyful, Confident Start to School Life",
        "intro": (
            "Flourish Early Learning Centre offers a warm, purpose-built kindergarten near Rossmoyne where "
            "children spend the year before school exactly as they should: playing, discovering, growing, and "
            "becoming. Our kindergarten programs are guided by the Early Years Learning Framework (EYLF): "
            "Belonging, Being and Becoming, and delivered by qualified educators who genuinely love the "
            "kindergarten years."
        ),
        "why_heading": "Why Choose Flourish ELC for Kindergarten Near Rossmoyne",
        "prog": 'A', "env": 'B', "why": 'A', "faq": 'kinder_D', "var": 0,
    },
    {
        "num": 26, "type": "kinder",
        "url": "https://flourishelc.com.au/kindergarten-shelley-wa/",
        "primary": "kindergarten near shelley", "secondary": "",
        "meta_title": "Kindergarten Near Shelley WA | Flourish Early Learning Centre",
        "meta_desc": "Looking for kindergarten near Shelley? Flourish ELC offers EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour today.",
        "h1": "Kindergarten Near Shelley: Where the Best Possible Year Before School Happens",
        "intro": (
            "Flourish Early Learning Centre offers a warm, play-based kindergarten near Shelley guided by "
            "the Early Years Learning Framework (EYLF): Belonging, Being and Becoming. The year before "
            "primary school is one of the most important a child will have, and our qualified, passionate "
            "educators pour everything into making it wonderful. Children in our kindergarten near Shelley "
            "grow academically, socially, and emotionally in an environment that is safe, joyful, and "
            "genuinely child-centred."
        ),
        "why_heading": "Why Choose Flourish ELC for Kindergarten Near Shelley",
        "prog": 'B', "env": 'C', "why": 'B', "faq": 'kinder_A', "var": 1,
    },
    {
        "num": 27, "type": "kinder",
        "url": "https://flourishelc.com.au/kindergarten-willetton-wa/",
        "primary": "kindergarten near willetton", "secondary": "",
        "meta_title": "Kindergarten Near Willetton WA | Flourish Early Learning Centre",
        "meta_desc": "Looking for kindergarten near Willetton? Flourish ELC offers EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour today.",
        "h1": "Kindergarten Near Willetton: Play-Based, Warm, and Built for the Best Start",
        "intro": (
            "Flourish Early Learning Centre offers a wonderful kindergarten near Willetton where children "
            "spend the year before primary school growing in every direction. Our kindergarten programs are "
            "guided by the Early Years Learning Framework (EYLF): Belonging, Being and Becoming, and "
            "delivered by qualified, passionate educators who make every day count. Families near Willetton "
            "choose our kindergarten because they see the quality in the detail: the care, the curriculum, "
            "the community, and the children who thrive here."
        ),
        "why_heading": "Why Choose Flourish ELC for Kindergarten Near Willetton",
        "prog": 'C', "env": 'D', "why": 'C', "faq": 'kinder_B', "var": 2,
    },
    {
        "num": 28, "type": "kinder",
        "url": "https://flourishelc.com.au/kindergarten-brentwood-wa/",
        "primary": "kindergarten near brentwood", "secondary": "",
        "meta_title": "Kindergarten Near Brentwood WA | Flourish Early Learning Centre",
        "meta_desc": "Looking for kindergarten near Brentwood? Flourish ELC offers EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour today.",
        "h1": "Kindergarten Near Brentwood: Your Child's Best Year Before School Starts Here",
        "intro": (
            "Flourish Early Learning Centre offers a warm, purpose-built kindergarten near Brentwood where "
            "the year before primary school is spent exactly as it should be: playing, learning, growing, "
            "and belonging. Guided by the Early Years Learning Framework (EYLF): Belonging, Being and "
            "Becoming, our kindergarten programs are shaped around each child's unique journey. Our qualified, "
            "passionate educators bring both skill and genuine heart to every kindergarten session."
        ),
        "why_heading": "Why Choose Flourish ELC for Kindergarten Near Brentwood",
        "prog": 'D', "env": 'A', "why": 'D', "faq": 'kinder_C', "var": 3,
    },
]

# FAQ lookup
faq_lookup = {}
faq_lookup.update({'childcare_' + k: v for k, v in faqs_childcare.items()})
faq_lookup.update({'preschool_' + k: v for k, v in faqs_preschool.items()})
faq_lookup.update({'kinder_' + k: v for k, v in faqs_kinder.items()})


def add_page(p):
    doc.add_heading(f"Page {p['num']}", level=1)

    # Info block
    info = f"URL: {p['url']}\nPage Type: Location Page\nPrimary Keyword: {p['primary']}"
    if p['secondary']:
        info += f"\nSecondary Keywords: {p['secondary']}"
    doc.add_paragraph(info)

    doc.add_paragraph("Content:")
    doc.add_paragraph(f"Meta Title: {p['meta_title']}")
    doc.add_paragraph(f"Meta Description: {p['meta_desc']}")

    # H1
    doc.add_heading(f"H1: {p['h1']}", level=2)
    doc.add_paragraph(p['intro'])

    # Build section order based on variant
    prog_sec = programs[p['prog']]
    env_sec = environment[p['env']]
    why_body = why_choose[p['why']]
    why_heading = f"H2: {p['why_heading']}"

    sections_prog = ('prog', prog_sec)
    sections_env = ('env', env_sec)
    sections_why = ('why', {'heading': why_heading, 'body': why_body})

    order_map = {
        0: [sections_prog, sections_env, sections_why],
        1: [sections_env, sections_why, sections_prog],
        2: [sections_why, sections_prog, sections_env],
        3: [sections_prog, sections_why, sections_env],
    }
    section_order = order_map[p['var']]

    for _, sec in section_order:
        heading_text = sec['heading']
        # Don't double-prefix if already has H2:
        if not heading_text.startswith('H2:'):
            heading_text = f"H2: {heading_text}"
        doc.add_heading(heading_text, level=3)
        doc.add_paragraph(sec['body'])

    # FAQs
    doc.add_heading("H2: Frequently Asked Questions", level=3)
    faq_list = faq_lookup[p['faq']]
    for q, a in faq_list:
        para = doc.add_paragraph()
        run = para.add_run(q)
        run.bold = True
        doc.add_paragraph(a)

    doc.add_page_break()


# Generate pages 1–28
for p in pages:
    add_page(p)
    print(f"Page {p['num']} done")

# ── PAGE 29 ── Areas We Serve
doc.add_heading("Page 29", level=1)
doc.add_paragraph(
    "URL: https://flourishelc.com.au/areas-we-serve/\n"
    "Page Type: Location Page\n"
    "Primary Keyword: childcare near me\n"
    "Secondary Keywords: daycare near me, preschool near me, kindergarten near me"
)
doc.add_paragraph("Content:")
doc.add_paragraph("Meta Title: Find Flourish Early Learning Centres Near You")
doc.add_paragraph(
    "Meta Description: Looking for childcare near me? Flourish ELC has warm, purpose-built centres "
    "with EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app."
)

doc.add_heading("H1: Childcare Near Me: Find a Flourish Early Learning Centre Close to Your Family", level=2)
doc.add_paragraph(
    "Whatever brought you to searching for childcare near me, a new baby, a house move, or simply "
    "the feeling that it is time, you have come to the right place. Flourish Early Learning Centre "
    "is a family of purpose-built early learning centres guided by the Early Years Learning Framework "
    "(EYLF): Belonging, Being and Becoming. Our qualified, passionate educators create environments "
    "where children feel at home, families feel supported, and learning feels like the joyful adventure "
    "it is meant to be."
)

doc.add_heading("H2: What Makes Flourish ELC Special Wherever You Find Us", level=3)
doc.add_paragraph(
    "Every Flourish ELC centre, whether you are searching for childcare, daycare, preschool, or "
    "kindergarten near you, is built around the same principles. Our programs cover Creative Arts, "
    "STEM Learning, Language, Mathematics, Fine Motor Skills, Music and Movement, Cultural "
    "Understanding, Social Justice, Independence, Exercise, Health, and Social and Emotional "
    "development. These are not just learning areas: they are the threads of a rich, whole-child "
    "experience that children and families love.\n\n"
    "Five nutritious meals are served fresh every weekday: breakfast, morning tea, lunch, afternoon "
    "tea, and a late snack. Our centres are open 7:00am to 6:00pm Monday to Friday. Our parent "
    "communication app keeps every family connected with real-time updates and photos throughout the "
    "day. School readiness programs are woven through our curriculum so that children transition to "
    "kindergarten and primary school with confidence, capability, and enthusiasm."
)

doc.add_heading("H2: Why Choose Flourish ELC for Childcare Near You", level=3)
doc.add_paragraph(
    "Flourish ELC is guided by the Early Years Learning Framework (EYLF): Belonging, Being and "
    "Becoming. That is the foundation of every decision we make, from how we design our spaces to "
    "how we hire and support our educators. Our centres meet and exceed the National Quality Framework "
    "(NQF), which means the care, the programs, and the environments your child experiences are "
    "assessed against Australia's highest national standards for early education.\n\n"
    "Our educators are qualified professionals who bring both formal training and genuine warmth to "
    "every interaction. They take the time to know each child: what makes them curious, what helps "
    "them feel secure, and what they are ready to discover next. School readiness programs ensure "
    "that children leave Flourish ELC ready for kindergarten and primary school with the skills and "
    "the confidence to embrace what comes next.\n\n"
    "Five nutritious meals are provided every weekday. Dietary requirements and allergies are always "
    "carefully managed. We are open Monday to Friday, 7:00am to 6:00pm, and our parent communication "
    "app delivers real-time updates and photos so you are never far from your child's world."
)

doc.add_heading("H2: Qualified Educators, Purpose-Built Spaces, Real-Time Connection", level=3)
doc.add_paragraph(
    "Every Flourish ELC centre is staffed by qualified educators who meet and exceed the National "
    "Quality Framework (NQF). Our facilities are purpose-built for children from 6 weeks to school "
    "age: not repurposed spaces, but environments designed from the ground up around how young "
    "children learn, rest, play, and grow. Families searching for preschool or kindergarten near me "
    "will find that same consistent standard of care and quality across every Flourish ELC location.\n\n"
    "Our parent communication app is one of the things families love most. It delivers real-time "
    "photos, updates, and observations throughout the day so even when you are at work, you are never "
    "far from your child's world. It is a small but meaningful reflection of how much we value the "
    "partnership between families and educators."
)

doc.add_heading("H2: Frequently Asked Questions", level=3)
areas_faqs = [
    ("What age groups do Flourish ELC centres cater for?",
     "All our centres welcome children from 6 weeks through to school age, with dedicated programs and spaces for each developmental stage."),
    ("What meals does Flourish ELC provide?",
     "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. All dietary requirements and allergies are carefully managed."),
    ("What are the opening hours across Flourish ELC centres?",
     "Our centres are open Monday to Friday, 7:00am to 6:00pm."),
    ("How does the parent communication app work?",
     "The app sends real-time updates, photos, and milestones throughout the day so families always feel genuinely connected to their child's world."),
    ("What does the EYLF mean for my child at Flourish ELC?",
     "The Early Years Learning Framework guides how our educators plan, observe, and respond to each child, always with belonging, wellbeing, and engagement at the heart of practice. It is the foundation of everything we do."),
]
for q, a in areas_faqs:
    para = doc.add_paragraph()
    run = para.add_run(q)
    run.bold = True
    doc.add_paragraph(a)

print("Page 29 done")

output_path = "/home/invoi/fahad_projects/clients/flourishelc.com.au/flourishelc.com.au_v3.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
