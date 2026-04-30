from docx import Document

doc = Document()

# ── VARIANT KEY SCHEME ──
# var_key (A-H) for cultural, school_readiness, why_bullets, environment
# Group offsets: NSW childcare=0, NSW preschool=1, WA childcare=2, WA kinder=3
# var_key = 'ABCDEFGH'[(group_offset + page_index_in_group) % 8]
# {suburb}, {service_type}, {state} placeholders rendered per page

# ─────────────────────────────────────────────
# PROGRAMS — CHILDCARE (7 variants A-G, {state} param)
# ─────────────────────────────────────────────
prog_childcare = {
    'A': {
        'heading': 'Programs That Cover Every Part of Childhood',
        'body': (
            "In {state}, Flourish ELC delivers a curriculum that is broad, intentional, and alive "
            "with possibility. Every day, children engage with STEM Learning, Creative Arts, Language, "
            "Mathematics, Fine Motor Skills, Music and Movement, Cultural Understanding, Social Justice, "
            "Independence, Exercise, Health, and Social and Emotional development through experiences "
            "that feel like adventure rather than instruction. These are not separate lessons delivered "
            "in sequence: they flow through play, projects, conversations, and discoveries that children "
            "genuinely love.\n\n"
            "Five nutritious meals are prepared and served every weekday: breakfast, morning tea, lunch, "
            "afternoon tea, and a late snack. Dietary requirements and allergies are always managed "
            "with care. School readiness is woven through everything so that children arrive at "
            "kindergarten with the skills, the confidence, and the enthusiasm to thrive."
        )
    },
    'B': {
        'heading': 'Learning Through Play, Growing in Every Direction',
        'body': (
            "Play is how children understand the world, and at Flourish ELC in {state}, play is the "
            "method, the medium, and the magic. Social and Emotional development sits at the heart of "
            "every day because children who feel secure and connected learn better and grow more "
            "confidently. STEM Learning, Creative Arts, Language, Mathematics, Fine Motor Skills, Music "
            "and Movement, Cultural Understanding, Social Justice, Independence, Exercise, and Health "
            "are all woven into joyful experiences that children find meaningful and exciting.\n\n"
            "Independence is nurtured from an early age, with educators guiding rather than directing, "
            "so children develop the judgment and self-belief that carry them forward. Five nutritious "
            "meals are served each weekday, and our school readiness programs ensure the transition to "
            "kindergarten is confident and joyful rather than uncertain."
        )
    },
    'C': {
        'heading': 'A Curriculum Designed for Curious, Capable Children',
        'body': (
            "The Flourish ELC childcare curriculum in {state} is designed around children who are "
            "curious, capable, and full of questions. Language and Mathematics build the academic "
            "foundations children carry into school. STEM Learning sparks inventive thinking. Creative "
            "Arts and Music and Movement give children expressive outlets that build confidence and "
            "joy. Cultural Understanding and Social Justice ensure that children develop empathy and a "
            "broad view of the world from the very beginning. Exercise and Health are part of every "
            "day, so children grow physically alongside everything else they discover.\n\n"
            "Five nutritious meals are provided each weekday: breakfast, morning tea, lunch, afternoon "
            "tea, and a late snack. Our school readiness programs are embedded throughout so children "
            "are genuinely prepared when kindergarten arrives."
        )
    },
    'D': {
        'heading': 'Rich Programs Across Every Learning Area',
        'body': (
            "At Flourish ELC in {state}, no dimension of a child's development is left to chance. "
            "Fine Motor Skills develop through hands-on creative activities. Language and Mathematics "
            "flow naturally through the daily rhythm of the centre. Cultural Understanding and Social "
            "Justice are embedded in everyday experiences because we want children to grow into people "
            "who are curious about the world and kind within it. Exercise and Health are woven into "
            "outdoor play and movement every single day.\n\n"
            "STEM Learning, Creative Arts, Music and Movement, Social and Emotional development, and "
            "Independence complete a curriculum designed to develop every child in every way. Five "
            "nutritious meals are served each weekday, and school readiness programs ensure every "
            "child is ready to embrace kindergarten with genuine excitement."
        )
    },
    'E': {
        'heading': 'Creative, Curious, and Ready for the World',
        'body': (
            "Imagination is the engine of learning, and at Flourish ELC in {state}, Creative Arts "
            "and STEM Learning are given the space and resources they deserve. Children paint, build, "
            "question, experiment, and create every single day, developing the critical thinking and "
            "problem-solving skills that serve them across every area of life. Music and Movement "
            "give children ways to express themselves that go far beyond words, while Language and "
            "Mathematics lay the academic groundwork for the years ahead.\n\n"
            "Cultural Understanding, Social Justice, Independence, Exercise, Health, and Social and "
            "Emotional development complete a holistic curriculum that develops the whole child. "
            "Five nutritious meals are served each weekday. Our school readiness programs ensure "
            "every child transitions to kindergarten feeling capable and excited."
        )
    },
    'F': {
        'heading': 'Language, Learning, and a Love of Discovery',
        'body': (
            "At Flourish ELC in {state}, language is everywhere. In stories, songs, conversations, "
            "and collaborative projects, children develop vocabulary, communication, and the "
            "confidence to express their ideas and feelings. Alongside Language, our daily programs "
            "weave in STEM Learning, Creative Arts, Mathematics, Fine Motor Skills, Music and "
            "Movement, Cultural Understanding, Social Justice, Independence, Exercise, Health, and "
            "Social and Emotional development into experiences children find genuinely engaging.\n\n"
            "Independence is cultivated gently, with educators supporting children in building "
            "self-reliance and pride in their own capabilities. Five nutritious meals are prepared "
            "each weekday, and our school readiness programs give every child the foundation they "
            "need to begin kindergarten with confidence."
        )
    },
    'G': {
        'heading': 'Programs Where Every Child Belongs and Grows',
        'body': (
            "At Flourish ELC in {state}, children come from many different backgrounds, and every "
            "one of them is celebrated. Cultural Understanding and Social Justice are not afterthoughts "
            "in our curriculum: they are woven into daily life through stories, discussions, art, and "
            "shared experiences that broaden every child's view of the world. Alongside this, STEM "
            "Learning, Creative Arts, Language, Mathematics, Fine Motor Skills, Music and Movement, "
            "Independence, Exercise, Health, and Social and Emotional development fill every day with "
            "rich, purposeful learning.\n\n"
            "Children build confidence, curiosity, and a genuine love of learning that stays with "
            "them long after they leave our centre. Five nutritious meals are served each weekday. "
            "School readiness is embedded throughout so every child transitions to kindergarten feeling "
            "truly ready."
        )
    },
}

# ─────────────────────────────────────────────
# PROGRAMS — PRESCHOOL (7 variants A-G)
# ─────────────────────────────────────────────
prog_preschool = {
    'A': {
        'heading': 'Preschool Programs That Develop the Whole Child',
        'body': (
            "The Flourish ELC preschool program is built around the understanding that 3 to 5 year "
            "olds are at their most curious, most creative, and most ready to grow. Every day, "
            "preschoolers engage with STEM Learning, Creative Arts, Language, Mathematics, Fine Motor "
            "Skills, Music and Movement, Cultural Understanding, Social Justice, Independence, Exercise, "
            "Health, and Social and Emotional development through experiences that feel like play "
            "because, at this age, the best learning always does.\n\n"
            "Five nutritious meals are prepared and served each weekday. Our preschool readiness "
            "programs progressively build the skills, confidence, and enthusiasm children need to "
            "transition into kindergarten feeling genuinely prepared for everything that comes next."
        )
    },
    'B': {
        'heading': 'A Preschool Curriculum Full of Wonder and Intention',
        'body': (
            "At Flourish ELC, the preschool years are understood as one of the richest and most "
            "important chapters of a child's life. Our curriculum gives preschoolers daily access to "
            "STEM Learning, Creative Arts, Language, Mathematics, Fine Motor Skills, Music and "
            "Movement, Cultural Understanding, Social Justice, Independence, Exercise, Health, and "
            "Social and Emotional development, all through experiences that are joyful, intentional, "
            "and deeply engaging for this age group.\n\n"
            "Social and Emotional development is given particular attention in our preschool program, "
            "because children who feel emotionally secure flourish academically and socially. Five "
            "nutritious meals are served each weekday, and school readiness programs ensure every "
            "preschooler transitions to kindergarten with confidence and genuine excitement."
        )
    },
    'C': {
        'heading': 'Preschool Learning Across Every Domain',
        'body': (
            "Flourish ELC's preschool program is designed around the remarkable potential of 3 to "
            "5 year old learners. Language and Mathematics lay the academic foundations preschoolers "
            "will build on for years. Creative Arts and Music and Movement nurture self-expression "
            "and joy. STEM Learning develops the inventive thinking that serves children across every "
            "discipline. Cultural Understanding and Social Justice broaden each child's view of the "
            "world with empathy and curiosity.\n\n"
            "Fine Motor Skills, Exercise, Health, Independence, and Social and Emotional development "
            "complete a preschool curriculum that addresses every dimension of childhood. Five "
            "nutritious meals are served each weekday. Our school readiness programs ensure that "
            "the transition to kindergarten is smooth, confident, and something children look forward to."
        )
    },
    'D': {
        'heading': 'Growing Confident, Curious Preschoolers Every Day',
        'body': (
            "Confidence and curiosity are the foundations of a great preschool experience, and at "
            "Flourish ELC, both are cultivated with care and intention every day. Our preschool "
            "program gives children daily engagement with STEM Learning, Creative Arts, Language, "
            "Mathematics, Fine Motor Skills, Music and Movement, Cultural Understanding, Social "
            "Justice, Independence, Exercise, Health, and Social and Emotional development through "
            "experiences that feel natural, exciting, and exactly right for this age.\n\n"
            "Independence is a particular focus in our preschool program, because children who can "
            "manage themselves feel capable and proud. Five nutritious meals are served each weekday. "
            "Our school readiness programs build every skill children need to walk into kindergarten "
            "feeling genuinely ready and genuinely excited."
        )
    },
    'E': {
        'heading': 'Preschool That Sparks a Lifelong Love of Learning',
        'body': (
            "The Flourish ELC preschool program is designed to do more than prepare children for "
            "kindergarten. It is designed to spark the kind of love of learning that stays with "
            "children throughout their lives. Every day, preschoolers explore STEM Learning and "
            "Creative Arts, build Language and Mathematics skills, develop Fine Motor abilities, "
            "engage with Music and Movement, learn about Cultural Understanding and Social Justice, "
            "grow their Independence, and develop the Health, Exercise habits, and Social and "
            "Emotional foundations that support everything else.\n\n"
            "Five nutritious meals are prepared fresh each weekday: breakfast, morning tea, lunch, "
            "afternoon tea, and a late snack. Our school readiness focus ensures every preschooler "
            "transitions to kindergarten not just ready but genuinely eager."
        )
    },
    'F': {
        'heading': 'A Preschool Program Rich in Language and Creativity',
        'body': (
            "Language is at the heart of the Flourish ELC preschool experience. In stories, songs, "
            "discussions, and collaborative play, preschoolers build vocabulary, communication skills, "
            "and the confidence to express who they are and what they think. Alongside Language, "
            "Creative Arts, STEM Learning, Mathematics, Fine Motor Skills, Music and Movement, "
            "Cultural Understanding, Social Justice, Independence, Exercise, Health, and Social and "
            "Emotional development make every preschool day rich and purposeful.\n\n"
            "Educators plan preschool experiences that reflect children's interests and developmental "
            "needs, so the curriculum always feels relevant and engaging. Five nutritious meals are "
            "served each weekday, and school readiness is embedded in everything we do."
        )
    },
    'G': {
        'heading': 'Preschool Programs for Every Child, Every Day',
        'body': (
            "Every child who walks through the door of Flourish ELC is unique, and our preschool "
            "program is designed to honour that. Cultural Understanding and Social Justice are woven "
            "into daily preschool life so that every child feels seen, celebrated, and genuinely "
            "included. STEM Learning, Creative Arts, Language, Mathematics, Fine Motor Skills, Music "
            "and Movement, Independence, Exercise, Health, and Social and Emotional development "
            "complete a curriculum that develops preschoolers in every direction.\n\n"
            "Five nutritious meals are served each weekday, and our school readiness programs give "
            "every preschooler the social, emotional, cognitive, and academic foundation they need "
            "to transition into kindergarten with confidence and genuine joy."
        )
    },
}

# ─────────────────────────────────────────────
# PROGRAMS — KINDERGARTEN (7 variants A-G)
# ─────────────────────────────────────────────
prog_kinder = {
    'A': {
        'heading': 'Kindergarten Programs Built for the Year Before School',
        'body': (
            "The kindergarten year at Flourish ELC is one of the most important a child will "
            "experience, and we take that responsibility joyfully and seriously. Our kindergarten "
            "program gives children daily engagement with STEM Learning, Creative Arts, Language, "
            "Mathematics, Fine Motor Skills, Music and Movement, Cultural Understanding, Social "
            "Justice, Independence, Exercise, Health, and Social and Emotional development through "
            "play-based experiences that are rich, intentional, and genuinely exciting for this age.\n\n"
            "Five nutritious meals are prepared and served each weekday. Our school readiness "
            "programs develop every skill and disposition children need to begin primary school "
            "feeling confident, curious, and genuinely prepared for the adventure ahead."
        )
    },
    'B': {
        'heading': 'A Kindergarten Curriculum That Prepares and Inspires',
        'body': (
            "Flourish ELC's kindergarten program is designed around the whole child, not just their "
            "academic preparation. Language and Mathematics build strong foundations for primary "
            "school. STEM Learning develops the creative problem-solving that serves children across "
            "every subject. Creative Arts and Music and Movement give children confidence and joy. "
            "Cultural Understanding, Social Justice, Independence, Fine Motor Skills, Exercise, "
            "Health, and Social and Emotional development ensure no dimension of childhood is left "
            "undeveloped in this crucial year.\n\n"
            "Five nutritious meals are served each weekday. Our school readiness programs are "
            "embedded throughout the kindergarten year so children transition to primary school "
            "with the skills, the confidence, and the enthusiasm they deserve."
        )
    },
    'C': {
        'heading': 'Learning Through Play in the Year That Matters Most',
        'body': (
            "Play-based learning is not a compromise in the Flourish ELC kindergarten program: it "
            "is the approach that gets the best results for this age group. Children engage with "
            "STEM Learning, Creative Arts, Language, Mathematics, Fine Motor Skills, Music and "
            "Movement, Cultural Understanding, Social Justice, Independence, Exercise, Health, and "
            "Social and Emotional development through experiences they find genuinely meaningful and "
            "joyful. The learning is real. The enjoyment is real.\n\n"
            "Five nutritious meals are prepared each weekday. Our school readiness programs "
            "systematically build the academic, social, and emotional skills children need to begin "
            "primary school as confident, capable learners who are ready for everything that comes next."
        )
    },
    'D': {
        'heading': 'Kindergarten Programs for Confident, School-Ready Children',
        'body': (
            "Flourish ELC's kindergarten program brings together play, intentional teaching, and a "
            "deep understanding of what children aged 4 to 5 need to thrive. Independence is given "
            "particular focus because kindergarteners who manage themselves confidently carry that "
            "capability into primary school and beyond. STEM Learning, Creative Arts, Language, "
            "Mathematics, Fine Motor Skills, Music and Movement, Cultural Understanding, Social "
            "Justice, Exercise, Health, and Social and Emotional development make every kindergarten "
            "day full and purposeful.\n\n"
            "Five nutritious meals are served each weekday. Our school readiness programs ensure "
            "every child who completes our kindergarten year walks into primary school genuinely "
            "ready, genuinely excited, and genuinely capable."
        )
    },
    'E': {
        'heading': 'A Kindergarten Program Rich in Creativity and Discovery',
        'body': (
            "Imagination and discovery are at the heart of the Flourish ELC kindergarten experience. "
            "Creative Arts and STEM Learning give kindergarteners space to invent, create, and "
            "question every day. Music and Movement build confidence and expression. Language and "
            "Mathematics develop the academic foundations that will serve children throughout primary "
            "school. Cultural Understanding, Social Justice, Fine Motor Skills, Independence, "
            "Exercise, Health, and Social and Emotional development complete a kindergarten "
            "curriculum that honours every dimension of this important year.\n\n"
            "Five nutritious meals are served each weekday. School readiness programs are woven "
            "through everything so children transition to primary school feeling genuinely prepared, "
            "excited, and proud of who they are."
        )
    },
    'F': {
        'heading': 'Language and Learning at the Heart of Kindergarten',
        'body': (
            "Language development is a priority in the Flourish ELC kindergarten program because "
            "strong language skills underpin success in every area of primary school and beyond. "
            "In stories, discussions, songs, and collaborative projects, kindergarteners build "
            "vocabulary, communication, and the confidence to share their ideas with the world. "
            "Alongside Language, STEM Learning, Creative Arts, Mathematics, Fine Motor Skills, "
            "Music and Movement, Cultural Understanding, Social Justice, Independence, Exercise, "
            "Health, and Social and Emotional development fill every kindergarten day.\n\n"
            "Five nutritious meals are served each weekday. Our school readiness programs build "
            "every skill children need to begin primary school as confident, enthusiastic learners."
        )
    },
    'G': {
        'heading': 'Kindergarten Where Every Child Belongs and Thrives',
        'body': (
            "At Flourish ELC, kindergarten is a time for belonging as much as it is for learning. "
            "Cultural Understanding and Social Justice are woven into daily kindergarten life so "
            "that every child feels included, respected, and genuinely valued. STEM Learning, "
            "Creative Arts, Language, Mathematics, Fine Motor Skills, Music and Movement, "
            "Independence, Exercise, Health, and Social and Emotional development make every "
            "kindergarten day rich, purposeful, and joyful.\n\n"
            "Five nutritious meals are prepared and served each weekday. Our school readiness "
            "programs ensure every child transitions to primary school with the academic, social, "
            "and emotional foundation they need to thrive from the very first day."
        )
    },
}

# ─────────────────────────────────────────────
# MUNCH & MOVE — NSW (7 variants A-G, {service_type} param)
# ─────────────────────────────────────────────
munch_nsw = {
    'A': (
        "The Munch & Move program is a NSW Health initiative that promotes healthy eating, physical "
        "activity, and reduced screen time for children in early childhood settings like our "
        "{service_type} program. Through garden-to-plate experiences, hands-on cooking activities, "
        "and fun movement sessions, children develop a positive relationship with food and an active "
        "lifestyle from the very beginning. Our educators use Munch & Move resources to make healthy "
        "choices feel natural, enjoyable, and genuinely part of everyday life at Flourish ELC."
    ),
    'B': (
        "Munch & Move is a NSW Health program that guides our {service_type} educators in promoting "
        "healthy eating and active play for young children. At Flourish ELC, Munch & Move comes to "
        "life through movement games, mindful mealtimes, and activities that teach children to love "
        "both nutritious food and physical activity. The program gives our team evidence-based "
        "resources and training so that healthy habits are woven into every part of the day, not "
        "treated as a separate lesson."
    ),
    'C': (
        "In our {service_type} program, the Munch & Move initiative from NSW Health supports "
        "children's physical and nutritional wellbeing through engaging, age-appropriate activities. "
        "Children learn about food choices through cooking, gardening, and mindful eating practices "
        "that make mealtimes a positive and exploratory experience. Movement is embedded in outdoor "
        "play, games, and structured activities that help children develop gross motor skills while "
        "building lifelong healthy habits in a fun, low-pressure environment."
    ),
    'D': (
        "The NSW Health Munch & Move program is an important part of how Flourish ELC supports "
        "the health and wellbeing of children in our {service_type} rooms. Our educators are trained "
        "in Munch & Move approaches that encourage active lifestyles and positive food relationships, "
        "reducing screen time and increasing meaningful physical engagement throughout the day. "
        "Children benefit from a balanced approach that celebrates healthy eating and joyful "
        "movement as natural parts of growing up."
    ),
    'E': (
        "Flourish ELC's {service_type} program incorporates Munch & Move, a NSW Health initiative "
        "that equips educators with tools to promote healthy food choices, physical activity, and "
        "reduced screen time. In practice, this means children engage in hands-on cooking "
        "experiences, active outdoor play, and mindful mealtimes that build positive habits they "
        "will carry into school and beyond. Munch & Move resources ensure our approach to health "
        "and nutrition is evidence-based, engaging, and genuinely beneficial for young children."
    ),
    'F': (
        "Through the Munch & Move program — a NSW Health initiative — our {service_type} educators "
        "actively promote physical activity, nutritious food choices, and reduced screen time for "
        "every child in our care. From movement games to garden-to-plate cooking, children develop "
        "a genuine love of being active and eating well. Munch & Move training ensures our team has "
        "the resources and knowledge to make health promotion a natural, joyful part of every day "
        "at Flourish ELC."
    ),
    'G': (
        "At Flourish ELC, our {service_type} program integrates the Munch & Move framework from "
        "NSW Health to support children's physical development and overall wellbeing. Children "
        "participate in active play sessions, mindfulness practices, and cooking experiences that "
        "build a lifelong appreciation for healthy living. Munch & Move gives our educators "
        "evidence-based strategies and training so that nutrition and physical activity are always "
        "age-appropriate, fun, and meaningfully connected to children's daily lives and interests."
    ),
}

# ─────────────────────────────────────────────
# MUNCH & MOVE — WA (7 variants A-G, {service_type} param, no NSW Health)
# ─────────────────────────────────────────────
munch_wa = {
    'A': (
        "The Munch & Move program supports healthy eating, physical activity, and reduced screen "
        "time for children in early childhood settings like our {service_type} program. Through "
        "garden-to-plate experiences, hands-on cooking, and fun movement sessions, children develop "
        "a positive relationship with food and an active lifestyle from the very beginning. Our "
        "educators use Munch & Move resources to make healthy choices feel natural, enjoyable, "
        "and genuinely part of every day at Flourish ELC."
    ),
    'B': (
        "Munch & Move guides our {service_type} educators in promoting healthy eating and active "
        "play for young children. At Flourish ELC, this comes to life through movement games, "
        "mindful mealtimes, and activities that teach children to love both nutritious food and "
        "physical activity. Our team has evidence-based resources and training so that healthy "
        "habits are woven into every part of the day, not treated as a separate lesson."
    ),
    'C': (
        "In our {service_type} program, Munch & Move supports children's physical and nutritional "
        "wellbeing through engaging, age-appropriate activities. Children learn about food choices "
        "through cooking, gardening, and mindful eating practices that make mealtimes a positive "
        "and exploratory experience. Movement is embedded in outdoor play, games, and structured "
        "activities that help children build healthy habits in a fun, low-pressure environment."
    ),
    'D': (
        "The Munch & Move program is an important part of how Flourish ELC supports the health "
        "and wellbeing of children in our {service_type} program. Our educators are trained in "
        "Munch & Move approaches that encourage active lifestyles and positive food relationships, "
        "reducing screen time and increasing meaningful physical engagement throughout the day. "
        "Children benefit from a balanced approach that celebrates healthy eating and joyful "
        "movement as natural parts of childhood."
    ),
    'E': (
        "Flourish ELC's {service_type} program incorporates Munch & Move, which equips educators "
        "with tools to promote healthy food choices, physical activity, and reduced screen time. "
        "In practice, this means children engage in hands-on cooking experiences, active outdoor "
        "play, and mindful mealtimes that build positive habits they will carry into school and "
        "beyond. Munch & Move resources ensure our approach to health and nutrition is "
        "evidence-based, engaging, and genuinely beneficial for young children."
    ),
    'F': (
        "Through the Munch & Move program, our {service_type} educators actively promote physical "
        "activity, nutritious food choices, and reduced screen time for every child. From movement "
        "games to garden-to-plate cooking, children develop a genuine love of being active and "
        "eating well. Munch & Move training ensures our team has the resources and knowledge to "
        "make health promotion a natural, joyful part of every day at Flourish ELC."
    ),
    'G': (
        "At Flourish ELC, our {service_type} program integrates Munch & Move to support children's "
        "physical development and overall wellbeing. Children participate in active play sessions, "
        "mindfulness practices, and cooking experiences that build a lifelong appreciation for "
        "healthy living. Munch & Move gives our educators evidence-based strategies so that "
        "nutrition and physical activity are always age-appropriate, fun, and meaningfully connected "
        "to each child's daily life."
    ),
}

# ─────────────────────────────────────────────
# CULTURAL EDUCATION (8 variants A-H, {suburb} + {service_type} params)
# ─────────────────────────────────────────────
cultural = {
    'A': (
        "In our {suburb} {service_type} program, diversity and inclusion are celebrated every day "
        "through our Cultural Education program. Children learn about Indigenous heritage, world "
        "cultures, and the importance of embracing differences through storytelling, music, art, "
        "and traditional celebrations that bring cultures to life in meaningful ways. Our program "
        "encourages inclusivity, understanding, and a global mindset, helping children develop a "
        "sense of belonging in a multicultural world where all cultures are genuinely valued and explored."
    ),
    'B': (
        "Our Cultural Education program is woven into the daily life of our {suburb} {service_type} "
        "program, where children explore Indigenous heritage, celebrate world cultures, and develop "
        "a deep respect for people from all backgrounds. Through music, art, storytelling, and "
        "traditional celebrations, children gain an appreciation for cultural diversity that shapes "
        "how they see and relate to the world. By fostering an environment where every culture is "
        "explored with curiosity and respect, we create meaningful connections that enrich every "
        "child's learning journey."
    ),
    'C': (
        "At our {suburb} {service_type} centre, Cultural Education means giving children genuine "
        "windows into the rich diversity of our world. Aboriginal and Torres Strait Islander "
        "perspectives are honoured alongside traditions from cultures around the globe, through "
        "storytelling, music, art, and shared celebrations that make learning personal and memorable. "
        "Our program builds inclusivity, understanding, and a global mindset so that every child "
        "grows up knowing how to embrace differences and find connection across all backgrounds."
    ),
    'D': (
        "Cultural Education at our {suburb} {service_type} program celebrates the diversity that "
        "makes our community extraordinary. Children explore Indigenous heritage, world traditions, "
        "and the beauty of cultural differences through creative activities including art, music, "
        "storytelling, and traditional celebrations. This immersive approach builds cultural empathy, "
        "respect, and a genuine global mindset, helping every child develop a strong sense of "
        "belonging in a multicultural world where all backgrounds are valued and explored with joy."
    ),
    'E': (
        "In our {suburb} {service_type} program, inclusivity and cultural understanding are not "
        "optional extras: they are fundamental to who we are. Our Cultural Education program "
        "introduces children to Indigenous Australian heritage and world cultures through hands-on, "
        "creative experiences including storytelling, art, music, and traditional celebrations. "
        "Children develop empathy, respect, and a genuine appreciation for diversity, growing into "
        "people who are curious about the world and kind within it, wherever they go."
    ),
    'F': (
        "Our {suburb} {service_type} program embraces cultural diversity through an ongoing "
        "Cultural Education program that honours Indigenous heritage, explores world traditions, "
        "and builds genuine respect for people of all backgrounds. Children engage with diverse "
        "cultures through art, music, storytelling, and shared celebrations that make cultural "
        "learning vivid and personal. The result is a community where every child feels seen, "
        "every culture is valued, and every family's identity is celebrated as a gift to the group."
    ),
    'G': (
        "Cultural Education connects children in our {suburb} {service_type} program to the wider "
        "world in meaningful and joyful ways. Through storytelling rooted in Indigenous Australian "
        "perspectives, world music, multicultural art projects, and traditional celebrations, "
        "children develop the cultural awareness and empathy that define global citizens. Our "
        "program fosters a genuine sense of belonging and shared respect, ensuring that every child "
        "and every family feels welcome, honoured, and part of something bigger than themselves."
    ),
    'H': (
        "At our {suburb} {service_type} centre, Cultural Education is a living thread through "
        "everything we do. Children explore Aboriginal and Torres Strait Islander heritage and "
        "world cultures through immersive creative activities, music, storytelling, and "
        "celebrations that make cultural understanding feel genuine rather than theoretical. "
        "We believe that when children grow up knowing how to embrace diversity, they grow into "
        "people who are compassionate, curious, and truly prepared for life in a multicultural world."
    ),
}

# ─────────────────────────────────────────────
# SCHOOL READINESS PROGRAM (8 variants A-H, {suburb} + {service_type} params)
# ─────────────────────────────────────────────
school_readiness = {
    'A': (
        "Our School Readiness Program at our {suburb} {service_type} centre is designed to give "
        "every child the skills, confidence, and love of learning they need to transition "
        "seamlessly to school. Educators design activities and experiences that develop and extend "
        "each child's thinking, skills, interests, and abilities across Language, Independence, "
        "Mathematics, Social & Emotional development, Fine Motor Skills, Music and Movement, "
        "Social Justice, and Health. We welcome you to book a tour and speak with our educators "
        "about how our School Readiness Program can support your child's unique journey."
    ),
    'B': (
        "In our {suburb} {service_type} program, school readiness is not a single program but a "
        "philosophy woven through every experience. Our educators build on each child's interests "
        "and abilities through play-based activities focused on Language, Mathematics, Fine Motor "
        "Skills, Independence, Social & Emotional development, Music and Movement, Social Justice, "
        "and Health. By the time children transition to school, they are not just academically "
        "prepared: they are socially confident, emotionally resilient, and genuinely excited about "
        "the next chapter of their learning journey."
    ),
    'C': (
        "School readiness at our {suburb} {service_type} centre means developing the whole child. "
        "Our educators plan experiences that extend children's thinking and build foundational "
        "skills across writing, science, and mathematics through programs grounded in children's "
        "own interests. Children graduate with confidence in Language, Independence, Mathematics, "
        "Social & Emotional skills, Fine Motor Skills, Music and Movement, Social Justice, and "
        "Health, giving them a strong and joyful foundation for every year of school that follows."
    ),
    'D': (
        "The School Readiness Program at our {suburb} {service_type} centre is built around each "
        "child's unique strengths and emerging interests. Educators introduce topics in writing, "
        "science, and mathematics through play-based experiences that feel engaging and purposeful "
        "for this age group. Outcomes we aim to have graduates confident in include Language, "
        "Independence, Mathematics, Social & Emotional development, Fine Motor Skills, Music and "
        "Movement, Social Justice, and Health. Every child's transition to school is supported "
        "with care, intention, and genuine pride in how far they have come."
    ),
    'E': (
        "At our {suburb} {service_type} centre, the School Readiness Program gives every child "
        "the skills they need to thrive from day one of school. Through intentional, play-based "
        "experiences in writing, science, and mathematics, educators build on each child's "
        "interests and abilities in a way that feels natural and exciting. Our graduates are "
        "confident in Language, Independence, Mathematics, Social & Emotional development, Fine "
        "Motor Skills, Music and Movement, Social Justice, and Health — the full range of skills "
        "that make for a genuinely successful school start."
    ),
    'F': (
        "Confidence is at the heart of our School Readiness Program at our {suburb} {service_type} "
        "centre. When children feel capable, they embrace school with enthusiasm rather than "
        "anxiety. Our educators design experiences that develop Language, Independence, Mathematics, "
        "Social & Emotional skills, Fine Motor Skills, Music and Movement, Social Justice, and "
        "Health through play-based programs rooted in children's own interests and developmental "
        "needs. The transition to school becomes something children and families look forward to "
        "rather than worry about."
    ),
    'G': (
        "Our {suburb} {service_type} program's School Readiness Program recognises that physical "
        "readiness matters as much as academic readiness. Fine Motor Skills are developed through "
        "hands-on creative activities that make writing, drawing, and cutting feel achievable and "
        "fun. Alongside this, our educators build Language, Independence, Mathematics, Social & "
        "Emotional development, Music and Movement, Social Justice, and Health skills so that every "
        "graduate is ready in every sense for the adventure of school."
    ),
    'H': (
        "Cognitive development is the engine of school readiness, and at our {suburb} {service_type} "
        "centre, we nurture children's curiosity, critical thinking, and love of discovery every "
        "single day. Our School Readiness Program builds foundational skills in Language, "
        "Mathematics, and Science through play-based experiences that children find genuinely "
        "exciting. Alongside cognitive development, our graduates are confident in Independence, "
        "Social & Emotional skills, Fine Motor Skills, Music and Movement, Social Justice, and "
        "Health — giving them a complete and joyful foundation for school."
    ),
}

# ─────────────────────────────────────────────
# WHY CHOOSE BULLETS (8 sets A-H, 9 bullets each)
# Each set: list of (label, description) tuples
# First bullet includes {suburb} for uniqueness
# ─────────────────────────────────────────────
why_bullets = {
    'A': [
        ("Child-Centred Approach", "In {suburb}, every child at Flourish ELC is seen as capable, curious, and confident, with a learning environment that makes them feel safe, valued, and empowered every day."),
        ("High-Quality Education", "Programs are guided by the Early Years Learning Framework (EYLF) and meet and exceed the National Quality Framework (NQF), ensuring holistic, play-based learning experiences."),
        ("School Readiness Program", "Children develop social, emotional, cognitive, and self-help skills to transition confidently to school, with readiness embedded in everything we do."),
        ("STEM & Creative Arts Integration", "Critical thinking, problem-solving, imagination, and emotional development are encouraged through hands-on STEM and creative activities every day."),
        ("Sustainability & Cultural Education", "Children learn to care for the environment and embrace Australia's cultural diversity, including Aboriginal and Torres Strait Islander perspectives."),
        ("Healthy Eating & Nutrition", "Garden-to-plate initiatives and mindful mealtime practices support positive food habits and overall wellbeing across five nutritious meals each weekday."),
        ("Passionate Educators", "Our team is committed to reflective practice, ongoing professional development, and upholding the highest child safety standards in everything they do."),
        ("Family Partnership", "Strong collaboration with families ensures every child's growth, milestones, and identity are recognised, celebrated, and supported."),
        ("Lifelong Love of Learning", "We nurture curiosity, exploration, and personal growth, helping every child flourish socially, emotionally, physically, and intellectually."),
    ],
    'B': [
        ("Every Child at the Centre", "Families in {suburb} choose Flourish ELC because every child is genuinely known here: seen as capable, curious, and worthy of an environment that makes them feel safe and empowered."),
        ("Excellence Guided by the EYLF & NQF", "Our programs are shaped by the Early Years Learning Framework (Belonging, Being and Becoming) and meet and exceed the National Quality Framework, ensuring the highest standard of early education."),
        ("Confident School Starters", "Our School Readiness Program progressively builds the social, emotional, cognitive, and self-help skills children need to walk into school feeling genuinely prepared and excited."),
        ("STEM Thinking & Creative Expression", "Hands-on STEM and Creative Arts experiences develop critical thinking, problem-solving, and imagination alongside emotional intelligence and a love of making things."),
        ("Culture & Sustainability Woven In", "Indigenous perspectives, world cultures, and environmental responsibility are embedded in everyday experiences so children grow up as global citizens who care."),
        ("Nourishing Every Child", "Five nutritious meals every weekday, mindful mealtimes, and garden-to-plate experiences build the healthy food habits and physical wellbeing every child deserves."),
        ("Educators Who Go the Extra Mile", "Ongoing professional development and reflective practice mean our educators are always growing, always learning, and always holding the highest standards of child safety."),
        ("Genuine Family Partnership", "We believe families and educators are co-authors of a child's story. Every milestone, achievement, and identity is honoured in genuine and ongoing partnership."),
        ("Growing Lifelong Learners", "Curiosity, exploration, and a deep love of learning are nurtured every day, giving children the social, emotional, physical, and intellectual foundations to flourish for life."),
    ],
    'C': [
        ("Children First, Always", "At Flourish ELC in {suburb}, children are seen as naturally capable and full of wonder. Every space, every relationship, and every experience is designed to make them feel safe, valued, and free to be themselves."),
        ("Programs Built on the EYLF & NQF", "Guided by the Early Years Learning Framework and exceeding the National Quality Framework, our curriculum delivers rich, play-based learning that develops the whole child."),
        ("Ready for School, Ready for Life", "Our School Readiness Program builds cognitive, social, emotional, and self-help skills progressively so children step into school feeling genuinely prepared and proud."),
        ("STEM + Creativity = Brilliant Futures", "Daily STEM and Creative Arts experiences encourage children to think critically, solve problems creatively, and develop the imagination that underpins lifelong success."),
        ("Culture and Environment Matter Here", "Aboriginal and Torres Strait Islander heritage, world cultures, and sustainability practices are celebrated and explored, helping children grow into responsible global citizens."),
        ("Food That Fuels Growing Bodies", "Garden-to-plate meals, mindful eating, and five nutritious daily meals build healthy habits and a genuinely positive relationship with food from the earliest years."),
        ("A Team That Keeps Getting Better", "Our educators engage in reflective practice and ongoing professional development because we believe the best educators never stop learning, just like the children they inspire."),
        ("Families Are Part of Our Community", "Every family at Flourish ELC is a genuine partner. Growth, milestones, and each child's unique identity are celebrated together, because it truly takes a village."),
        ("Curiosity That Lasts a Lifetime", "By honouring every child's natural curiosity and love of discovery, we plant the seeds of a lifelong love of learning that grows long after they leave our centre."),
    ],
    'D': [
        ("A Safe Place to Be Exactly Who You Are", "Families in {suburb} trust Flourish ELC because every child is welcomed exactly as they are, with a learning environment that celebrates individuality and builds genuine confidence."),
        ("Quality You Can See and Feel", "Our programs are guided by the EYLF (Belonging, Being and Becoming) and meet and exceed the NQF, so families know the quality of care and education their children receive is Australia's best."),
        ("School Readiness That Really Works", "Social, emotional, cognitive, and self-help skills are developed through intentional, play-based programs that make school readiness something children genuinely enjoy building."),
        ("Thinking, Creating, Discovering", "STEM Learning and Creative Arts are given real time and real resources at Flourish ELC, sparking the critical thinking, imagination, and problem-solving skills children carry for life."),
        ("Rooted in Culture, Caring for Country", "Indigenous Australian perspectives, world cultures, and a genuine commitment to environmental sustainability are part of everything we do, every day."),
        ("Healthy and Happy from the Inside Out", "Five nutritious meals, garden-to-plate cooking experiences, and mindful mealtime practices build the food habits and physical wellbeing that support children's growth and learning."),
        ("Educators Who Are Truly Passionate", "Every member of our team is deeply committed to ongoing professional development, reflective practice, and the highest standards of child safety and wellbeing."),
        ("Your Family Is Our Family", "Strong, genuine collaboration with families means every child's milestones, achievements, and identity are noticed, celebrated, and woven into their learning journey."),
        ("A Love of Learning That Never Stops", "We plant the seeds of curiosity, exploration, and intellectual joy early, so every child grows up loving to learn, socially, emotionally, physically, and academically."),
    ],
    'E': [
        ("Every Child Is Extraordinary", "In {suburb}, Flourish ELC sees every child as uniquely capable, curious, and full of potential. Our environment is built to help them feel safe, known, and empowered every single day."),
        ("Guided by Australia's Best Frameworks", "Our programs are shaped by the EYLF (Belonging, Being and Becoming) and meet and exceed the National Quality Framework, setting a standard families can truly trust."),
        ("School-Ready in Every Way", "Our School Readiness Program develops the full range of skills children need for school: cognitive, social, emotional, and self-help, through play-based learning that children genuinely love."),
        ("STEM & Arts: A Powerful Combination", "When children build, design, paint, and create, they develop critical thinking, problem-solving, and emotional intelligence alongside the joy and confidence that come from making something real."),
        ("Proud of Where We Come From", "Aboriginal and Torres Strait Islander perspectives are honoured alongside world cultures and sustainability education, because we want children to grow up proud of Australia's diversity and committed to its future."),
        ("Meals That Matter", "Garden-to-plate initiatives, mindful eating practices, and five nutritious meals every weekday build a lifelong healthy relationship with food and support the energy children need to learn and play."),
        ("Committed to Being the Best", "Our educators engage in ongoing professional development and reflective practice, always pushing themselves to deliver the highest quality of care and the highest standards of child safety."),
        ("Families Are Celebrated Here", "We partner closely with families to make sure every child's growth, identity, and milestones are recognised and celebrated as the remarkable achievements they truly are."),
        ("Flourishing, in Every Sense", "Curiosity, wonder, and personal growth are nurtured across every domain: social, emotional, physical, and intellectual, giving every child the foundation to truly flourish for life."),
    ],
    'F': [
        ("Seen, Known, and Valued", "At Flourish ELC in {suburb}, every child is treated as an individual: seen, known by name, celebrated for their strengths, and supported in their growth every single day."),
        ("The Gold Standard in Early Education", "Programs guided by the EYLF and meeting and exceeding the NQF mean families in our community can trust that their children are receiving Australia's highest standard of early education."),
        ("Prepared for School, Excited for Life", "Through our School Readiness Program, children build the social, emotional, cognitive, and self-help skills that make the transition to school a moment of genuine excitement rather than anxiety."),
        ("Hands-On STEM & Creative Learning", "From building blocks to art studios, STEM and Creative Arts experiences give children daily opportunities to think, create, problem-solve, and develop the skills that matter most."),
        ("Celebrating Every Culture, Caring for the Earth", "Cultural diversity, including Aboriginal and Torres Strait Islander perspectives, and environmental sustainability are celebrated and embedded in our everyday curriculum."),
        ("Good Food, Great Habits", "Five nutritious daily meals, garden-to-plate cooking, and mindful mealtimes teach children to love food that loves them back, building positive habits that last a lifetime."),
        ("Educators Who Grow Too", "Our team's commitment to ongoing professional development and reflective practice means children always benefit from educators who are engaged, growing, and giving their very best."),
        ("Partnership From Day One", "We believe parents and carers are essential partners in every child's growth. Milestones, achievements, and each child's unique identity are celebrated together, every step of the way."),
        ("Lifelong Curiosity, Lifelong Success", "By nurturing curiosity and exploration from the earliest years, we help every child develop a love of learning that carries them through school and shapes who they become."),
    ],
    'G': [
        ("A Warm Welcome Every Single Day", "Families in {suburb} choose Flourish ELC because children love coming here. Every child is greeted warmly, celebrated genuinely, and helped to feel like they belong from day one."),
        ("EYLF & NQF: Our North Star", "The Early Years Learning Framework and National Quality Framework guide every decision we make, from how we design our spaces to how we hire and support our educators."),
        ("School Readiness Done Right", "Our School Readiness Program builds social, emotional, cognitive, and self-help skills through intentional play-based experiences that children find genuinely meaningful and enjoyable."),
        ("Creativity and STEM: Every Day", "STEM Learning and Creative Arts are not special events at Flourish ELC: they are part of every day, because critical thinking, imagination, and creativity matter every day."),
        ("Proud Custodians of Culture and Country", "Indigenous perspectives, multicultural education, and environmental sustainability are woven into daily life, because we want children to grow up respectful, aware, and proud."),
        ("Nourishment That Goes Beyond Meals", "Five nutritious daily meals, garden-to-plate experiences, and mindful eating practices nourish both body and mind, building healthy habits that children carry into adulthood."),
        ("Educators Who Are Truly Committed", "Our team's dedication to professional growth, reflective practice, and child safety is unwavering, because the best early learning starts with the best educators."),
        ("Together Every Step", "Flourish ELC families are genuine partners in their child's learning. Every milestone, growth moment, and achievement is shared, celebrated, and built upon together."),
        ("Curiosity That Grows With Them", "We nurture each child's natural curiosity and love of discovery, so the love of learning they develop here follows them through school and throughout their lives."),
    ],
    'H': [
        ("Children Are Our Whole World", "At Flourish ELC in {suburb}, every room, every program, and every decision is built around one simple priority: the children. They are seen, celebrated, and genuinely loved here."),
        ("Quality in Everything We Do", "From EYLF-guided programs to NQF-exceeding standards, the quality of care, education, and environment at Flourish ELC is something families notice and trust from the very first visit."),
        ("Ready, Willing, and Able for School", "Our School Readiness Program gives every child the confidence, skills, and genuine excitement they need to transition to school as a capable, joyful learner."),
        ("STEM + Arts: Where Wonder Lives", "Creative problem-solving, scientific curiosity, and artistic expression are nurtured every day through hands-on STEM and Creative Arts experiences that children find exciting and meaningful."),
        ("Culture, Country, and Community", "We honour Aboriginal and Torres Strait Islander perspectives, celebrate world cultures, and embed environmental stewardship in everything we do, because these values matter to us deeply."),
        ("Food That Fuels, Habits That Last", "Five nutritious daily meals and a garden-to-plate philosophy give children the fuel they need to learn, play, and grow, while building healthy habits that will serve them for life."),
        ("Educators Who Bring Their Best", "Ongoing learning, reflective practice, and an uncompromising commitment to child safety define every member of the Flourish ELC team, every single day."),
        ("Families Are Part of the Story", "Every child's family is a genuine partner at Flourish ELC. We celebrate growth, honour identity, and share milestones together, because a child's story belongs to everyone who loves them."),
        ("Learning That Lasts a Lifetime", "From the earliest age, we nurture the curiosity, resilience, and love of discovery that turn into a lifelong love of learning, socially, emotionally, physically, and intellectually."),
    ],
}

# ─────────────────────────────────────────────
# ENVIRONMENT (8 variants A-H, {suburb} param)
# ─────────────────────────────────────────────
environment = {
    'A': {
        'heading': 'Purpose-Built, NQF-Aligned, and Always Welcoming',
        'body': (
            "Our {suburb} facility is purpose-built for children aged 6 weeks to school age. Every "
            "room, every outdoor space, and every learning corner has been thoughtfully designed "
            "with children's safety, wellbeing, and development in mind: not repurposed from "
            "something else, but built from the ground up for early learning. We meet and exceed "
            "the National Quality Framework (NQF), which means our educators, programs, and "
            "environments are assessed against Australia's highest national standards for early "
            "education and care.\n\n"
            "Our educators are qualified professionals who bring both formal training and genuine "
            "warmth to every interaction. They take the time to know each child as an individual: "
            "what makes them curious, what makes them feel safe, and what they are ready to explore "
            "next. We are open Monday to Friday, 7:00am to 6:00pm, and our parent communication "
            "app delivers real-time updates and photos throughout the day."
        )
    },
    'B': {
        'heading': 'Qualified Educators in a Space Designed for Children',
        'body': (
            "Flourish ELC's {suburb} centre is purpose-built. Every space, from indoor learning "
            "rooms to outdoor play areas, has been designed for the specific needs of young children "
            "at different developmental stages. Our qualified educators meet and exceed the National "
            "Quality Framework (NQF), and their dedication shows in the way they engage with "
            "children: attentively, warmly, and with genuine professional skill.\n\n"
            "The NQF is not just a compliance standard at Flourish ELC: it is a reflection of our "
            "values. Every element of the centre, the spaces, the programs, and the relationships "
            "we build with families, has been assessed against Australia's national quality standards. "
            "We are open 7:00am to 6:00pm Monday to Friday, and our parent communication app "
            "keeps families genuinely connected throughout the day."
        )
    },
    'C': {
        'heading': 'A Safe, Warm Environment That Meets the NQF',
        'body': (
            "The Flourish ELC {suburb} environment has been purpose-built for children from 6 weeks "
            "to school age. Safe, bright, and welcoming, it is a place where children feel at home "
            "from the moment they arrive. Meeting and exceeding the National Quality Framework (NQF) "
            "is not a target we reach for: it is the standard we live every single day. Our "
            "educators hold recognised qualifications and bring a passion for early childhood to "
            "every interaction.\n\n"
            "Children spend their days in spaces designed to inspire movement, creativity, and "
            "connection. Outdoor areas give children room to run, explore, and experience nature "
            "as part of their learning. Indoor environments are warm and purposeful, with activities "
            "that reflect children's interests and developmental needs. We are open Monday to Friday, "
            "7:00am to 6:00pm, and our parent communication app sends real-time updates and photos "
            "throughout the day."
        )
    },
    'D': {
        'heading': 'An Environment Built Around Children, Open Every Weekday',
        'body': (
            "Everything about the Flourish ELC {suburb} facility has been built with children in "
            "mind. Our purpose-built centre meets and exceeds the National Quality Framework (NQF), "
            "giving families confidence that the care, the people, and the programs their children "
            "experience have been assessed against Australia's highest national standards. That "
            "standard is one we are genuinely proud of and work hard to uphold every day.\n\n"
            "Our educators are at the heart of everything. They are qualified, warm, and genuinely "
            "committed to the children in their care. They observe carefully, listen closely, and "
            "respond thoughtfully, creating a daily experience that is both nurturing and stimulating. "
            "We are open 7:00am to 6:00pm Monday to Friday, with a parent communication app that "
            "delivers real-time photos and updates throughout the day."
        )
    },
    'E': {
        'heading': 'Outdoor Spaces and Indoor Warmth for Growing Children',
        'body': (
            "At Flourish ELC {suburb}, the learning environment extends beyond four walls. Our "
            "purpose-built outdoor spaces give children room to run, investigate, garden, and "
            "connect with the natural world every day, because outdoor play is not a break from "
            "learning: it is one of the richest forms of it. Indoor environments are designed "
            "with equal care: warm, purposeful spaces that invite creativity, movement, and "
            "collaboration at every turn.\n\n"
            "Our qualified educators bring expertise and genuine warmth to every indoor and outdoor "
            "experience, meeting and exceeding the National Quality Framework (NQF) in everything "
            "they do. We are open Monday to Friday, 7:00am to 6:00pm, and our parent communication "
            "app ensures families are always connected to their child's day with real-time updates "
            "and photos throughout the day."
        )
    },
    'F': {
        'heading': 'A Thoughtful Daily Routine in a Purpose-Built Centre',
        'body': (
            "The Flourish ELC {suburb} centre runs on a thoughtful daily routine that balances "
            "structured learning with free play, indoor exploration with outdoor adventure, and "
            "group activity with quiet individual time. This rhythm gives children a sense of "
            "security and predictability that is foundational to their emotional wellbeing and "
            "readiness to learn. Every element of the routine reflects our understanding of child "
            "development and our commitment to each child's individual needs.\n\n"
            "Our purpose-built facility meets and exceeds the National Quality Framework (NQF), "
            "and our qualified educators bring both professional expertise and genuine care to every "
            "part of the day. We are open 7:00am to 6:00pm, Monday to Friday. Our parent "
            "communication app keeps families connected with real-time updates throughout the day."
        )
    },
    'G': {
        'heading': 'A Community Where Children and Families Feel at Home',
        'body': (
            "Flourish ELC {suburb} is more than a centre: it is a community. Families here "
            "describe a warmth and sense of belonging that is genuinely hard to find elsewhere, "
            "and our qualified educators work every day to make that feeling real. Our "
            "purpose-built facility meets and exceeds the National Quality Framework (NQF), "
            "which means every aspect of what we offer has been assessed against Australia's "
            "highest standards for early education and care.\n\n"
            "Our parent communication app is one of the things families love most: real-time "
            "photos, updates, and observations throughout the day so you are never far from "
            "your child's world. We are open Monday to Friday, 7:00am to 6:00pm, giving "
            "families a generous and flexible window for drop-off and pick-up every weekday."
        )
    },
    'H': {
        'heading': 'Wellbeing, Settling In, and Feeling Truly Welcome',
        'body': (
            "At Flourish ELC {suburb}, every new child is welcomed with care, patience, and a "
            "genuine commitment to making the settling-in process as smooth and positive as "
            "possible. Our educators take the time to get to know each child before they "
            "start, building the trust and familiarity that help children feel secure from "
            "the very beginning. A child who feels safe and settled is a child who is free "
            "to explore, connect, and learn.\n\n"
            "Our purpose-built facility meets and exceeds the National Quality Framework (NQF), "
            "and our qualified, passionate educators are the heart of everything we offer. "
            "We are open 7:00am to 6:00pm Monday to Friday, with a parent communication app "
            "that delivers real-time updates and photos so families always feel close to their "
            "child's world, even on the busiest of days."
        )
    },
}


# ─────────────────────────────────────────────
# PAGES
# ─────────────────────────────────────────────
pages = [
  # ── NSW CHILDCARE (group_offset=0, var_key A-G) ──
  {
    "num": 1, "type": "childcare", "state": "NSW", "suburb": "Oakville",
    "prog_type": "childcare", "prog_key": "A", "munch_key": "A", "var_key": "A", "var": 0,
    "url": "https://flourishelc.com.au/daycare-childcare-oakville-nsw/",
    "primary": "childcare oakville", "secondary": "daycare oakville",
    "meta_title": "Childcare Oakville | Flourish Early Learning Centre",
    "meta_desc": "Warm, purpose-built childcare in Oakville with qualified educators, five daily meals, and programs guided by the EYLF. Book your tour today.",
    "h1": "Childcare Oakville Families Love: Flourish Early Learning Centre",
    "intro": "Flourish Early Learning Centre is a purpose-built childcare in Oakville designed around children thriving in every sense of the word. Guided by the Early Years Learning Framework (EYLF): Belonging, Being and Becoming, our qualified educators bring genuine passion to every single day. Oakville families choose us because we go beyond care: we create a place where children light up, grow, and flourish.",
    "why_heading": "Why Choose Flourish ELC for Childcare in Oakville",
    "faqs": [
        ("What age groups does Flourish ELC in Oakville cater for?", "We welcome children from 6 weeks through to school age, with dedicated programs and environments for each developmental stage."),
        ("What meals are served each day at the Oakville centre?", "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. All dietary requirements and allergies are carefully accommodated."),
        ("What are your opening hours in Oakville?", "Our centre is open Monday to Friday, 7:00am to 6:00pm, giving families a flexible window for drop-off and pick-up."),
        ("How does the parent communication app work?", "The app delivers real-time photos, updates, and observations throughout the day so you always feel genuinely connected to your child's world, even during the busiest workday."),
        ("What is the EYLF and how does it guide your programs?", "The Early Years Learning Framework (Belonging, Being and Becoming) shapes how our educators plan, observe, and respond to children, always placing identity, wellbeing, and engagement at the centre."),
        ("How does Flourish ELC support school readiness in Oakville?", "Our school readiness programs develop language, numeracy, independence, and social-emotional skills progressively so children transition to kindergarten with confidence, capability, and genuine excitement."),
        ("What qualifications do your educators hold?", "All educators hold recognised early childhood qualifications and meet and exceed National Quality Framework (NQF) standards for professional development and educator ratios."),
        ("How do I book a tour of your Oakville centre?", "Contact us through our website or by phone to arrange a tour. We warmly welcome all families to come and experience Flourish ELC firsthand before enrolling."),
    ],
  },
  {
    "num": 2, "type": "childcare", "state": "NSW", "suburb": "Box Hill",
    "prog_type": "childcare", "prog_key": "B", "munch_key": "B", "var_key": "B", "var": 1,
    "url": "https://flourishelc.com.au/daycare-childcare-box-hill-nsw/",
    "primary": "childcare near box hill", "secondary": "daycare near box hill",
    "meta_title": "Childcare Near Box Hill NSW | Flourish Early Learning Centre",
    "meta_desc": "Searching for childcare near Box Hill? Flourish ELC is a purpose-built centre with qualified educators, five daily meals, and EYLF-guided programs. Book your tour today.",
    "h1": "Childcare Near Box Hill: Where Every Child Belongs, Grows and Flourishes",
    "intro": "Flourish Early Learning Centre is a warm, purpose-built early learning centre conveniently located for families in and around Box Hill. Guided by the Early Years Learning Framework (EYLF): Belonging, Being and Becoming, our qualified, passionate educators bring that philosophy to life every single day. If you are looking for childcare near Box Hill that feels like a second home for your child, you have found it.",
    "why_heading": "Why Choose Flourish ELC for Childcare Near Box Hill",
    "faqs": [
        ("Do you cater for newborns at your centre near Box Hill?", "Yes, we welcome children from as young as 6 weeks. Our nursery spaces are purpose-built to support the specific needs of infants and very young babies."),
        ("What kinds of meals are provided throughout the day near Box Hill?", "Children receive five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. All allergies and dietary requirements are always managed with care."),
        ("How long is the childcare day near Box Hill?", "Our centre is open 7:00am to 6:00pm, Monday to Friday, giving families a full and flexible window of care that fits around most working schedules."),
        ("How are families kept updated while children are in care near Box Hill?", "Our parent communication app sends real-time photos and updates so families always know what their child is doing, experiencing, and enjoying throughout the day."),
        ("What is the National Quality Framework and why does it matter?", "The NQF is Australia's national quality standard for early childhood education. Meeting and exceeding it means our programs, educators, and environments are held to the country's highest benchmarks."),
        ("How does Flourish ELC support children with allergies near Box Hill?", "Dietary requirements and allergies are documented at enrolment and managed carefully every day by our kitchen and educator teams, ensuring every child is safe and well nourished."),
        ("What outdoor play opportunities are available near Box Hill?", "Our purpose-built outdoor spaces give children daily opportunities to run, explore, garden, and engage with nature as a genuine part of their learning experience."),
        ("Can I visit the centre near Box Hill before my child starts?", "Absolutely. We encourage all families to visit before enrolling so they can meet our educators, see our spaces, and feel confident their child will thrive at Flourish ELC."),
    ],
  },
  {
    "num": 3, "type": "childcare", "state": "NSW", "suburb": "Vineyard",
    "prog_type": "childcare", "prog_key": "C", "munch_key": "C", "var_key": "C", "var": 2,
    "url": "https://flourishelc.com.au/daycare-childcare-vineyard-nsw/",
    "primary": "childcare near vineyard", "secondary": "daycare near vineyard",
    "meta_title": "Childcare Near Vineyard NSW | Flourish Early Learning Centre",
    "meta_desc": "Looking for childcare near Vineyard? Flourish ELC is purpose-built, guided by the EYLF, and offers five daily meals and qualified educators. Book your tour today.",
    "h1": "Childcare Near Vineyard That Puts Children First",
    "intro": "Flourish Early Learning Centre is a purpose-built early learning centre serving families in and around Vineyard. Built around the Early Years Learning Framework (EYLF): Belonging, Being and Becoming, our centre is a place where children are celebrated for who they are, supported in what they are becoming, and given every opportunity to belong. Our qualified, passionate educators make Flourish ELC a truly special place for Vineyard families.",
    "why_heading": "Why Choose Flourish ELC for Childcare Near Vineyard",
    "faqs": [
        ("What is the youngest age you accept for childcare near Vineyard?", "We welcome children from 6 weeks. Our infant rooms are purpose-built and staffed by educators with specific expertise in caring for very young children."),
        ("Are there healthy food options provided throughout the day near Vineyard?", "Five nutritious meals are served every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. Our menu is designed to support healthy development and growth."),
        ("What hours is the Flourish ELC centre open near Vineyard?", "We are open Monday to Friday, 7:00am to 6:00pm, a generous window that accommodates most working hours and commute times."),
        ("How does Flourish ELC communicate with parents during the day near Vineyard?", "Our parent communication app delivers real-time updates, photos, and observations so families are always connected to their child's day, no matter where they are."),
        ("How are children supported in developing independence at Flourish ELC near Vineyard?", "Independence is nurtured gently through age-appropriate choices, self-help skill building, and educator support that encourages children to try, discover, and feel proud of their abilities."),
        ("Does Flourish ELC incorporate STEM learning for young children near Vineyard?", "Yes, STEM Learning is woven into daily experiences through hands-on exploration, building, experimenting, and problem-solving that children find genuinely exciting and engaging."),
        ("How does Flourish ELC help children settle in when they start near Vineyard?", "We offer flexible settling-in visits and work closely with families to build routines that help each child feel safe and comfortable from their very first day at our centre."),
        ("How is the curriculum planned for children near Vineyard?", "Our educators observe children closely, follow their interests, and plan EYLF-guided experiences that are personalised, engaging, and developmentally appropriate for every child."),
    ],
  },
  {
    "num": 4, "type": "childcare", "state": "NSW", "suburb": "Gables",
    "prog_type": "childcare", "prog_key": "D", "munch_key": "D", "var_key": "D", "var": 3,
    "url": "https://flourishelc.com.au/daycare-childcare-gables-nsw/",
    "primary": "childcare near gables", "secondary": "daycare near gables",
    "meta_title": "Childcare Near Gables NSW | Flourish Early Learning Centre",
    "meta_desc": "Warm, purpose-built childcare near Gables with EYLF-guided programs, qualified educators, and five nutritious daily meals. Book your family tour today.",
    "h1": "Childcare Near Gables: A Place Where Children Truly Flourish",
    "intro": "Flourish Early Learning Centre is a purpose-built childcare centre welcoming families in and around Gables. Everything here, from the way our educators engage with children to the way our spaces are designed, is shaped by the Early Years Learning Framework (EYLF): Belonging, Being and Becoming. Children at Flourish ELC do not just attend a centre: they belong to a warm, vibrant community that celebrates every step of their journey.",
    "why_heading": "Why Choose Flourish ELC for Childcare Near Gables",
    "faqs": [
        ("At what age can children start childcare near Gables?", "We welcome children from 6 weeks through to school age, with age-appropriate programs, rooms, and outdoor spaces for every developmental stage."),
        ("What food and meals does Flourish ELC provide near Gables?", "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. We always accommodate dietary requirements and allergies with care and consistency."),
        ("What are the childcare centre's operating hours near Gables?", "Monday to Friday, 7:00am to 6:00pm. We also have a parent communication app so families are connected throughout the day with real-time photos and updates."),
        ("What is the best way to stay in touch with my child's educators near Gables?", "Our parent communication app provides real-time updates and photos, and our educators are available for conversations at drop-off, pick-up, and via the app messaging feature."),
        ("How does Flourish ELC celebrate cultural diversity near Gables?", "Our Cultural Education program incorporates Indigenous Australian perspectives, world cultures, and multicultural celebrations through storytelling, art, music, and shared experiences every week."),
        ("What does the school readiness program involve for childcare children near Gables?", "Intentional, play-based programs that build language, maths, fine motor skills, independence, and social-emotional abilities so children are genuinely prepared for kindergarten."),
        ("Are CCS (Child Care Subsidy) payments applicable at Flourish ELC near Gables?", "Yes, Flourish ELC is an approved provider, meaning eligible families can access the Child Care Subsidy to help with the cost of care. Our team can guide you through the process."),
        ("How do educators support children's social and emotional development near Gables?", "Educators build warm, trusting relationships, use positive guidance strategies, and create a safe environment where children feel comfortable expressing feelings and building friendships."),
    ],
  },
  {
    "num": 5, "type": "childcare", "state": "NSW", "suburb": "Grantham Farm",
    "prog_type": "childcare", "prog_key": "E", "munch_key": "E", "var_key": "E", "var": 0,
    "url": "https://flourishelc.com.au/daycare-childcare-grantham-farm-nsw/",
    "primary": "childcare near grantham farm", "secondary": "daycare near grantham farm",
    "meta_title": "Childcare Near Grantham Farm | Flourish Early Learning Centre",
    "meta_desc": "Purpose-built childcare near Grantham Farm with EYLF-guided learning, qualified educators, five daily meals, and a real-time parent app. Book your tour today.",
    "h1": "Childcare Near Grantham Farm: Warm, Qualified, and Built for Growing Families",
    "intro": "Families near Grantham Farm looking for childcare that goes beyond the basics will find exactly that at Flourish Early Learning Centre. Our purpose-built centre is guided by the Early Years Learning Framework (EYLF): Belonging, Being and Becoming, and staffed by qualified educators who genuinely love what they do. At Flourish ELC, children are known, celebrated, and given every opportunity to thrive.",
    "why_heading": "Why Choose Flourish ELC for Childcare Near Grantham Farm",
    "faqs": [
        ("What age does Flourish ELC accept children from near Grantham Farm?", "Children are welcome from 6 weeks of age. Our purpose-built centre has dedicated rooms and programs for infants, toddlers, and preschool-aged children through to school age."),
        ("What nutritional support does Flourish ELC provide near Grantham Farm?", "Five nutritious meals every weekday, planned with children's health and development in mind. All dietary requirements, allergies, and cultural food preferences are carefully accommodated."),
        ("What childcare hours are available near Grantham Farm?", "We are open Monday to Friday, 7:00am to 6:00pm. Our parent communication app keeps you connected to your child's day with real-time photos and updates."),
        ("What creative arts programs are on offer at Flourish ELC near Grantham Farm?", "Creative Arts are part of every day at Flourish ELC, with children engaging in painting, drawing, music, craft, and collaborative projects that build imagination, expression, and fine motor skills."),
        ("How does Flourish ELC approach language development near Grantham Farm?", "Language is woven through every experience: stories, songs, conversations, and collaborative play all build vocabulary, communication skills, and the confidence to express ideas and feelings."),
        ("What happens when a child is unwell and cannot attend near Grantham Farm?", "We ask families to keep unwell children at home to protect the health of all children and staff. Our team communicates openly about illness management and exclusion guidelines."),
        ("How are staff ratios maintained at Flourish ELC near Grantham Farm?", "We meet and exceed NQF-mandated educator-to-child ratios at all times, ensuring every child receives the individual attention and supervision they need throughout the day."),
        ("How can I learn more about enrolling near Grantham Farm?", "We welcome families to book a tour, meet our educators, and explore our centre firsthand. Contact us through our website or phone to arrange a convenient time to visit."),
    ],
  },
  {
    "num": 6, "type": "childcare", "state": "NSW", "suburb": "Riverstone",
    "prog_type": "childcare", "prog_key": "F", "munch_key": "F", "var_key": "F", "var": 1,
    "url": "https://flourishelc.com.au/daycare-childcare-riverstone-nsw/",
    "primary": "childcare near riverstone", "secondary": "daycare near riverstone",
    "meta_title": "Childcare Near Riverstone NSW | Flourish Early Learning Centre",
    "meta_desc": "Warm, purpose-built childcare near Riverstone with EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour.",
    "h1": "Childcare Near Riverstone: Belonging, Growing, and Flourishing Together",
    "intro": "Flourish Early Learning Centre offers families near Riverstone a warm, purpose-built early learning environment where children are known by name, celebrated for who they are, and gently guided toward all they are becoming. Our programs are shaped by the Early Years Learning Framework (EYLF): Belonging, Being and Becoming, and delivered by qualified educators who bring both skill and heart to every day.",
    "why_heading": "Why Choose Flourish ELC for Childcare Near Riverstone",
    "faqs": [
        ("Is there childcare available for babies near Riverstone at Flourish ELC?", "Yes, we welcome babies from 6 weeks of age. Our infant rooms are specifically designed and staffed to provide the nurturing, responsive care that very young children need."),
        ("How many meals does Flourish ELC provide each day near Riverstone?", "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. Our kitchen accommodates all dietary requirements and allergies with care."),
        ("When is Flourish ELC open for childcare near Riverstone?", "Our centre is open Monday to Friday, 7:00am to 6:00pm, providing families with a full working-day window of care."),
        ("How does the Flourish ELC parent app work near Riverstone?", "The app sends real-time photos, milestone updates, and daily observations from educators, giving you a genuine window into what your child is doing and learning each day."),
        ("How does Flourish ELC incorporate music and movement near Riverstone?", "Music and Movement is part of every day, with children singing, dancing, playing instruments, and moving their bodies in ways that build rhythm, coordination, confidence, and creative expression."),
        ("How are children's individual learning needs supported near Riverstone?", "Our educators observe each child closely, plan individually tailored experiences, and work with families to ensure every child's unique strengths, interests, and needs are recognised and supported."),
        ("Does Flourish ELC have a waiting list near Riverstone?", "Places are in demand, so we encourage families to enquire early. Contact our centre to discuss availability and to add your child to our enrolment waitlist if needed."),
        ("What social justice education is offered at Flourish ELC near Riverstone?", "Social Justice is woven into our curriculum through discussions, stories, and activities that help children understand fairness, empathy, and the importance of standing up for what is right."),
    ],
  },
  {
    "num": 7, "type": "childcare", "state": "NSW", "suburb": "Mulgrave",
    "prog_type": "childcare", "prog_key": "G", "munch_key": "G", "var_key": "G", "var": 2,
    "url": "https://flourishelc.com.au/daycare-childcare-mulgrave-nsw/",
    "primary": "childcare near mulgrave", "secondary": "daycare near mulgrave",
    "meta_title": "Childcare Near Mulgrave NSW | Flourish Early Learning Centre",
    "meta_desc": "Looking for childcare near Mulgrave? Flourish ELC is purpose-built with EYLF-guided programs, five nutritious daily meals, and qualified educators. Book your tour.",
    "h1": "Childcare Near Mulgrave That Families Trust and Children Love",
    "intro": "Flourish Early Learning Centre is a purpose-built early learning centre welcoming families near Mulgrave into a warm, vibrant community. Guided by the Early Years Learning Framework (EYLF): Belonging, Being and Becoming, our centre is a place where children feel secure, excited, and free to explore who they are. Our qualified, passionate educators are the heart of everything we do.",
    "why_heading": "Why Choose Flourish ELC for Childcare Near Mulgrave",
    "faqs": [
        ("What ages do you cater for at your childcare near Mulgrave?", "We welcome children from 6 weeks through to school age, with purpose-built rooms, programs, and outdoor spaces designed for every developmental stage."),
        ("Are all meals included in the childcare fees near Mulgrave?", "Yes, five nutritious meals are included every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. Dietary requirements and allergies are always carefully managed."),
        ("What are the opening times at Flourish ELC near Mulgrave?", "Monday to Friday, 7:00am to 6:00pm. We also use a parent communication app to keep families informed and connected throughout the day."),
        ("How does Flourish ELC build fine motor skills near Mulgrave?", "Fine Motor Skills are developed through daily drawing, cutting, threading, clay, painting, and building activities, all designed to build the physical dexterity children need for writing in an enjoyable way."),
        ("What health and wellbeing practices are followed at Flourish ELC near Mulgrave?", "Our approach includes regular outdoor play, the Munch & Move program, five nutritious daily meals, handwashing routines, and educators trained in first aid and child safety."),
        ("How is exercise incorporated into the childcare program near Mulgrave?", "Physical activity is part of every day through outdoor play, movement games, dance, and structured activities that keep children active and channel their energy positively."),
        ("How does Flourish ELC communicate children's progress near Mulgrave?", "Educators document learning through observations, photos, and learning stories shared via our parent communication app, giving families genuine insight into their child's development and growth."),
        ("What should my child bring to childcare near Mulgrave?", "We recommend a spare change of clothes, a hat for outdoor play, a water bottle, and any comfort items for rest time. We provide all meals, sunscreen, and nappy products for enrolled children."),
    ],
  },
  # ── NSW PRESCHOOL (group_offset=1, var_key B-H) ──
  {
    "num": 8, "type": "preschool", "state": "NSW", "suburb": "Oakville",
    "prog_type": "preschool", "prog_key": "A", "munch_key": "A", "var_key": "B", "var": 3,
    "url": "https://flourishelc.com.au/preschool-oakville-nsw/",
    "primary": "preschool oakville", "secondary": "",
    "meta_title": "Preschool Oakville NSW | Flourish Early Learning Centre",
    "meta_desc": "Looking for a preschool in Oakville? Flourish ELC offers EYLF-guided programs, qualified educators, five daily meals, and a warm, purpose-built environment. Enquire today.",
    "h1": "Preschool Oakville Families Choose: Flourish Early Learning Centre",
    "intro": "Flourish Early Learning Centre is a warm, purpose-built preschool in Oakville where children step into a world designed entirely for them. Our preschool programs are guided by the Early Years Learning Framework (EYLF): Belonging, Being and Becoming, and delivered by qualified educators who genuinely care about every child's growth. Families in Oakville choose Flourish ELC because it is the kind of place that children look forward to, and parents feel proud of.",
    "why_heading": "Why Choose Flourish ELC for Preschool in Oakville",
    "faqs": [
        ("What age is the preschool program designed for at Flourish ELC Oakville?", "Our preschool program is designed for children aged 3 to 5. Our full centre also welcomes children from 6 weeks through to school age across all age groups."),
        ("What meals are included in the preschool day at Oakville?", "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. All dietary requirements and allergies are always managed with care and consistency."),
        ("What are the preschool hours at Flourish ELC in Oakville?", "We are open Monday to Friday, 7:00am to 6:00pm, giving preschool families in Oakville a generous window for drop-off and pick-up."),
        ("How does the parent communication app support preschool families in Oakville?", "The app delivers real-time photos, updates, and observations from educators throughout the day so you always have a genuine window into your child's preschool experience."),
        ("How does Flourish ELC prepare preschoolers in Oakville for kindergarten?", "Our school readiness programs build literacy, numeracy, independence, and social-emotional skills progressively throughout the preschool years so children transition to kindergarten feeling confident and genuinely ready."),
        ("What does play-based learning look like in the Oakville preschool program?", "Play-based learning means children build skills through joyful, purposeful experiences: building, creating, investigating, and collaborating in ways that feel exciting and natural for this age group."),
        ("How does Flourish ELC involve families in preschool learning at Oakville?", "Families are genuine partners at Flourish ELC. We share learning stories, invite families to centre events, and hold regular conversations to ensure parents feel informed and part of their child's preschool journey."),
        ("What STEM experiences do preschoolers access in Oakville?", "Preschoolers engage in STEM Learning through hands-on building, simple experiments, nature investigation, and problem-solving activities that spark curiosity and develop critical thinking skills."),
    ],
  },
  {
    "num": 9, "type": "preschool", "state": "NSW", "suburb": "Box Hill",
    "prog_type": "preschool", "prog_key": "B", "munch_key": "B", "var_key": "C", "var": 0,
    "url": "https://flourishelc.com.au/preschool-box-hill-nsw/",
    "primary": "preschool near box hill", "secondary": "",
    "meta_title": "Preschool Near Box Hill NSW | Flourish Early Learning Centre",
    "meta_desc": "Searching for a preschool near Box Hill? Flourish ELC is purpose-built with qualified educators, EYLF-guided programs, and five daily meals. Book your tour today.",
    "h1": "Preschool Near Box Hill: Where Children Grow Into Themselves",
    "intro": "Flourish Early Learning Centre is a purpose-built preschool near Box Hill where the Early Years Learning Framework (EYLF): Belonging, Being and Becoming, is not just a document on a wall but the living philosophy behind every experience. Our qualified, passionate educators know each child by name, by interest, and by strength. Families near Box Hill choose Flourish ELC because their children flourish here in every sense of the word.",
    "why_heading": "Why Choose Flourish ELC for Preschool Near Box Hill",
    "faqs": [
        ("What age range does the Box Hill preschool program cover?", "Our preschool program caters for children aged 3 to 5, within a full centre that welcomes children from 6 weeks through to school age."),
        ("How many meals are provided during preschool near Box Hill?", "Five nutritious meals are served every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. Our kitchen team always accommodates dietary requirements and allergies."),
        ("How long is the preschool day near Box Hill?", "Our centre is open Monday to Friday, 7:00am to 6:00pm, giving families a full and flexible day of preschool care and learning."),
        ("What qualifications do preschool educators hold near Box Hill?", "Our educators hold recognised early childhood qualifications and meet and exceed the National Quality Framework (NQF) standards, bringing both professional expertise and genuine warmth to every preschool session."),
        ("How is the EYLF incorporated into the preschool program near Box Hill?", "The EYLF (Belonging, Being and Becoming) guides how our educators plan preschool experiences, observe children's growth, and respond to each child's individual interests and developmental needs."),
        ("What creative arts experiences are available in the preschool program near Box Hill?", "Creative Arts are part of every preschool day: painting, drawing, craft, sculpture, music, and dramatic play give children expressive outlets that build confidence, fine motor skills, and a love of making."),
        ("How does Flourish ELC support children starting preschool for the first time near Box Hill?", "Our educators meet each new child and family before they start, offering settling-in visits and building routines that help children feel safe, excited, and confident from their very first preschool day."),
        ("How are preschool children's milestones tracked and shared near Box Hill?", "Educators document learning through observations, photos, and learning stories that are shared with families via our parent communication app, giving parents an ongoing picture of their child's preschool growth."),
    ],
  },
  {
    "num": 10, "type": "preschool", "state": "NSW", "suburb": "Vineyard",
    "prog_type": "preschool", "prog_key": "C", "munch_key": "C", "var_key": "D", "var": 1,
    "url": "https://flourishelc.com.au/preschool-vineyard-nsw/",
    "primary": "preschool near vineyard", "secondary": "",
    "meta_title": "Preschool Near Vineyard NSW | Flourish Early Learning Centre",
    "meta_desc": "Looking for a preschool near Vineyard? Flourish ELC is purpose-built with EYLF-guided programs, qualified educators, and five daily nutritious meals. Enquire today.",
    "h1": "Preschool Near Vineyard: A Warm Place Where Learning Comes Alive",
    "intro": "Flourish Early Learning Centre is a purpose-built preschool near Vineyard where the Early Years Learning Framework (EYLF): Belonging, Being and Becoming, shapes every experience, every relationship, and every room. Our qualified educators bring knowledge and genuine warmth to the preschool years, creating an environment where children discover their strengths, build their confidence, and develop a love of learning that will stay with them for life.",
    "why_heading": "Why Choose Flourish ELC for Preschool Near Vineyard",
    "faqs": [
        ("Is the Flourish ELC preschool near Vineyard suitable for 3 year olds?", "Yes, our preschool program is designed for children aged 3 to 5. Younger children from 6 weeks are also welcome across our other age groups and programs."),
        ("What nutritional support is provided during preschool near Vineyard?", "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. All allergies and dietary requirements are documented and managed carefully every day."),
        ("When is the preschool program available near Vineyard?", "Monday to Friday, 7:00am to 6:00pm. Families can also track their child's day in real time through our parent communication app."),
        ("How is social and emotional development supported in the preschool program near Vineyard?", "Educators build warm, trusting relationships with every preschooler, use positive guidance strategies, and create a safe space where children can express feelings, navigate friendships, and grow emotionally."),
        ("What school readiness activities happen in the preschool program near Vineyard?", "Writing, mathematics, science, language, and independence skills are woven into the play-based preschool curriculum every day, always connected to children's genuine interests and curiosities."),
        ("How does the Flourish ELC preschool near Vineyard celebrate Indigenous culture?", "Aboriginal and Torres Strait Islander perspectives are embedded in our Cultural Education program through storytelling, art, music, and celebrations that help children develop respect for Indigenous heritage."),
        ("What music and movement activities are part of preschool near Vineyard?", "Music and Movement is a daily feature: children sing, dance, play percussion instruments, and move their bodies in ways that build rhythm, coordination, and genuine joy."),
        ("How can I arrange a preschool tour near Vineyard?", "Contact us through our website or phone to book a tour. We love welcoming families to see the centre, meet our educators, and experience the warm Flourish ELC environment firsthand."),
    ],
  },
  {
    "num": 11, "type": "preschool", "state": "NSW", "suburb": "Gables",
    "prog_type": "preschool", "prog_key": "D", "munch_key": "D", "var_key": "E", "var": 2,
    "url": "https://flourishelc.com.au/preschool-gables-nsw/",
    "primary": "preschool near gables", "secondary": "",
    "meta_title": "Preschool Near Gables NSW | Flourish Early Learning Centre",
    "meta_desc": "Looking for a preschool near Gables? Flourish ELC is purpose-built with EYLF-guided learning, qualified educators, and five daily meals. Book your tour today.",
    "h1": "Preschool Near Gables: Where Little Learners Truly Belong",
    "intro": "Flourish Early Learning Centre is a warm, purpose-built preschool near Gables where the Early Years Learning Framework (EYLF): Belonging, Being and Becoming, is more than a framework: it is the spirit of everything we do. Our qualified, passionate educators create an environment where preschoolers feel safe, celebrated, and inspired to explore the world around them. Families near Gables choose Flourish ELC because their children thrive here, and they can see it every day.",
    "why_heading": "Why Choose Flourish ELC for Preschool Near Gables",
    "faqs": [
        ("What is the age range for preschool at Flourish ELC near Gables?", "3 to 5 years. Our full centre caters for children from 6 weeks right through to school age, with dedicated programs for each stage of early childhood."),
        ("Are meals provided as part of the preschool program near Gables?", "Yes, five nutritious meals are served every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. Allergies and dietary requirements are always managed with care."),
        ("What hours does the preschool program run near Gables?", "Our centre is open Monday to Friday, 7:00am to 6:00pm. We keep families connected throughout the preschool day with our real-time parent communication app."),
        ("How does Flourish ELC build independence in preschoolers near Gables?", "Independence is nurtured through age-appropriate self-help tasks, choices within the preschool routine, and educator encouragement that helps children feel capable, confident, and proud of managing themselves."),
        ("What does the NQF mean for the preschool program near Gables?", "The National Quality Framework sets Australia's quality standards for early education. Exceeding it means our preschool educators, programs, and environment are held to the highest national benchmarks."),
        ("How is sustainability education incorporated into the preschool program near Gables?", "Children learn to care for the environment through gardening, composting, recycling activities, and discussions about caring for Country, woven into the daily preschool routine in meaningful ways."),
        ("What fine motor skill activities are part of the preschool experience near Gables?", "Drawing, cutting, threading, clay, painting, and writing activities are all part of our daily preschool program, helping children develop the physical dexterity needed for school in fun and creative ways."),
        ("How are preschool families kept informed near Gables?", "Our parent communication app delivers real-time photos and educator updates throughout the day, and our team is always available at drop-off and pick-up for personal conversations about your child's preschool progress."),
    ],
  },
  {
    "num": 12, "type": "preschool", "state": "NSW", "suburb": "Grantham Farm",
    "prog_type": "preschool", "prog_key": "E", "munch_key": "E", "var_key": "F", "var": 3,
    "url": "https://flourishelc.com.au/preschool-grantham-farm-nsw/",
    "primary": "preschool near grantham farm", "secondary": "",
    "meta_title": "Preschool Near Grantham Farm | Flourish Early Learning Centre",
    "meta_desc": "Looking for a preschool near Grantham Farm? Flourish ELC is purpose-built with EYLF-guided programs, five daily meals, and qualified educators. Book your tour.",
    "h1": "Preschool Near Grantham Farm: Where Every Child Gets the Best Start",
    "intro": "Flourish Early Learning Centre is a purpose-built preschool near Grantham Farm where the Early Years Learning Framework (EYLF): Belonging, Being and Becoming, guides every experience. Our qualified educators are passionate about early childhood and bring that passion to life in programs that are joyful, intentional, and genuinely child-centred. Families near Grantham Farm trust Flourish ELC because it is the kind of preschool that makes a real difference.",
    "why_heading": "Why Choose Flourish ELC for Preschool Near Grantham Farm",
    "faqs": [
        ("What age is Flourish ELC's preschool program designed for near Grantham Farm?", "Our preschool is designed for children aged 3 to 5, within a centre that welcomes all children from 6 weeks to school age with dedicated programs at every stage."),
        ("What food is served during the preschool day near Grantham Farm?", "Breakfast, morning tea, lunch, afternoon tea, and a late snack: five nutritious meals prepared fresh every weekday. Dietary needs and allergies are always carefully managed."),
        ("What are the operating hours for preschool near Grantham Farm?", "Monday to Friday, 7:00am to 6:00pm. Our parent communication app keeps families connected with real-time photos and updates throughout the day."),
        ("How does Flourish ELC approach STEM in the preschool program near Grantham Farm?", "STEM Learning is woven into daily preschool experiences through science explorations, mathematical thinking, building challenges, and nature investigations that make critical thinking genuinely exciting for this age group."),
        ("How does Flourish ELC support preschoolers with additional learning needs near Grantham Farm?", "We work closely with families, specialists, and relevant support services to create an inclusive preschool environment where every child's individual needs are understood, respected, and well supported."),
        ("What garden or environmental activities take place in the preschool near Grantham Farm?", "Children participate in garden-to-plate activities, composting, nature walks, and sustainability projects through both our Munch & Move program and our broader environmental education curriculum."),
        ("How do educators communicate the preschool curriculum to families near Grantham Farm?", "We share learning stories, program plans, and observations through our parent app and at regular educator-family conversations, so parents always understand what their preschooler is working on and why."),
        ("How does the Flourish ELC preschool program transition children to kindergarten near Grantham Farm?", "School readiness is embedded throughout the preschool year, with language, mathematics, independence, and social skills built progressively so every child arrives at kindergarten feeling confident and excited."),
    ],
  },
  {
    "num": 13, "type": "preschool", "state": "NSW", "suburb": "Riverstone",
    "prog_type": "preschool", "prog_key": "F", "munch_key": "F", "var_key": "G", "var": 0,
    "url": "https://flourishelc.com.au/preschool-riverstone-nsw/",
    "primary": "preschool near riverstone", "secondary": "",
    "meta_title": "Preschool Near Riverstone NSW | Flourish Early Learning Centre",
    "meta_desc": "Looking for a preschool near Riverstone? Flourish ELC is purpose-built with EYLF-guided programs, five daily meals, and passionate qualified educators. Book today.",
    "h1": "Preschool Near Riverstone That Families Are Proud to Choose",
    "intro": "Flourish Early Learning Centre is a purpose-built preschool near Riverstone guided by the Early Years Learning Framework (EYLF): Belonging, Being and Becoming. Our qualified educators bring both professional expertise and genuine warmth to preschool, creating an environment where children feel at home, feel inspired, and feel proud of who they are. Families near Riverstone choose Flourish ELC because they can see, and feel, the difference a truly great preschool makes.",
    "why_heading": "Why Choose Flourish ELC for Preschool Near Riverstone",
    "faqs": [
        ("When can my child start the preschool program near Riverstone?", "Our preschool program is available for children from 3 years of age. If your child is younger, our childcare programs from 6 weeks are also available at the same centre."),
        ("How many daily meals are included in the preschool program near Riverstone?", "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. All dietary requirements and allergies are accommodated with care and attention."),
        ("What hours does Flourish ELC's preschool operate near Riverstone?", "We are open Monday to Friday, 7:00am to 6:00pm, and our parent communication app ensures you are always connected to your preschooler's day."),
        ("How does Flourish ELC foster a love of reading and storytelling in preschool near Riverstone?", "Storytelling, shared reading, and language-rich conversations are part of every preschool day, with educators using books, puppets, and dramatic play to build vocabulary and a genuine love of stories."),
        ("How is the Munch & Move program delivered in preschool near Riverstone?", "Our preschool educators use Munch & Move resources to make healthy eating and physical activity engaging and natural parts of every day, through cooking activities, movement games, and mindful mealtimes."),
        ("What does a typical preschool day look like near Riverstone?", "Preschool days at Flourish ELC flow between guided and child-led activities, indoor and outdoor play, group experiences, mealtimes, and quiet rest periods, planned around children's interests and developmental needs."),
        ("Are there specialist programs or incursions for preschoolers near Riverstone?", "We regularly incorporate incursions including music, movement, science, and cultural programs that enrich our preschool curriculum and give children experiences that extend their learning in exciting ways."),
        ("How does Flourish ELC ensure the safety of preschoolers near Riverstone?", "Child safety is our highest priority. All educators hold current first aid and child protection training, our facilities meet NQF safety standards, and our policies are reviewed and updated regularly."),
    ],
  },
  {
    "num": 14, "type": "preschool", "state": "NSW", "suburb": "Mulgrave",
    "prog_type": "preschool", "prog_key": "G", "munch_key": "G", "var_key": "H", "var": 1,
    "url": "https://flourishelc.com.au/preschool-mulgrave-nsw/",
    "primary": "preschool near mulgrave", "secondary": "",
    "meta_title": "Preschool Near Mulgrave NSW | Flourish Early Learning Centre",
    "meta_desc": "Looking for a preschool near Mulgrave? Flourish ELC offers EYLF-guided programs, five daily meals, qualified educators, and a warm purpose-built environment. Enquire today.",
    "h1": "Preschool Near Mulgrave: Joyful Learning, Qualified Educators, Happy Children",
    "intro": "Flourish Early Learning Centre is a warm, purpose-built preschool near Mulgrave where children are welcomed into an environment shaped by the Early Years Learning Framework (EYLF): Belonging, Being and Becoming. Our qualified educators are passionate about the preschool years and create daily experiences that are rich, joyful, and genuinely meaningful. Families near Mulgrave trust Flourish ELC because children love being here, and the learning shows.",
    "why_heading": "Why Choose Flourish ELC for Preschool Near Mulgrave",
    "faqs": [
        ("Is Flourish ELC's preschool near Mulgrave suitable for children who have not been in care before?", "Absolutely. We warmly welcome children starting early education for the first time. Our educators are experienced in settling in new children and making the transition feel safe and joyful."),
        ("What are the daily meals like in the preschool program near Mulgrave?", "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack, all prepared fresh with children's health and growth in mind. Dietary requirements are always accommodated."),
        ("What are the preschool operating hours near Mulgrave?", "Monday to Friday, 7:00am to 6:00pm. Families can also access real-time updates and photos via our parent communication app throughout the preschool day."),
        ("How does Flourish ELC support preschoolers' cognitive development near Mulgrave?", "Cognitive development is nurtured through inquiry-based play, open-ended materials, problem-solving challenges, and educator questioning that encourages children to think deeply and discover answers for themselves."),
        ("What technology is used in the preschool program near Mulgrave?", "We take a balanced approach: technology is used purposefully as one of many learning tools, while screen time is kept minimal in favour of hands-on, creative, social, and outdoor experiences."),
        ("How does Flourish ELC support children's transition from preschool to kindergarten near Mulgrave?", "Our school readiness programs build language, numeracy, independence, fine motor, and social-emotional skills throughout the preschool year so children arrive at kindergarten genuinely prepared and excited."),
        ("How are family cultures and languages respected in the preschool program near Mulgrave?", "We celebrate every family's cultural background and home language. Families are invited to share their traditions and languages with the preschool group, enriching the cultural fabric of our whole community."),
        ("What is the process for enrolling in preschool near Mulgrave?", "Start by booking a tour of our centre. Our team will walk you through the enrolment process, answer all your questions, and help you understand what to expect from your child's first days in preschool."),
    ],
  },
  # ── WA CHILDCARE (group_offset=2, var_key C-A cycling) ──
  {
    "num": 15, "type": "childcare", "state": "WA", "suburb": "Bull Creek",
    "prog_type": "childcare", "prog_key": "A", "munch_key": "A", "var_key": "C", "var": 2,
    "url": "https://flourishelc.com.au/daycare-childcare-bull-creek-wa/",
    "primary": "childcare bull creek", "secondary": "daycare bull creek",
    "meta_title": "Childcare Bull Creek WA | Flourish Early Learning Centre",
    "meta_desc": "Warm, purpose-built childcare in Bull Creek with EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour today.",
    "h1": "Childcare in Bull Creek: Warm, Qualified, and Built Around Your Child",
    "intro": "Flourish Early Learning Centre offers families in Bull Creek a purpose-built early learning environment where children are guided by the best, and so are their educators. Our programs are shaped by the Early Years Learning Framework (EYLF): Belonging, Being and Becoming, and delivered by qualified, passionate educators who genuinely love what they do. Bull Creek families choose Flourish ELC because it is a place where children thrive, and parents feel completely at ease.",
    "why_heading": "Why Choose Flourish ELC for Childcare in Bull Creek",
    "faqs": [
        ("What ages does Flourish ELC in Bull Creek cater for in childcare?", "We welcome children from 6 weeks through to school age, with purpose-built rooms, outdoor spaces, and programs designed specifically for each developmental stage."),
        ("What meals are provided at Flourish ELC Bull Creek?", "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. All dietary requirements and allergies are always accommodated with care."),
        ("What childcare hours are available in Bull Creek?", "Our centre is open Monday to Friday, 7:00am to 6:00pm, giving families in Bull Creek a generous window for drop-off and pick-up."),
        ("How does the parent communication app work at the Bull Creek centre?", "The app sends real-time photos, updates, and observations from educators so families always feel genuinely connected to what their child is doing and learning throughout the day."),
        ("How does Flourish ELC approach cultural education in Bull Creek?", "Our Cultural Education program celebrates Indigenous Australian heritage, world cultures, and multicultural traditions through art, music, storytelling, and shared celebrations that bring diversity to life meaningfully."),
        ("What school readiness support is offered in childcare at Bull Creek?", "School readiness is embedded in our daily childcare program, with language, mathematics, independence, fine motor, and social-emotional skills developed progressively through intentional, play-based experiences."),
        ("What outdoor play is available at the Bull Creek childcare centre?", "Our purpose-built outdoor spaces give children daily opportunities to run, climb, garden, and explore nature, because outdoor play is one of the richest learning experiences for young children."),
        ("How do I enquire about enrolling my child at Bull Creek?", "Contact us via our website or phone to book a tour, discuss availability, and learn more about the enrolment process. We warmly welcome all families to come and see Flourish ELC firsthand."),
    ],
  },
  {
    "num": 16, "type": "childcare", "state": "WA", "suburb": "Leeming",
    "prog_type": "childcare", "prog_key": "B", "munch_key": "B", "var_key": "D", "var": 3,
    "url": "https://flourishelc.com.au/daycare-childcare-leeming-wa/",
    "primary": "childcare near leeming", "secondary": "daycare near leeming",
    "meta_title": "Childcare Near Leeming WA | Flourish Early Learning Centre",
    "meta_desc": "Looking for childcare near Leeming? Flourish ELC is purpose-built with EYLF-guided programs, five daily meals, qualified educators, and a parent communication app. Enquire today.",
    "h1": "Childcare Near Leeming: Where Children Belong, Grow, and Flourish",
    "intro": "Flourish Early Learning Centre is a purpose-built early learning centre near Leeming where the Early Years Learning Framework (EYLF): Belonging, Being and Becoming, is the foundation of every program, every interaction, and every space. Our qualified, passionate educators create an environment where children feel genuinely at home, genuinely seen, and genuinely free to learn at their own wonderful pace. Families near Leeming trust Flourish ELC because the warmth and quality are real.",
    "why_heading": "Why Choose Flourish ELC for Childcare Near Leeming",
    "faqs": [
        ("What age groups are catered for in childcare near Leeming at Flourish ELC?", "We welcome children from 6 weeks through to school age. Each age group has its own dedicated program, environment, and team of qualified educators."),
        ("Are meals included in childcare fees near Leeming?", "Yes, five nutritious meals are provided every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. All dietary requirements and allergies are carefully managed every day."),
        ("What are the operating hours at Flourish ELC near Leeming?", "Monday to Friday, 7:00am to 6:00pm. Our parent communication app also keeps families connected with real-time updates and photos throughout the childcare day."),
        ("How does Flourish ELC support children's social skills near Leeming?", "Social and emotional development is a core focus of our childcare program. Educators model positive social behaviour, facilitate peer relationships, and create a warm environment where every child feels they belong."),
        ("What is the National Quality Framework and how does it apply to the Leeming centre?", "The NQF is Australia's quality standard for early childhood education. Meeting and exceeding it means our Leeming childcare educators, programs, and environment are consistently held to the country's highest benchmarks."),
        ("How is physical activity incorporated into the childcare program near Leeming?", "Exercise and movement are part of every childcare day through outdoor play, movement games, dance, and structured activities that keep children physically active, healthy, and energised to learn."),
        ("How does Flourish ELC handle emergency situations at the Leeming centre?", "All educators hold current first aid and emergency response training. Our emergency procedures are regularly practised, and families are immediately notified in the event of any incident."),
        ("What music activities take place in childcare near Leeming?", "Music and Movement is woven through the childcare day, with children singing, playing rhythm instruments, dancing, and engaging with music in ways that build creativity, coordination, language, and genuine joy."),
    ],
  },
  {
    "num": 17, "type": "childcare", "state": "WA", "suburb": "Bateman",
    "prog_type": "childcare", "prog_key": "C", "munch_key": "C", "var_key": "E", "var": 0,
    "url": "https://flourishelc.com.au/daycare-childcare-bateman-wa/",
    "primary": "childcare near bateman", "secondary": "daycare near bateman",
    "meta_title": "Childcare Near Bateman WA | Flourish Early Learning Centre",
    "meta_desc": "Looking for childcare near Bateman? Flourish ELC is purpose-built with EYLF-guided learning, five daily meals, qualified educators, and a real-time parent app. Book a tour.",
    "h1": "Childcare Near Bateman: Qualified Educators, Warm Community, Happy Children",
    "intro": "Flourish Early Learning Centre is a purpose-built early learning centre near Bateman where families find exactly what they have been looking for: a warm, qualified, and genuinely child-centred environment. Guided by the Early Years Learning Framework (EYLF): Belonging, Being and Becoming, our programs and our people are dedicated to making each child's early years as rich and joyful as possible. Families near Bateman choose Flourish ELC because the quality is evident from the moment you walk through the door.",
    "why_heading": "Why Choose Flourish ELC for Childcare Near Bateman",
    "faqs": [
        ("Is Flourish ELC near Bateman an approved childcare provider?", "Yes, we are an approved childcare provider, which means eligible families can access the Child Care Subsidy (CCS) to help with the cost of care. Our team can help you navigate this process."),
        ("What meals and snacks are provided at the Bateman childcare centre?", "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. Our kitchen team accommodates all dietary requirements and allergies with care."),
        ("What hours does the Flourish ELC centre near Bateman operate?", "Our centre is open Monday to Friday, 7:00am to 6:00pm, and our parent communication app keeps families connected throughout the childcare day with real-time photos and updates."),
        ("How does Flourish ELC document children's learning near Bateman?", "Educators create detailed learning stories, observations, and photos that capture each child's milestones and growth, shared with families through our parent communication app on an ongoing basis."),
        ("How does Flourish ELC integrate sustainability into the childcare program near Bateman?", "Children learn to care for the environment through gardening, composting, recycling, nature play, and discussions about caring for Country woven meaningfully into the everyday childcare curriculum."),
        ("What creative projects do children do in childcare near Bateman?", "Creative Arts experiences include painting, drawing, sculpture, craft, collage, and collaborative art projects that develop fine motor skills, self-expression, and creative confidence children carry throughout life."),
        ("How are children's learning goals set and tracked near Bateman?", "Educators observe children carefully, identify their interests and developmental priorities, and use the EYLF to set and track meaningful learning goals that are shared and discussed with families."),
        ("How does Flourish ELC support language-rich play in childcare near Bateman?", "Language is embedded in everything: storytelling, conversations at mealtimes, songs, puppet play, books, and collaborative projects all build the vocabulary and communication skills children need for school."),
    ],
  },
  {
    "num": 18, "type": "childcare", "state": "WA", "suburb": "Rossmoyne",
    "prog_type": "childcare", "prog_key": "D", "munch_key": "D", "var_key": "F", "var": 1,
    "url": "https://flourishelc.com.au/daycare-childcare-rossmoyne-wa/",
    "primary": "childcare near rossmoyne", "secondary": "daycare near rossmoyne",
    "meta_title": "Childcare Near Rossmoyne WA | Flourish Early Learning Centre",
    "meta_desc": "Looking for childcare near Rossmoyne? Flourish ELC is purpose-built with EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book today.",
    "h1": "Childcare Near Rossmoyne: Where Learning Feels Like Home",
    "intro": "Flourish Early Learning Centre is a warm, purpose-built early learning centre near Rossmoyne where children are welcomed into an environment guided by the Early Years Learning Framework (EYLF): Belonging, Being and Becoming. Our qualified, passionate educators are the heart of Flourish ELC, and it shows. Families near Rossmoyne choose us because their children genuinely love coming, and the development they witness speaks for itself.",
    "why_heading": "Why Choose Flourish ELC for Childcare Near Rossmoyne",
    "faqs": [
        ("Does Flourish ELC near Rossmoyne accept very young babies in childcare?", "Yes, we welcome babies from 6 weeks of age. Our infant rooms are purpose-built with everything needed to provide the safe, nurturing, responsive care that very young babies require."),
        ("What types of food are served at the Rossmoyne childcare centre?", "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. Our menu is designed around children's nutritional needs, and all dietary requirements are managed carefully."),
        ("What time does Flourish ELC near Rossmoyne open and close?", "We are open Monday to Friday, 7:00am to 6:00pm. Families can access real-time updates via our parent communication app throughout the childcare day."),
        ("How are children's transitions between rooms managed near Rossmoyne?", "Room transitions are handled gradually and with care. Educators from both rooms work together to familiarise children with their new environment, new peers, and new educators before any formal transition."),
        ("What is the settling-in process at the Rossmoyne childcare centre?", "We offer flexible settling-in visits prior to a child's start date, encouraging families to visit multiple times so the centre feels familiar, safe, and welcoming before the first full childcare day."),
        ("How does Flourish ELC approach diversity and inclusion near Rossmoyne?", "Diversity is celebrated through our Cultural Education program, which incorporates Indigenous Australian perspectives and world cultures. Every child and every family's cultural identity is genuinely honoured at our centre."),
        ("Are there any sibling enrolment benefits at Flourish ELC near Rossmoyne?", "We give preference to siblings of currently enrolled children wherever possible, so families can have their children in the same centre. Please enquire with our team about current sibling enrolment options."),
        ("How does Flourish ELC communicate about public holiday closures near Rossmoyne?", "Families are notified well in advance of all public holiday centre closures through the parent app and via email, ensuring families always have adequate notice to arrange alternative care."),
    ],
  },
  {
    "num": 19, "type": "childcare", "state": "WA", "suburb": "Shelley",
    "prog_type": "childcare", "prog_key": "E", "munch_key": "E", "var_key": "G", "var": 2,
    "url": "https://flourishelc.com.au/daycare-childcare-shelley-wa/",
    "primary": "childcare near shelley", "secondary": "daycare near shelley",
    "meta_title": "Childcare Near Shelley WA | Flourish Early Learning Centre",
    "meta_desc": "Looking for childcare near Shelley? Flourish ELC is purpose-built with EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour.",
    "h1": "Childcare Near Shelley: A Place Where Children Thrive",
    "intro": "Flourish Early Learning Centre is a purpose-built early learning centre near Shelley where families find something genuinely special: a warm community, qualified educators, and programs shaped by the Early Years Learning Framework (EYLF): Belonging, Being and Becoming. At Flourish ELC, children are known, celebrated, and supported in becoming the best version of themselves. Families near Shelley trust us because their children love it here, and the learning is real.",
    "why_heading": "Why Choose Flourish ELC for Childcare Near Shelley",
    "faqs": [
        ("What is the minimum enrolment age at Flourish ELC near Shelley?", "We welcome children from 6 weeks of age. Our purpose-built infant rooms and programs are designed specifically for the needs of very young babies and their families."),
        ("How is the childcare menu planned near Shelley?", "Our kitchen team plans a rotating nutritious menu covering all five daily meals: breakfast, morning tea, lunch, afternoon tea, and a late snack. Dietary requirements, allergies, and cultural preferences are always accommodated."),
        ("What childcare hours are available at Flourish ELC near Shelley?", "Monday to Friday, 7:00am to 6:00pm. Our parent communication app ensures families are always connected to their child's childcare day with real-time updates and photos."),
        ("How are children's special needs and support requirements handled near Shelley?", "We work closely with families, paediatricians, and relevant specialist support services to create inclusive childcare environments where every child's individual needs are understood, respected, and well supported."),
        ("How does Flourish ELC celebrate important milestones in childcare near Shelley?", "Educators document and celebrate each child's milestones, capturing first steps, new words, friendships formed, and skills mastered in learning stories and photos shared joyfully with families via our app."),
        ("What weather and outdoor play policies apply at the Shelley childcare centre?", "Outdoor play is encouraged every day within appropriate weather conditions. On extreme heat or cold days, outdoor time is modified or moved indoors, always prioritising children's safety and comfort."),
        ("How does Flourish ELC support healthy sleep and rest routines in childcare near Shelley?", "Rest periods are built into the daily routine for younger children. Educators follow each child's individual sleep patterns and family preferences while ensuring a quiet, comfortable space to rest."),
        ("What is the best way to communicate concerns or feedback as a parent near Shelley?", "We welcome open, ongoing dialogue. Families can speak with educators at drop-off or pick-up, send messages through the parent communication app, or arrange a private meeting with our centre director."),
    ],
  },
  {
    "num": 20, "type": "childcare", "state": "WA", "suburb": "Willetton",
    "prog_type": "childcare", "prog_key": "F", "munch_key": "F", "var_key": "H", "var": 3,
    "url": "https://flourishelc.com.au/daycare-childcare-willetton-wa/",
    "primary": "childcare near willetton", "secondary": "daycare near willetton",
    "meta_title": "Childcare Near Willetton WA | Flourish Early Learning Centre",
    "meta_desc": "Looking for childcare near Willetton? Flourish ELC is purpose-built with EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book today.",
    "h1": "Childcare Near Willetton: Warmth, Quality, and Children Who Flourish",
    "intro": "Flourish Early Learning Centre is a purpose-built childcare centre near Willetton where everything is designed around children having the very best start in life. Guided by the Early Years Learning Framework (EYLF): Belonging, Being and Becoming, our programs are thoughtful, joyful, and genuinely focused on each child. Our qualified, passionate educators are what families near Willetton talk about most, because great educators make everything else possible.",
    "why_heading": "Why Choose Flourish ELC for Childcare Near Willetton",
    "faqs": [
        ("What age groups does Flourish ELC accommodate in childcare near Willetton?", "All children from 6 weeks through to school age are welcome, with dedicated rooms, programs, and outdoor spaces designed for infants, toddlers, and preschool-aged children."),
        ("What is included in the daily childcare fees near Willetton?", "Five nutritious daily meals are included: breakfast, morning tea, lunch, afternoon tea, and a late snack. All meals are prepared fresh and all dietary requirements are accommodated at no extra cost."),
        ("What are the childcare hours at Flourish ELC near Willetton?", "We are open Monday to Friday, 7:00am to 6:00pm. Our parent communication app keeps families connected with real-time photos and updates throughout the day."),
        ("How does Flourish ELC support mathematical thinking in childcare near Willetton?", "Mathematics is embedded in everyday play: counting, sorting, measuring, building, and problem-solving make mathematical thinking natural, enjoyable, and relevant for children at every age in our childcare program."),
        ("How are children's friendships and peer relationships supported near Willetton?", "Educators facilitate positive peer interactions, support conflict resolution with kindness, and create group experiences that help children develop social skills and friendship-building confidence they need for school."),
        ("What professional development do Flourish ELC educators undertake near Willetton?", "Our educators engage in ongoing professional development aligned with the NQF and EYLF, including reflective practice, child protection training, first aid updates, and specialised early childhood education programs."),
        ("What is a typical morning like in childcare near Willetton?", "Mornings begin with a warm welcome, breakfast together, and then a mix of child-led free play and educator-planned experiences, with outdoor play and morning tea woven through a thoughtful daily rhythm."),
        ("How does Flourish ELC handle the transition to preschool within the centre near Willetton?", "Room transitions are managed gradually and gently, with educators from both rooms working together so children feel familiar with their new environment and peers well before any formal transition occurs."),
    ],
  },
  {
    "num": 21, "type": "childcare", "state": "WA", "suburb": "Brentwood",
    "prog_type": "childcare", "prog_key": "G", "munch_key": "G", "var_key": "A", "var": 0,
    "url": "https://flourishelc.com.au/daycare-childcare-brentwood-wa/",
    "primary": "childcare near brentwood", "secondary": "daycare near brentwood",
    "meta_title": "Childcare Near Brentwood WA | Flourish Early Learning Centre",
    "meta_desc": "Looking for childcare near Brentwood? Flourish ELC is purpose-built with EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book today.",
    "h1": "Childcare Near Brentwood: Where Children Are Known, Loved, and Inspired",
    "intro": "Flourish Early Learning Centre is a purpose-built early learning centre near Brentwood where families find something they often say is hard to describe but impossible to miss: the feeling that their child truly belongs here. Our programs are guided by the Early Years Learning Framework (EYLF): Belonging, Being and Becoming, and our qualified, passionate educators bring that philosophy to life with warmth and expertise every single day.",
    "why_heading": "Why Choose Flourish ELC for Childcare Near Brentwood",
    "faqs": [
        ("What programs does Flourish ELC offer for different age groups in childcare near Brentwood?", "We have purpose-built programs for infants from 6 weeks, toddlers, and preschool-aged children through to school age, with environments and resources tailored to each developmental stage."),
        ("Are allergen-free meal options available at Flourish ELC near Brentwood?", "Yes, our kitchen team is experienced in managing all dietary requirements and allergies. All allergen information is documented at enrolment and carefully managed across all five daily meals every weekday."),
        ("What are the opening hours at the Flourish ELC centre near Brentwood?", "Monday to Friday, 7:00am to 6:00pm. Families can also use our parent communication app to stay connected with their child's childcare day in real time."),
        ("How does Flourish ELC approach inquiry-based learning in childcare near Brentwood?", "Educators observe children's curiosities and interests, then design open-ended learning experiences that invite questions, exploration, and discovery, supporting children to think critically and love learning."),
        ("What gardening activities are available in childcare near Brentwood?", "Garden-to-plate experiences are part of our Munch & Move program and everyday curriculum, with children planting, tending, harvesting, and cooking with produce from our centre garden throughout the year."),
        ("How are public holidays and centre closure days communicated near Brentwood?", "All public holiday closures and planned centre closure days are communicated well in advance via our parent app and email notifications, giving families plenty of time to arrange alternative care."),
        ("What does the Child Care Subsidy process look like at Flourish ELC near Brentwood?", "As an approved childcare provider, we can help eligible families understand and apply for the Child Care Subsidy through MyGov. Our team is happy to guide you through the process at enrolment."),
        ("How does Flourish ELC embed social justice concepts in childcare near Brentwood?", "Social Justice is introduced through age-appropriate stories, discussions, and activities that help children understand fairness, kindness, and the importance of standing up for others, building the values of compassionate citizens."),
    ],
  },
  # ── WA KINDERGARTEN (group_offset=3, var_key D-B cycling) ──
  {
    "num": 22, "type": "kinder", "state": "WA", "suburb": "Bull Creek",
    "prog_type": "kinder", "prog_key": "A", "munch_key": "A", "var_key": "D", "var": 1,
    "url": "https://flourishelc.com.au/kindergarten-bull-creek-wa/",
    "primary": "kindergarten bull creek", "secondary": "",
    "meta_title": "Kindergarten Bull Creek WA | Flourish Early Learning Centre",
    "meta_desc": "Looking for kindergarten in Bull Creek? Flourish ELC offers EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour today.",
    "h1": "Kindergarten in Bull Creek: A Joyful, Qualified Start to School Life",
    "intro": "Flourish Early Learning Centre offers a warm, EYLF-guided kindergarten program in Bull Creek where children are prepared for school in the best possible way: through joyful, intentional, play-based learning. Our Early Years Learning Framework, Belonging, Being and Becoming, is not just a concept here: it is the lived experience of every child in our care. Our qualified educators know each child well and create experiences that light them up.",
    "why_heading": "Why Choose Flourish ELC for Kindergarten in Bull Creek",
    "faqs": [
        ("What age is the kindergarten program at Flourish ELC in Bull Creek?", "Our kindergarten program is designed for children aged 4 to 5, in the year before they begin primary school. Our full centre also caters for children from 6 weeks across all programs."),
        ("What meals are included in the kindergarten program at Bull Creek?", "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. All dietary requirements and allergies are managed carefully every day."),
        ("What are the kindergarten hours at Flourish ELC in Bull Creek?", "Our centre is open Monday to Friday, 7:00am to 6:00pm, giving kindergarten families in Bull Creek a full and flexible day of quality care and learning."),
        ("How does the parent communication app work during the kindergarten year at Bull Creek?", "The app sends real-time photos, milestone updates, and educator observations so families are always connected to what their kindergartener is doing and learning throughout the day."),
        ("How does Flourish ELC prepare kindergarteners for primary school in Bull Creek?", "Our school readiness programs develop literacy, numeracy, independence, social-emotional skills, and fine motor abilities systematically through the kindergarten year so children begin primary school genuinely ready and excited."),
        ("What is the EYLF and how does it guide kindergarten at Bull Creek?", "The Early Years Learning Framework (Belonging, Being and Becoming) shapes how our educators plan kindergarten experiences and observe children's growth, placing each child's identity, wellbeing, and engagement at the centre."),
        ("What outdoor learning takes place in kindergarten at Bull Creek?", "Our Bull Creek outdoor spaces are used daily for nature exploration, physical play, gardening, and investigations that build children's curiosity, physical development, and environmental awareness."),
        ("How do I enrol my child in kindergarten at Flourish ELC Bull Creek?", "Book a tour through our website or phone to meet our team, see the centre, and discuss enrolment. Places in our kindergarten program are in demand, so we encourage families to enquire early."),
    ],
  },
  {
    "num": 23, "type": "kinder", "state": "WA", "suburb": "Leeming",
    "prog_type": "kinder", "prog_key": "B", "munch_key": "B", "var_key": "E", "var": 2,
    "url": "https://flourishelc.com.au/kindergarten-leeming-wa/",
    "primary": "kindergarten near leeming", "secondary": "",
    "meta_title": "Kindergarten Near Leeming WA | Flourish Early Learning Centre",
    "meta_desc": "Looking for kindergarten near Leeming? Flourish ELC offers EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour today.",
    "h1": "Kindergarten Near Leeming: Where Children Grow Ready, Willing, and Able",
    "intro": "Flourish Early Learning Centre offers a warm, play-based kindergarten near Leeming guided by the Early Years Learning Framework (EYLF): Belonging, Being and Becoming. Our kindergarten year is one of the most important a child will have, and we take that responsibility joyfully and seriously. Our qualified educators design kindergarten experiences that build confident, curious, capable children who walk into primary school genuinely ready for everything that comes next.",
    "why_heading": "Why Choose Flourish ELC for Kindergarten Near Leeming",
    "faqs": [
        ("When does kindergarten at Flourish ELC near Leeming typically begin?", "Kindergarten is designed for children in the year before primary school begins, typically aged 4 to 5. Contact us to discuss your child's specific start date and enrolment options."),
        ("Are meals part of the kindergarten day near Leeming?", "Yes, five nutritious meals are provided every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. Dietary needs and allergies are always managed with care."),
        ("What hours does the kindergarten program run near Leeming?", "Monday to Friday, 7:00am to 6:00pm. Families can also stay connected via our parent communication app throughout the kindergarten day."),
        ("How are kindergarteners supported in building literacy skills near Leeming?", "Literacy is developed through daily storytelling, reading, phonological awareness games, mark-making, and writing activities that make early literacy skills feel exciting and achievable for kindergarten-aged children."),
        ("What STEM activities are available to kindergarteners near Leeming?", "Kindergarteners engage in STEM Learning through building and engineering challenges, simple science experiments, mathematical investigations, and technology exploration that develop critical thinking in genuinely fun ways."),
        ("How does Flourish ELC's kindergarten near Leeming differ from a public kindergarten?", "Flourish ELC offers a full-day, 7:00am to 6:00pm kindergarten program with five daily meals, a parent communication app, purpose-built environments, and an EYLF-guided curriculum."),
        ("How is the Munch & Move program delivered in kindergarten near Leeming?", "Our kindergarten educators integrate Munch & Move through cooking activities, movement games, active outdoor play, and mindful mealtimes that build healthy eating and physical activity habits that last a lifetime."),
        ("How does Flourish ELC celebrate the transition from kindergarten to primary school near Leeming?", "The transition to primary school is celebrated as a significant milestone. We hold special events, share learning portfolios with families, and ensure every graduate leaves our centre feeling proud, confident, and excited."),
    ],
  },
  {
    "num": 24, "type": "kinder", "state": "WA", "suburb": "Bateman",
    "prog_type": "kinder", "prog_key": "C", "munch_key": "C", "var_key": "F", "var": 3,
    "url": "https://flourishelc.com.au/kindergarten-bateman-wa/",
    "primary": "kindergarten near bateman", "secondary": "",
    "meta_title": "Kindergarten Near Bateman WA | Flourish Early Learning Centre",
    "meta_desc": "Looking for kindergarten near Bateman? Flourish ELC offers EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour today.",
    "h1": "Kindergarten Near Bateman: Play-Based, Purposeful, and Genuinely Wonderful",
    "intro": "Flourish Early Learning Centre offers a warm, EYLF-guided kindergarten program near Bateman where the year before primary school is spent in the very best way: learning through play, connecting with others, and growing into a confident, capable child. The Early Years Learning Framework, Belonging, Being and Becoming, shapes every experience in our kindergarten, and our qualified, passionate educators bring it to life with skill and joy.",
    "why_heading": "Why Choose Flourish ELC for Kindergarten Near Bateman",
    "faqs": [
        ("What is the age requirement for kindergarten at Flourish ELC near Bateman?", "Kindergarten is designed for children aged 4 to 5, typically in the year before primary school begins. Our broader centre also welcomes children from 6 weeks across all age groups."),
        ("What food is provided during the kindergarten day near Bateman?", "Breakfast, morning tea, lunch, afternoon tea, and a late snack: five nutritious meals prepared fresh every weekday. All dietary requirements and allergies are always accommodated."),
        ("How long is the kindergarten day near Bateman?", "Our centre is open Monday to Friday, 7:00am to 6:00pm, providing a full and flexible kindergarten day that supports working families in the Bateman area."),
        ("How does Flourish ELC build mathematical skills in kindergarten near Bateman?", "Mathematical thinking is developed through counting, pattern recognition, measurement, sorting, and problem-solving embedded in play, cooking, building, and outdoor investigations, making maths feel relevant and exciting."),
        ("How does Flourish ELC support kindergarteners who are anxious about starting primary school near Bateman?", "Our school readiness programs build confidence and capability progressively so the transition feels manageable and exciting. Educators also work with families to address specific anxieties and prepare children emotionally."),
        ("What social-emotional skills does kindergarten at Flourish ELC develop near Bateman?", "Kindergarteners develop empathy, friendship skills, emotional regulation, resilience, and a growing ability to manage themselves in group settings, all foundational to thriving in primary school."),
        ("Are there any kindergarten orientation programs near Bateman before enrolment?", "We encourage families to book a tour before enrolling, and we offer settling-in visits once enrolled so children can familiarise themselves with the kindergarten environment before their official start."),
        ("How does Flourish ELC incorporate Cultural Education into the kindergarten program near Bateman?", "Indigenous Australian perspectives and world cultures are woven into the kindergarten curriculum through storytelling, art, music, and traditional celebrations that build cultural awareness, respect, and a global mindset."),
    ],
  },
  {
    "num": 25, "type": "kinder", "state": "WA", "suburb": "Rossmoyne",
    "prog_type": "kinder", "prog_key": "D", "munch_key": "D", "var_key": "G", "var": 0,
    "url": "https://flourishelc.com.au/kindergarten-rossmoyne-wa/",
    "primary": "kindergarten near rossmoyne", "secondary": "",
    "meta_title": "Kindergarten Near Rossmoyne WA | Flourish Early Learning Centre",
    "meta_desc": "Looking for kindergarten near Rossmoyne? Flourish ELC offers EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour today.",
    "h1": "Kindergarten Near Rossmoyne: A Joyful, Confident Start to School Life",
    "intro": "Flourish Early Learning Centre offers a warm, purpose-built kindergarten near Rossmoyne where children spend the year before school exactly as they should: playing, discovering, growing, and becoming. Our kindergarten programs are guided by the Early Years Learning Framework (EYLF): Belonging, Being and Becoming, and delivered by qualified educators who genuinely love the kindergarten years.",
    "why_heading": "Why Choose Flourish ELC for Kindergarten Near Rossmoyne",
    "faqs": [
        ("Is the kindergarten program at Flourish ELC near Rossmoyne play-based?", "Yes, our kindergarten program is deeply play-based because children aged 4 to 5 learn most effectively through joyful, purposeful, hands-on play experiences guided by skilled educators."),
        ("How many meals are served during kindergarten near Rossmoyne?", "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. Dietary requirements and allergies are always accommodated with care and consistency."),
        ("What are the hours for kindergarten near Rossmoyne?", "Monday to Friday, 7:00am to 6:00pm. Our parent communication app keeps kindergarten families connected with real-time photos and educator updates throughout the day."),
        ("What qualifications do kindergarten educators hold near Rossmoyne?", "Our educators hold recognised early childhood qualifications and meet and exceed the National Quality Framework (NQF) requirements, bringing professional expertise and genuine passion for the kindergarten years."),
        ("How does Flourish ELC develop children's sense of identity in kindergarten near Rossmoyne?", "Identity is celebrated and nurtured through activities that help children explore who they are, where they come from, and what they value, building the self-awareness and confidence they carry into school and life."),
        ("What science exploration activities take place in kindergarten near Rossmoyne?", "Kindergarteners investigate the natural world through experiments, nature collections, garden observations, and STEM challenges that build scientific curiosity, vocabulary, and the habit of asking questions."),
        ("How does Flourish ELC manage the handover of learning portfolios to primary school near Rossmoyne?", "We compile learning portfolios that capture each kindergartener's growth, achievements, and individual learning journey throughout the year, shared with families at the end of the kindergarten year as a treasured keepsake."),
        ("How is music integrated into the kindergarten program near Rossmoyne?", "Music and Movement is a daily feature of kindergarten, with children singing, playing percussion, creating rhythms, and moving their bodies in ways that build creativity, language, coordination, and pure joy."),
    ],
  },
  {
    "num": 26, "type": "kinder", "state": "WA", "suburb": "Shelley",
    "prog_type": "kinder", "prog_key": "E", "munch_key": "E", "var_key": "H", "var": 1,
    "url": "https://flourishelc.com.au/kindergarten-shelley-wa/",
    "primary": "kindergarten near shelley", "secondary": "",
    "meta_title": "Kindergarten Near Shelley WA | Flourish Early Learning Centre",
    "meta_desc": "Looking for kindergarten near Shelley? Flourish ELC offers EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour today.",
    "h1": "Kindergarten Near Shelley: Where the Best Possible Year Before School Happens",
    "intro": "Flourish Early Learning Centre offers a warm, play-based kindergarten near Shelley guided by the Early Years Learning Framework (EYLF): Belonging, Being and Becoming. The year before primary school is one of the most important a child will have, and our qualified, passionate educators pour everything into making it wonderful. Children in our kindergarten near Shelley grow academically, socially, and emotionally in an environment that is safe, joyful, and genuinely child-centred.",
    "why_heading": "Why Choose Flourish ELC for Kindergarten Near Shelley",
    "faqs": [
        ("What makes Flourish ELC's kindergarten near Shelley different from other options?", "Our EYLF-guided curriculum, qualified educators, purpose-built environment, five daily meals, and embedded school readiness programs combine to offer a warm, comprehensive, and genuinely excellent kindergarten experience."),
        ("Are snacks and meals included in the kindergarten fees near Shelley?", "Yes, all five daily meals are included: breakfast, morning tea, lunch, afternoon tea, and a late snack. All dietary requirements, allergies, and cultural food preferences are always managed with care."),
        ("What time does kindergarten start and finish near Shelley?", "Our centre operates Monday to Friday, 7:00am to 6:00pm, providing a full kindergarten day that accommodates families with a wide range of working schedules."),
        ("How does Flourish ELC approach creative arts in kindergarten near Shelley?", "Creative Arts are integral to our kindergarten curriculum: children paint, sculpt, draw, create music, engage in dramatic play, and collaborate on artistic projects that build expression, fine motor skills, and creative confidence."),
        ("How is independence built in the kindergarten year near Shelley?", "Independence is a core focus in our kindergarten program. Children are supported to manage their belongings, make choices, problem-solve independently, and contribute to the care of their environment, building real self-sufficiency."),
        ("How does Flourish ELC partner with families during the kindergarten year near Shelley?", "Families are genuine partners throughout the kindergarten year. We share learning stories, hold regular conversations, and invite parents into centre events so families always feel informed and celebrated."),
        ("What is the daily outdoor routine like for kindergarteners near Shelley?", "Kindergarteners access outdoor spaces daily for free play, nature investigations, gardening, and structured physical activities. Outdoor learning is valued as highly as indoor learning and planned with equal intention."),
        ("How does Flourish ELC ensure kindergarteners are truly school-ready by the end of the year near Shelley?", "Our School Readiness Program systematically builds language, mathematics, independence, fine motor, social, and emotional skills throughout the kindergarten year so every child leaves genuinely prepared and excited for primary school."),
    ],
  },
  {
    "num": 27, "type": "kinder", "state": "WA", "suburb": "Willetton",
    "prog_type": "kinder", "prog_key": "F", "munch_key": "F", "var_key": "A", "var": 2,
    "url": "https://flourishelc.com.au/kindergarten-willetton-wa/",
    "primary": "kindergarten near willetton", "secondary": "",
    "meta_title": "Kindergarten Near Willetton WA | Flourish Early Learning Centre",
    "meta_desc": "Looking for kindergarten near Willetton? Flourish ELC offers EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour today.",
    "h1": "Kindergarten Near Willetton: Play-Based, Warm, and Built for the Best Start",
    "intro": "Flourish Early Learning Centre offers a wonderful kindergarten near Willetton where children spend the year before primary school growing in every direction. Our kindergarten programs are guided by the Early Years Learning Framework (EYLF): Belonging, Being and Becoming, and delivered by qualified, passionate educators who make every day count. Families near Willetton choose our kindergarten because they see the quality in the detail: the care, the curriculum, the community, and the children who thrive here.",
    "why_heading": "Why Choose Flourish ELC for Kindergarten Near Willetton",
    "faqs": [
        ("How does Flourish ELC's kindergarten near Willetton approach school readiness?", "School readiness is woven into every day of our kindergarten program through play-based experiences that build language, numeracy, independence, and social-emotional skills progressively throughout the year."),
        ("What food is served during the kindergarten day near Willetton?", "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack, all prepared fresh. Dietary requirements and allergies are always carefully accommodated at no extra cost."),
        ("What are the kindergarten operating hours near Willetton?", "Monday to Friday, 7:00am to 6:00pm. Our parent communication app keeps families connected with real-time updates, photos, and educator observations throughout the kindergarten day."),
        ("What does the kindergarten curriculum look like in terms of daily structure near Willetton?", "Kindergarten days flow between child-led play, educator-guided experiences, group activities, outdoor time, mealtimes, and quieter moments, planned around a daily rhythm that gives children security, variety, and genuine engagement."),
        ("How does Flourish ELC incorporate Environmental Education in kindergarten near Willetton?", "Sustainability and care for Country are embedded in the kindergarten curriculum through gardening, composting, nature exploration, and discussions about the environment that build genuine environmental responsibility."),
        ("What language development strategies are used in kindergarten near Willetton?", "Storytelling, daily reading, dramatic play, vocabulary-rich conversations, phonological awareness games, and writing explorations are all used to build the strong language foundations that underpin success across every subject in primary school."),
        ("How does Flourish ELC support kindergarteners who are learning English as an additional language near Willetton?", "Our educators create language-rich environments that support all children, and we work closely with families to understand each child's language background and provide additional support and scaffolding where needed."),
        ("Is there a specific kindergarten orientation or orientation visit at the Willetton centre?", "Yes, we encourage all new kindergarten families to visit before their child begins. We offer orientation sessions so children can meet educators, explore the space, and feel comfortable before their first full kindergarten day."),
    ],
  },
  {
    "num": 28, "type": "kinder", "state": "WA", "suburb": "Brentwood",
    "prog_type": "kinder", "prog_key": "G", "munch_key": "G", "var_key": "B", "var": 3,
    "url": "https://flourishelc.com.au/kindergarten-brentwood-wa/",
    "primary": "kindergarten near brentwood", "secondary": "",
    "meta_title": "Kindergarten Near Brentwood WA | Flourish Early Learning Centre",
    "meta_desc": "Looking for kindergarten near Brentwood? Flourish ELC offers EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour today.",
    "h1": "Kindergarten Near Brentwood: Your Child's Best Year Before School Starts Here",
    "intro": "Flourish Early Learning Centre offers a warm, purpose-built kindergarten near Brentwood where the year before primary school is spent exactly as it should be: playing, learning, growing, and belonging. Guided by the Early Years Learning Framework (EYLF): Belonging, Being and Becoming, our kindergarten programs are shaped around each child's unique journey. Our qualified, passionate educators bring both skill and genuine heart to every kindergarten session.",
    "why_heading": "Why Choose Flourish ELC for Kindergarten Near Brentwood",
    "faqs": [
        ("What can my child expect from their first week of kindergarten at Flourish ELC near Brentwood?", "The first week focuses on settling in, building relationships with educators and peers, and exploring the kindergarten environment. Our educators work closely with each new child to ensure they feel safe and genuinely excited to be here."),
        ("What meals are served in the kindergarten program near Brentwood?", "Breakfast, morning tea, lunch, afternoon tea, and a late snack: five nutritious meals every weekday, prepared fresh. Dietary requirements and allergies are always managed with care and consistency."),
        ("What are the hours for kindergarten at Flourish ELC near Brentwood?", "Monday to Friday, 7:00am to 6:00pm. Our parent communication app delivers real-time photos and updates throughout the kindergarten day so families always feel genuinely close to their child's world."),
        ("What role do kindergarten educators play in preparing children for primary school near Brentwood?", "Our educators are active partners in every child's school readiness journey, designing intentional experiences that build language, mathematics, independence, social skills, and fine motor abilities progressively throughout the kindergarten year."),
        ("How does Flourish ELC's kindergarten near Brentwood incorporate the Munch & Move program?", "Munch & Move is integrated through cooking activities, movement sessions, mindful mealtimes, and active play experiences that build healthy eating habits and a love of physical activity from the earliest years."),
        ("How are kindergarten educators at Brentwood selected and trained?", "All educators hold recognised early childhood qualifications, meet and exceed the NQF, and engage in ongoing professional development. We select for both academic qualification and genuine passion for working with children."),
        ("How does the Flourish ELC kindergarten near Brentwood support children's curiosity?", "Curiosity is celebrated and encouraged through open-ended materials, inquiry-based projects, educator questioning, and an environment that invites children to wonder, investigate, and discover at every turn of every day."),
        ("How does Flourish ELC communicate with kindergarten families about their child's progress near Brentwood?", "Learning stories, milestone observations, and regular educator conversations ensure families are always informed and celebrated as partners. Our parent app makes real-time sharing easy and joyful throughout the kindergarten year."),
    ],
  },
]



# ── SERVICE TYPE MAP ──
service_type_map = {
    'childcare': 'childcare',
    'preschool': 'preschool',
    'kinder': 'kindergarten',
}

# ── PROG DICT MAP ──
prog_dict_map = {
    'childcare': prog_childcare,
    'preschool': prog_preschool,
    'kinder': prog_kinder,
}


def add_page(p):
    suburb = p['suburb']
    state = p['state']
    service_type = service_type_map[p['type']]
    var_key = p['var_key']
    prog_key = p['prog_key']
    munch_key = p['munch_key']

    doc.add_heading(f"Page {p['num']}", level=1)

    # Info block
    info = f"URL: {p['url']}\nPage Type: Location Page\nPrimary Keyword: {p['primary']}"
    if p.get('secondary'):
        info += f"\nSecondary Keywords: {p['secondary']}"
    doc.add_paragraph(info)

    doc.add_paragraph("Content:")
    doc.add_paragraph(f"Meta Title: {p['meta_title']}")
    doc.add_paragraph(f"Meta Description: {p['meta_desc']}")

    # H1
    doc.add_heading(f"H1: {p['h1']}", level=2)
    doc.add_paragraph(p['intro'])

    # Resolve content variants
    prog_sec = prog_dict_map[p['prog_type']][prog_key]
    munch_text = (munch_nsw if state == 'NSW' else munch_wa)[munch_key].replace('{service_type}', service_type)
    cultural_text = cultural[var_key].replace('{suburb}', suburb).replace('{service_type}', service_type)
    sr_text = school_readiness[var_key].replace('{suburb}', suburb).replace('{service_type}', service_type)
    env_sec = environment[var_key]
    why_set = why_bullets[var_key]

    def render_prog():
        doc.add_heading(f"H2: {prog_sec['heading']}", level=3)
        doc.add_paragraph(prog_sec['body'].replace('{state}', state))
        doc.add_heading("H3: Munch & Move", level=4)
        doc.add_paragraph(munch_text)
        doc.add_heading("H3: Cultural Education", level=4)
        doc.add_paragraph(cultural_text)
        doc.add_heading("H3: School Readiness Program", level=4)
        doc.add_paragraph(sr_text)

    def render_why():
        doc.add_heading(f"H2: {p['why_heading']}", level=3)
        for label, desc in why_set:
            bullet = doc.add_paragraph(style='List Bullet')
            run = bullet.add_run(f"{label}: ")
            run.bold = True
            bullet.add_run(desc.replace('{suburb}', suburb))

    def render_env():
        doc.add_heading(f"H2: {env_sec['heading']}", level=3)
        doc.add_paragraph(env_sec['body'].replace('{suburb}', suburb))

    order_map = {
        0: [render_prog, render_env, render_why],
        1: [render_env, render_why, render_prog],
        2: [render_why, render_prog, render_env],
        3: [render_prog, render_why, render_env],
    }

    for fn in order_map[p['var']]:
        fn()

    # FAQs
    doc.add_heading("H2: Frequently Asked Questions", level=3)
    for q, a in p['faqs']:
        para = doc.add_paragraph()
        run = para.add_run(q)
        run.bold = True
        doc.add_paragraph(a)

    doc.add_page_break()


# Generate pages 1–28
for p in pages:
    add_page(p)
    print(f"Page {p['num']} done")

# ── PAGE 29 ── Areas We Serve hub
doc.add_heading("Page 29", level=1)
doc.add_paragraph(
    "URL: https://flourishelc.com.au/areas-we-serve/\n"
    "Page Type: Hub Page\n"
    "Primary Keyword: childcare near me\n"
    "Secondary Keywords: daycare near me, preschool near me, kindergarten near me"
)
doc.add_paragraph("Content:")
doc.add_paragraph("Meta Title: Find Flourish Early Learning Centres Near You | Flourish ELC")
doc.add_paragraph(
    "Meta Description: Looking for childcare near me? Flourish ELC runs warm, purpose-built centres "
    "in NSW and WA with EYLF-guided programs, five fresh daily meals, qualified educators, and a "
    "real-time parent app. Find your nearest centre and book a tour today."
)

doc.add_heading(
    "H1: Childcare Near Me: Find a Flourish Early Learning Centre That Feels Like Home",
    level=2
)
doc.add_paragraph(
    "Searching for childcare near me is one of the most important searches a family will ever make. "
    "Flourish Early Learning Centre is a family of warm, purpose-built early learning centres "
    "across New South Wales and Western Australia, each guided by the Early Years Learning Framework "
    "(EYLF): Belonging, Being and Becoming. Whether you are looking for daycare, preschool, or "
    "kindergarten, every Flourish ELC centre is designed to make children feel safe, seen, and "
    "genuinely excited to learn. Families who find us tend to stay."
)

doc.add_heading("H2: What Every Flourish ELC Centre Offers Your Family", level=3)
doc.add_paragraph(
    "Every Flourish ELC location, whether you are searching for childcare, daycare, preschool, or "
    "kindergarten near you, is built around the same core experience. Programs are guided by the EYLF "
    "and the National Quality Framework (NQF), covering Creative Arts, STEM Learning, Language, "
    "Mathematics, Fine Motor Skills, Music and Movement, Cultural Understanding, Social Justice, "
    "Independence, Exercise, Health, and Social and Emotional development.\n\n"
    "Five nutritious meals are served fresh every weekday: breakfast, morning tea, lunch, afternoon tea, "
    "and a late snack. Dietary requirements and allergies are always carefully managed. Our centres are "
    "open Monday to Friday, 7:00am to 6:00pm. A real-time parent communication app keeps every family "
    "close to their child's world throughout the day."
)

doc.add_heading("H2: Why Choose Flourish ELC", level=3)
for label, desc in why_bullets['A']:
    bullet = doc.add_paragraph(style='List Bullet')
    run = bullet.add_run(f"{label}: ")
    run.bold = True
    bullet.add_run(desc.replace('{suburb}', 'your area'))

doc.add_heading("H2: School Readiness, Munch & Move, and Cultural Education — Everywhere", level=3)
doc.add_paragraph(
    "Every Flourish ELC centre integrates three signature program elements across childcare, preschool, "
    "and kindergarten. The Munch & Move program (a NSW Health initiative available across our network) "
    "promotes healthy eating, physical activity, and reduced screen time through hands-on cooking, "
    "movement games, and mindful mealtimes. Cultural Education celebrates Australia's rich diversity "
    "through storytelling, music, art, and traditional practices that build empathy and a global "
    "mindset in every child. Our School Readiness Program builds language, mathematics, independence, "
    "fine motor skills, and social-emotional capabilities so every child transitions to primary school "
    "with confidence and genuine enthusiasm."
)

doc.add_heading("H2: Frequently Asked Questions", level=3)
hub_faqs = [
    ("What age groups do Flourish ELC centres accept?",
     "All Flourish ELC centres welcome children from 6 weeks through to school age, with tailored programs and environments for each developmental stage."),
    ("What meals does Flourish ELC provide each day?",
     "Five fresh nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. All dietary requirements and allergies are managed with care."),
    ("What are the opening hours at Flourish ELC centres?",
     "All centres are open Monday to Friday, 7:00am to 6:00pm."),
    ("How does the parent communication app work at Flourish ELC?",
     "The app sends real-time updates, photos, and milestone observations throughout the day, so families stay genuinely close to their child's world even while at work."),
    ("What is the EYLF and how does it shape my child's experience?",
     "The Early Years Learning Framework: Belonging, Being and Becoming guides every decision our educators make — from environment design to curriculum planning — keeping each child's wellbeing and development at the centre."),
    ("Does Flourish ELC offer a school readiness program?",
     "Yes. School readiness is embedded across all age groups and includes a dedicated program focusing on language, mathematics, independence, fine motor skills, social-emotional skills, and confidence for the transition to primary school."),
    ("How do I find the Flourish ELC centre closest to me?",
     "Browse the areas listed on this page to find your nearest location, then book a tour. Our team will be delighted to show you the centre, answer your questions, and help you decide if Flourish ELC is the right fit for your family."),
    ("What makes Flourish ELC different from other childcare providers?",
     "Our combination of EYLF-guided curriculum, NQF-exceeding quality, five fresh daily meals, real-time parent app, dedicated school readiness and cultural education programs, and genuinely passionate educators creates an experience families consistently describe as exceptional."),
]
for q, a in hub_faqs:
    para = doc.add_paragraph()
    run = para.add_run(q)
    run.bold = True
    doc.add_paragraph(a)

print("Page 29 done")

output_path = "/home/invoi/fahad_projects/clients/flourishelc.com.au/flourishelc.com.au_v4.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
