from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

def add_page(doc, page_num, url, page_type, primary_kw, secondary_kw,
             meta_title, meta_desc, h1_text, h1_body,
             h2_1_title, h2_1_body,
             h2_2_title, h2_2_body,
             h2_3_title, faqs):
    """Add a single page block to the document."""
    # Page heading
    doc.add_heading(f"Page {page_num}", level=1)

    # Info paragraph
    info = f"URL: {url}\nPage Type: {page_type}\nPrimary Keyword: {primary_kw}"
    if secondary_kw:
        info += f"\nSecondary Keywords: {secondary_kw}"
    doc.add_paragraph(info)

    doc.add_paragraph("Content:")
    doc.add_paragraph(f"Meta Title: {meta_title}")
    doc.add_paragraph(f"Meta Description: {meta_desc}")

    # H1
    doc.add_heading("H1: " + h1_text, level=2)
    doc.add_paragraph(h1_body)

    # H2 section 1
    doc.add_heading(h2_1_title, level=3)
    doc.add_paragraph(h2_1_body)

    # H2 section 2
    doc.add_heading(h2_2_title, level=3)
    doc.add_paragraph(h2_2_body)

    # H2 section 3 (FAQs or benefits)
    doc.add_heading(h2_3_title, level=3)
    for q, a in faqs:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.bold = True
        doc.add_paragraph(a)

    doc.add_page_break()


# ── PAGE 1 ── Oakville childcare
add_page(
    doc, 1,
    "https://flourishelc.com.au/daycare-childcare-oakville-nsw/",
    "Location Page", "childcare oakville", "daycare oakville",
    "Childcare Oakville | Flourish Early Learning Centre",
    "Warm, purpose-built childcare in Oakville with qualified educators, five daily meals, and programs guided by the EYLF. Book your tour today.",
    "Childcare Oakville Families Love: Flourish Early Learning Centre",
    "Flourish Early Learning Centre is a purpose-built childcare in Oakville designed around children thriving in every sense of the word. Guided by the Early Years Learning Framework (EYLF) -- Belonging, Being and Becoming -- our qualified educators bring genuine passion to every single day. Oakville families choose us because we go beyond care: we create a place where children light up, grow, and flourish.",
    "What Makes Flourish ELC Special for Oakville Families",
    "Our programs span a rich range of learning areas that cover the whole child. Creative Arts, Music and Movement, STEM Learning, Language, Mathematics, and Fine Motor Skills are woven into daily experiences -- not as separate subjects, but as joyful, connected play. Cultural Understanding and Social Justice are part of our curriculum too, because we believe children flourish when they understand and celebrate the world around them.\n\nWe also serve five nutritious meals every day: breakfast, morning tea, lunch, afternoon tea, and a late snack. Our doors are open 7:00am to 6:00pm Monday to Friday, and our parent communication app keeps you connected with real-time updates throughout the day. School readiness programs mean your child transitions to kindergarten with confidence, curiosity, and a love of learning already in full bloom.",
    "Our Nurturing Environment",
    "The Flourish ELC environment is purpose-built -- every room, every outdoor space, and every learning corner has been thoughtfully designed for children aged 0 to 5. Our educators meet and exceed the National Quality Framework (NQF), bringing both professional excellence and genuine warmth to every interaction.\n\nExercise, Health, and Social and Emotional development are central to our daily rhythm. Children have space to move, rest, create, and connect. Independence is nurtured gently, with educators standing alongside children as they build confidence in their own abilities. The result is a community within a centre -- one that Oakville families feel proud to be part of.",
    "Frequently Asked Questions",
    [
        ("What age groups does Flourish ELC Oakville cater to?", "We welcome children from 6 weeks through to school age, with dedicated programs tailored to each developmental stage."),
        ("What meals does Flourish ELC provide?", "We serve five nutritious meals daily: breakfast, morning tea, lunch, afternoon tea, and a late snack. All dietary requirements and allergies are carefully catered for."),
        ("What are your opening hours?", "We are open Monday to Friday, 7:00am to 6:00pm, giving families a generous window for drop-off and pick-up."),
        ("How does the parent communication app work?", "Our app sends you real-time updates, photos, and milestones throughout the day so you stay connected to your child's world even while you're at work."),
        ("What is the Early Years Learning Framework?", "The EYLF -- Belonging, Being and Becoming -- is the national framework that guides how we plan learning experiences. It places children's identity, wellbeing, and engagement at the heart of everything we do."),
    ]
)

# ── PAGE 2 ── Box Hill childcare
add_page(
    doc, 2,
    "https://flourishelc.com.au/daycare-childcare-box-hill-nsw/",
    "Location Page", "childcare near box hill", "daycare near box hill",
    "Childcare Near Box Hill NSW | Flourish Early Learning Centre",
    "Searching for childcare near Box Hill? Flourish ELC is a purpose-built centre with qualified educators, five daily meals, and EYLF-guided programs. Book your tour today.",
    "Childcare Near Box Hill: Where Every Child Belongs, Grows and Flourishes",
    "Flourish Early Learning Centre is a warm, purpose-built early learning centre conveniently located for families in and around Box Hill. We are guided by the Early Years Learning Framework (EYLF) -- Belonging, Being and Becoming -- and our qualified, passionate educators bring that philosophy to life every single day. If you are looking for childcare near Box Hill that feels like a second home for your child, you have found it.",
    "Learning Programs That Cover the Whole Child",
    "At Flourish ELC, learning is never one-dimensional. Our curriculum weaves together Creative Arts, STEM Learning, Language, Mathematics, Music and Movement, Fine Motor Skills, Cultural Understanding, Social Justice, Exercise, Health, Independence, and Social and Emotional development. These aren't boxes to tick -- they are the threads of a rich, joyful childhood.\n\nChildren learn by doing, by playing, by wondering, and by connecting with each other and their educators. Our school readiness programs make sure that when the time comes for kindergarten, children step through those gates feeling ready, capable, and excited. Five nutritious meals are prepared daily -- breakfast, morning tea, lunch, afternoon tea, and a late snack -- so children are always fuelled for discovery.",
    "A Safe, Welcoming Place That Meets and Exceeds the NQF",
    "Our purpose-built facility is designed with children in mind from the ground up. Every space -- from bright indoor classrooms to outdoor play areas -- has been crafted to inspire curiosity and support development. We meet and exceed the National Quality Framework (NQF), which means families can trust that our educators, programs, and environments are held to the highest national standards.\n\nWe are open 7:00am to 6:00pm Monday to Friday, and our parent communication app delivers real-time updates, photos, and milestone moments straight to your phone. The Flourish ELC community near Box Hill is one where children thrive and families feel genuinely supported.",
    "Frequently Asked Questions",
    [
        ("What age groups do you welcome at Flourish ELC near Box Hill?", "We cater for children from 6 weeks to school age, with dedicated rooms and programs for each stage of development."),
        ("What meals are included?", "Five nutritious meals each day: breakfast, morning tea, lunch, afternoon tea, and a late snack. We cater for all dietary requirements and allergies."),
        ("What are your opening hours?", "7:00am to 6:00pm, Monday to Friday."),
        ("How do you keep families connected during the day?", "Our parent communication app sends real-time updates, including photos and notes about your child's day, so you never feel out of the loop."),
        ("How does the EYLF shape your programs?", "The Early Years Learning Framework guides how we plan, observe, and respond to each child's learning journey, placing belonging, wellbeing, and curiosity at the centre of everything."),
    ]
)

# ── PAGE 3 ── Vineyard childcare
add_page(
    doc, 3,
    "https://flourishelc.com.au/daycare-childcare-vineyard-nsw/",
    "Location Page", "childcare near vineyard", "daycare near vineyard",
    "Childcare Near Vineyard NSW | Flourish Early Learning Centre",
    "Looking for childcare near Vineyard? Flourish ELC is purpose-built, guided by the EYLF, and offers five daily meals and qualified educators. Book your tour today.",
    "Childcare Near Vineyard That Puts Children First",
    "Flourish Early Learning Centre is a purpose-built early learning centre serving families in and around Vineyard. Built around the Early Years Learning Framework (EYLF) -- Belonging, Being and Becoming -- our centre is a place where children are celebrated for who they are, supported in what they're becoming, and given every opportunity to belong. Our qualified, passionate educators make Flourish ELC a truly special place for Vineyard families.",
    "Rich Programs for Growing Minds",
    "Our curriculum at Flourish ELC covers the full spectrum of early childhood development. STEM Learning sparks curiosity about how the world works. Creative Arts and Music and Movement give children expressive outlets that build confidence. Language, Mathematics, and Fine Motor Skills lay the academic foundations that serve children well into their school years. Cultural Understanding, Social Justice, and Social and Emotional development nurture empathy, resilience, and a sense of self.\n\nExercise and Health are built into every day -- children move, play outdoors, and develop physical confidence alongside all the other wonderful things they are discovering. Five nutritious meals are served daily, from breakfast through to a late afternoon snack, so growing bodies are always well nourished.",
    "Purpose-Built and Open Every Weekday",
    "Our facility has been designed from the ground up with children and families in mind. Meeting and exceeding the National Quality Framework (NQF) is not just a checkbox for us -- it is a commitment that runs through everything we do. Bright spaces, thoughtful outdoor areas, and warm, capable educators create an environment families near Vineyard trust completely.\n\nWe are open 7:00am to 6:00pm Monday to Friday. Our parent communication app keeps you genuinely connected with your child's day through real-time updates and photos. School readiness programs mean children transition to kindergarten with the skills and the heart to embrace what comes next.",
    "Frequently Asked Questions",
    [
        ("What ages do you cater for at Flourish ELC near Vineyard?", "From 6 weeks through to school age, with age-appropriate programs and dedicated spaces for each developmental stage."),
        ("What food does Flourish ELC provide?", "Five nutritious meals every day: breakfast, morning tea, lunch, afternoon tea, and a late snack. Dietary requirements and allergies are carefully managed."),
        ("What are your hours?", "Monday to Friday, 7:00am to 6:00pm."),
        ("How will I know what my child has been doing during the day?", "Our parent communication app sends real-time updates, photos, and observations so you stay connected throughout the day."),
        ("What is school readiness at Flourish ELC?", "Our school readiness programs develop the academic, social, and emotional skills children need to transition confidently into kindergarten."),
    ]
)

# ── PAGE 4 ── Gables childcare
add_page(
    doc, 4,
    "https://flourishelc.com.au/daycare-childcare-gables-nsw/",
    "Location Page", "childcare near gables", "daycare near gables",
    "Childcare Near Gables NSW | Flourish Early Learning Centre",
    "Warm, purpose-built childcare near Gables with EYLF-guided programs, qualified educators, and five nutritious daily meals. Book your family tour today.",
    "Childcare Near Gables: A Place Where Children Truly Flourish",
    "Flourish Early Learning Centre is a purpose-built childcare centre welcoming families in and around Gables. Everything here -- from the way our educators engage with children to the way our spaces are designed -- is shaped by the Early Years Learning Framework (EYLF): Belonging, Being and Becoming. Children at Flourish ELC don't just attend a centre; they belong to a warm, vibrant community that celebrates every step of their journey.",
    "Programs That Nurture Every Dimension of Childhood",
    "Our curriculum covers the learning areas that matter most in a child's early years. STEM Learning, Creative Arts, Language, Mathematics, Fine Motor Skills, and Music and Movement build the foundations children carry into school and beyond. Cultural Understanding, Social Justice, Independence, Exercise, Health, and Social and Emotional development round out a program that sees children as whole people -- curious, capable, and full of potential.\n\nFive nutritious meals are prepared and served each day: breakfast, morning tea, lunch, afternoon tea, and a late snack. School readiness programs are threaded through our curriculum so that when children move on to kindergarten, they do so with confidence and genuine enthusiasm for learning.",
    "A Qualified Team in a Purpose-Built Home",
    "Our team of qualified, passionate educators meets and exceeds the National Quality Framework (NQF). Their dedication shows in every interaction -- in the way they notice a child's spark of curiosity, the way they support a child who needs a moment, and the way they celebrate every milestone. The Flourish ELC environment itself has been purpose-built: safe, welcoming, and designed to let children explore and grow.\n\nWe are open 7:00am to 6:00pm Monday to Friday. Our parent communication app sends real-time updates so families near Gables stay genuinely connected to their child's day, even on the busiest mornings.",
    "Frequently Asked Questions",
    [
        ("What ages does Flourish ELC near Gables cater for?", "We welcome children from 6 weeks through to school age, with dedicated rooms and programs tailored to each developmental stage."),
        ("What meals are served daily?", "Breakfast, morning tea, lunch, afternoon tea, and a late snack -- five nutritious meals every weekday."),
        ("What are your opening hours near Gables?", "Our centre is open Monday to Friday, 7:00am to 6:00pm."),
        ("How does the parent app work?", "The app delivers real-time updates, photos, and notes about your child's activities and meals throughout the day."),
        ("How does the NQF affect the quality of care?", "Meeting and exceeding the National Quality Framework means our programs, educators, and environments are assessed against Australia's highest national standards for early education and care."),
    ]
)

# ── PAGE 5 ── Grantham Farm childcare
add_page(
    doc, 5,
    "https://flourishelc.com.au/daycare-childcare-grantham farm-nsw/",
    "Location Page", "childcare near grantham farm", "daycare near grantham farm",
    "Childcare Near Grantham Farm | Flourish Early Learning Centre",
    "Purpose-built childcare near Grantham Farm with EYLF-guided learning, qualified educators, five daily meals, and a real-time parent app. Book your tour today.",
    "Childcare Near Grantham Farm: Warm, Qualified, and Built for Growing Families",
    "Families near Grantham Farm looking for childcare that goes beyond the basics will find exactly that at Flourish Early Learning Centre. Our purpose-built centre is guided by the Early Years Learning Framework (EYLF) -- Belonging, Being and Becoming -- and staffed by qualified educators who genuinely love what they do. At Flourish ELC, children are known, celebrated, and given every opportunity to thrive.",
    "Learning That Covers the Whole Child",
    "Our daily programs are anything but ordinary. STEM Learning, Creative Arts, Language, Mathematics, Music and Movement, and Fine Motor Skills are woven into play-based experiences that make learning feel like the adventure it should be. Cultural Understanding, Social Justice, Social and Emotional development, Independence, Exercise, and Health complete a curriculum designed to grow children in every sense of the word.\n\nFive nutritious meals are prepared and served each day -- breakfast, morning tea, lunch, afternoon tea, and a late snack. Our school readiness programs ensure that when children move on to kindergarten, they are ready academically, socially, and emotionally. Nothing is left to chance.",
    "Purpose-Built, NQF-Aligned, and Open Every Weekday",
    "Our facility near Grantham Farm has been designed from the ground up for children aged 0 to 5. Safe, bright, and welcoming, it meets and exceeds the National Quality Framework (NQF) in every area. Families can trust that the people caring for their children hold formal qualifications and are committed to continuous professional growth.\n\nWe are open 7:00am to 6:00pm Monday to Friday. The Flourish ELC parent communication app keeps families connected throughout the day with real-time updates, photos, and observations -- a small but meaningful way of saying your child's day matters to us as much as it does to you.",
    "Frequently Asked Questions",
    [
        ("What age groups do you cater for near Grantham Farm?", "We welcome children from 6 weeks to school age, with age-appropriate rooms and programs for each developmental stage."),
        ("What meals does Flourish ELC provide?", "Five nutritious meals daily: breakfast, morning tea, lunch, afternoon tea, and a late snack. All dietary needs and allergies are accommodated."),
        ("What are the centre's opening hours?", "Monday to Friday, 7:00am to 6:00pm."),
        ("How do parents stay informed throughout the day?", "Our parent communication app provides real-time updates, photos, and milestones so you always know what your child has been up to."),
        ("What does school readiness look like at Flourish ELC?", "Our school readiness programs develop the literacy, numeracy, social, and emotional skills children need to transition into kindergarten with confidence."),
    ]
)

# ── PAGE 6 ── Riverstone childcare
add_page(
    doc, 6,
    "https://flourishelc.com.au/daycare-childcare-riverstone-nsw/",
    "Location Page", "childcare near riverstone", "daycare near riverstone",
    "Childcare Near Riverstone NSW | Flourish Early Learning Centre",
    "Warm, purpose-built childcare near Riverstone with EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour.",
    "Childcare Near Riverstone: Belonging, Growing, and Flourishing Together",
    "Flourish Early Learning Centre offers families near Riverstone a warm, purpose-built early learning environment where children are known by name, celebrated for who they are, and gently guided toward all they are becoming. Our programs are shaped by the Early Years Learning Framework (EYLF) -- Belonging, Being and Becoming -- and delivered by qualified educators who bring both skill and heart to every day.",
    "Programs Full of Purpose and Joy",
    "At Flourish ELC, every day is packed with rich learning experiences. Creative Arts, STEM Learning, Language, Mathematics, Fine Motor Skills, Music and Movement, Cultural Understanding, Social Justice, Independence, Exercise, Health, and Social and Emotional development are all part of our curriculum. These areas are not taught in isolation -- they come to life through play, projects, conversations, and shared discovery.\n\nFive nutritious meals are served each day: breakfast, morning tea, lunch, afternoon tea, and a late snack. Our school readiness programs prepare children thoroughly for the transition to kindergarten, focusing on the skills and confidence they need to step into formal education with a smile.",
    "A Caring Environment Held to the Highest Standards",
    "Our purpose-built centre near Riverstone meets and exceeds the National Quality Framework (NQF). That means every aspect of the physical environment, our educator qualifications, our program quality, and our family partnerships has been thoughtfully considered and regularly assessed. Families choose Flourish ELC because they trust that standard of care.\n\nWe are open Monday to Friday, 7:00am to 6:00pm. Our parent communication app delivers real-time updates and photos throughout the day so families always feel connected, no matter how busy things get.",
    "Frequently Asked Questions",
    [
        ("What ages do you cater for near Riverstone?", "From 6 weeks through to school age, with dedicated programs and spaces for infants, toddlers, and preschoolers."),
        ("What meals are included?", "Five nutritious meals daily: breakfast, morning tea, lunch, afternoon tea, and a late snack. Dietary requirements and allergies are always catered for."),
        ("What are your opening hours?", "Monday to Friday, 7:00am to 6:00pm."),
        ("How does the parent app work?", "It sends real-time updates, photos, and observations straight to your phone, keeping you connected to your child's day from anywhere."),
        ("How does Flourish ELC approach school readiness?", "School readiness is woven through our programs from an early age, building literacy, numeracy, independence, and social skills so children transition to kindergarten feeling genuinely ready."),
    ]
)

# ── PAGE 7 ── Mulgrave childcare
add_page(
    doc, 7,
    "https://flourishelc.com.au/daycare-childcare-mulgrave-nsw/",
    "Location Page", "childcare near mulgrave", "daycare near mulgrave",
    "Childcare Near Mulgrave NSW | Flourish Early Learning Centre",
    "Looking for childcare near Mulgrave? Flourish ELC is purpose-built with EYLF-guided programs, five nutritious daily meals, and qualified educators. Book your tour.",
    "Childcare Near Mulgrave That Families Trust and Children Love",
    "Flourish Early Learning Centre is a purpose-built early learning centre welcoming families near Mulgrave into a warm, vibrant community. Guided by the Early Years Learning Framework (EYLF) -- Belonging, Being and Becoming -- our centre is a place where children feel secure, excited, and free to explore who they are. Our qualified, passionate educators are the heart of everything we do.",
    "A Curriculum That Celebrates the Whole Child",
    "Learning at Flourish ELC stretches across the full range of early childhood development. STEM Learning, Creative Arts, Language, Mathematics, Music and Movement, Fine Motor Skills, Cultural Understanding, Social Justice, Exercise, Health, Independence, and Social and Emotional development are all part of our daily rhythm. Children discover, create, move, connect, and grow -- and every experience is intentional.\n\nFive nutritious meals are prepared and served each weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. Our school readiness programs are thoughtfully integrated into the curriculum so that children are well-prepared for the next chapter of their education when the time comes.",
    "Purpose-Built, NQF-Certified, and Genuinely Welcoming",
    "Our facility near Mulgrave has been purpose-built for children from 6 weeks to school age. It meets and exceeds the National Quality Framework (NQF), which gives families confidence that everything -- the spaces, the programs, the educator qualifications, and the family partnerships -- is held to Australia's highest national standards for early education.\n\nWe are open 7:00am to 6:00pm Monday to Friday. Our parent communication app keeps families connected throughout the day with real-time updates, photos, and observations. It is one of the many ways Flourish ELC makes families feel truly part of their child's world, even when they are apart.",
    "Frequently Asked Questions",
    [
        ("What ages do you cater for near Mulgrave?", "We welcome children from 6 weeks through to school age, with dedicated rooms and programs tailored to each stage."),
        ("Are meals provided?", "Yes, five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. All dietary requirements are carefully accommodated."),
        ("What are your opening hours?", "Monday to Friday, 7:00am to 6:00pm."),
        ("How do I stay connected to my child's day?", "Our parent communication app sends real-time updates, photos, and notes so you are always in the loop."),
        ("What is the NQF?", "The National Quality Framework sets and measures quality standards for early education and care in Australia. We meet and exceed those standards across every area of our centre."),
    ]
)

# ── PAGE 8 ── Oakville preschool
add_page(
    doc, 8,
    "https://flourishelc.com.au/preschool-oakville-nsw/",
    "Location Page", "preschool oakville", "",
    "Preschool Oakville NSW | Flourish Early Learning Centre",
    "Looking for a preschool in Oakville? Flourish ELC offers EYLF-guided programs, qualified educators, five daily meals, and a warm, purpose-built environment. Enquire today.",
    "Preschool Oakville Families Choose: Flourish Early Learning Centre",
    "Flourish Early Learning Centre is a warm, purpose-built preschool in Oakville where children step into a world designed entirely for them. Our preschool programs are guided by the Early Years Learning Framework (EYLF) -- Belonging, Being and Becoming -- and delivered by qualified educators who genuinely care about every child's growth. Families in Oakville choose Flourish ELC because it is the kind of place that children look forward to, and parents feel proud of.",
    "Preschool Programs That Build Bright Futures",
    "Our preschool curriculum is designed to develop children across every dimension. Language, Mathematics, and Fine Motor Skills lay the academic groundwork that children carry into kindergarten. Creative Arts, Music and Movement, and STEM Learning spark the curiosity and creative thinking that make learning joyful. Cultural Understanding, Social Justice, Social and Emotional development, Independence, Exercise, and Health complete a program that sees each child as a whole person.\n\nChildren are served five nutritious meals every weekday -- breakfast, morning tea, lunch, afternoon tea, and a late snack. Our dedicated school readiness programs ensure preschoolers are socially, emotionally, and academically prepared to transition into kindergarten with genuine excitement and confidence.",
    "A Safe, Purposefully Designed Preschool Environment",
    "Our purpose-built centre in Oakville meets and exceeds the National Quality Framework (NQF), which means our educators, spaces, and programs are all assessed against Australia's highest standards for early education. Children spend their preschool years in an environment that has been thoughtfully designed for their age and stage -- safe, stimulating, and genuinely welcoming.\n\nWe are open Monday to Friday, 7:00am to 6:00pm. Our parent communication app keeps families connected throughout the day with real-time updates and photos, so even when you are far away, you feel close to what your child is experiencing.",
    "Frequently Asked Questions",
    [
        ("What age does the preschool program cater to at Flourish ELC Oakville?", "Our preschool programs are designed for children aged 3 to 5, though our centre welcomes children from 6 weeks to school age."),
        ("What meals are included in the preschool program?", "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. All allergies and dietary requirements are accommodated."),
        ("What are the preschool hours?", "Our centre is open Monday to Friday, 7:00am to 6:00pm."),
        ("How do I know my child is progressing well?", "Our educators share observations and milestones through the parent communication app, and formal school readiness assessments are available as children approach kindergarten age."),
        ("How does the EYLF guide the preschool program?", "The EYLF shapes how our educators plan experiences, observe children, and support learning -- always with belonging, wellbeing, and engagement at the heart of practice."),
    ]
)

# ── PAGE 9 ── Box Hill preschool
add_page(
    doc, 9,
    "https://flourishelc.com.au/preschool-box-hill-nsw/",
    "Location Page", "preschool near box hill", "",
    "Preschool Near Box Hill NSW | Flourish Early Learning Centre",
    "Searching for a preschool near Box Hill? Flourish ELC is purpose-built with qualified educators, EYLF-guided programs, and five daily meals. Book your tour today.",
    "Preschool Near Box Hill: Where Children Grow Into Themselves",
    "Flourish Early Learning Centre is a purpose-built preschool near Box Hill where the Early Years Learning Framework (EYLF) -- Belonging, Being and Becoming -- is not just a document on a wall but the living philosophy behind every experience. Our qualified, passionate educators know each child by name, by interest, and by strength. Families near Box Hill choose Flourish ELC because their children flourish here in every sense of the word.",
    "A Preschool Curriculum Full of Discovery",
    "Our preschool programs near Box Hill are rich and varied. Language, Mathematics, and Fine Motor Skills build the academic foundations children need. STEM Learning, Creative Arts, and Music and Movement make every day an adventure. Cultural Understanding, Social Justice, Social and Emotional development, Independence, Exercise, and Health round out a program where the whole child -- not just the academic child -- is celebrated and supported.\n\nFive nutritious meals are prepared and served each weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. School readiness programs are thoughtfully woven through our preschool curriculum so that children are genuinely ready -- socially, emotionally, and academically -- when kindergarten arrives.",
    "Purpose-Built for Preschoolers, Held to the Highest Standards",
    "Our purpose-built centre meets and exceeds the National Quality Framework (NQF), which means families near Box Hill can trust that every aspect of our preschool -- the educators, the spaces, the programs, and the partnerships with families -- has been carefully assessed against Australia's national quality standards. We take that responsibility seriously and work hard to exceed it.\n\nWe are open Monday to Friday, 7:00am to 6:00pm. Our parent communication app delivers real-time updates, photos, and observations throughout the preschool day, so families always feel connected to their child's world.",
    "Frequently Asked Questions",
    [
        ("What ages does the preschool near Box Hill cater for?", "Our preschool programs are designed for children aged 3 to 5, and our full centre caters for children from 6 weeks to school age."),
        ("What meals are included?", "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. Dietary needs and allergies are always accommodated."),
        ("What are the opening hours?", "Monday to Friday, 7:00am to 6:00pm."),
        ("How does Flourish ELC prepare children for school?", "School readiness programs develop literacy, numeracy, independence, and social skills progressively so children transition to kindergarten with confidence and curiosity."),
        ("How do educators stay connected with families?", "Through daily updates, photos, and observations shared via our parent communication app, as well as open conversations at drop-off and pick-up."),
    ]
)

# ── PAGE 10 ── Vineyard preschool
add_page(
    doc, 10,
    "https://flourishelc.com.au/preschool-vineyard-nsw/",
    "Location Page", "preschool near vineyard", "",
    "Preschool Near Vineyard NSW | Flourish Early Learning Centre",
    "Looking for a preschool near Vineyard? Flourish ELC is purpose-built with EYLF-guided programs, qualified educators, and five daily nutritious meals. Enquire today.",
    "Preschool Near Vineyard: A Warm Place Where Learning Comes Alive",
    "Flourish Early Learning Centre is a purpose-built preschool near Vineyard where the Early Years Learning Framework (EYLF) -- Belonging, Being and Becoming -- shapes every experience, every relationship, and every room. Our qualified educators bring knowledge and genuine warmth to the preschool years, creating an environment where children discover their strengths, build their confidence, and develop a love of learning that will stay with them for life. Families near Vineyard trust Flourish ELC because they can see the difference it makes.",
    "Rich Preschool Programs Across Every Learning Area",
    "Learning at Flourish ELC is broad, deep, and joyful. Language and Mathematics build strong academic foundations. STEM Learning and Creative Arts cultivate the problem-solving and creative thinking children will use their whole lives. Fine Motor Skills, Music and Movement, and Exercise support physical development in ways that are fun and confidence-building. Cultural Understanding, Social Justice, Independence, Health, and Social and Emotional development round out a curriculum where the whole child is cherished.\n\nFive nutritious meals are served each weekday -- breakfast, morning tea, lunch, afternoon tea, and a late snack. School readiness is embedded in our preschool programs, so by the time children move on to kindergarten, they carry both the skills and the spark to embrace it.",
    "A Purposefully Designed Space, Open Every Weekday",
    "Our purpose-built centre near Vineyard meets and exceeds the National Quality Framework (NQF). Families can be confident that the people caring for their children hold recognised qualifications and are supported by a centre environment that has been designed from the ground up for children's safety, wellbeing, and delight.\n\nWe are open Monday to Friday, 7:00am to 6:00pm. Our parent communication app sends real-time photos and updates throughout the preschool day, keeping families near Vineyard genuinely close to their child's world even when apart.",
    "Frequently Asked Questions",
    [
        ("What age does the preschool near Vineyard cater for?", "Preschool programs are designed for children aged 3 to 5, and we welcome children from 6 weeks to school age across our full centre."),
        ("What meals does Flourish ELC provide?", "Five nutritious meals each weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. All dietary requirements and allergies are carefully managed."),
        ("What are the hours at the Vineyard preschool?", "Monday to Friday, 7:00am to 6:00pm."),
        ("How are families kept informed?", "Our parent communication app delivers real-time updates and photos throughout the day so you always know what your child has been exploring and learning."),
        ("What does school readiness look like?", "We build literacy, numeracy, independence, and social and emotional skills progressively through the preschool years, so children transition to kindergarten with confidence and readiness."),
    ]
)

print("Pages 1-10 done")

# ── PAGE 11 ── Gables preschool
add_page(
    doc, 11,
    "https://flourishelc.com.au/preschool-gables-nsw/",
    "Location Page", "preschool near gables", "",
    "Preschool Near Gables NSW | Flourish Early Learning Centre",
    "Looking for a preschool near Gables? Flourish ELC is purpose-built with EYLF-guided learning, qualified educators, and five daily meals. Book your tour today.",
    "Preschool Near Gables: Where Little Learners Truly Belong",
    "Flourish Early Learning Centre is a warm, purpose-built preschool near Gables where the Early Years Learning Framework (EYLF) -- Belonging, Being and Becoming -- is more than a framework: it is the spirit of everything we do. Our qualified, passionate educators create an environment where preschoolers feel safe, celebrated, and inspired to explore the world around them. Families near Gables love Flourish ELC because their children thrive here, and they can see it every single day.",
    "Preschool Programs That Develop the Whole Child",
    "Our preschool curriculum is carefully designed to nurture children across every dimension of development. Language, Mathematics, and Fine Motor Skills build the academic skills children need for kindergarten. Creative Arts, Music and Movement, and STEM Learning bring the joy and wonder that keep children genuinely engaged. Cultural Understanding, Social Justice, Social and Emotional development, Independence, Exercise, and Health complete a program that looks at the whole child -- not just what they know, but who they are becoming.\n\nFive nutritious meals are prepared fresh each weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. Our school readiness programs are woven through the preschool years so children are thoroughly prepared -- academically and socially -- for their next adventure.",
    "Purpose-Built and Held to National Quality Standards",
    "Our purpose-built centre near Gables meets and exceeds the National Quality Framework (NQF). Every element of the centre -- from the learning spaces and outdoor environments to the qualifications our educators hold and the partnerships we build with families -- is assessed against Australia's highest national standards. We take pride in not just meeting those standards, but genuinely surpassing them.\n\nWe are open Monday to Friday, 7:00am to 6:00pm. Our parent communication app shares real-time updates, photos, and observations throughout the day, keeping families near Gables close to their child's world even during a busy workday.",
    "Frequently Asked Questions",
    [
        ("What age groups does the preschool near Gables serve?", "Our preschool programs cater for children aged 3 to 5. Our full centre welcomes children from 6 weeks to school age."),
        ("What meals are included each day?", "Breakfast, morning tea, lunch, afternoon tea, and a late snack -- five nutritious meals prepared fresh every weekday."),
        ("What are the preschool opening hours?", "7:00am to 6:00pm, Monday to Friday."),
        ("How do families stay connected during the day?", "Our parent communication app delivers real-time updates and photos so you always know how your child's day is unfolding."),
        ("How does Flourish ELC prepare preschoolers for kindergarten?", "School readiness programs build literacy, numeracy, social and emotional skills, and independence progressively, so children step into kindergarten feeling capable and excited."),
    ]
)

# ── PAGE 12 ── Grantham Farm preschool
add_page(
    doc, 12,
    "https://flourishelc.com.au/preschool-grantham-farm-nsw/",
    "Location Page", "preschool near grantham farm", "",
    "Preschool Near Grantham Farm | Flourish Early Learning Centre",
    "Looking for a preschool near Grantham Farm? Flourish ELC is purpose-built with EYLF-guided programs, five daily meals, and qualified educators. Book your tour.",
    "Preschool Near Grantham Farm: Where Every Child Gets the Best Start",
    "Flourish Early Learning Centre is a purpose-built preschool near Grantham Farm where the Early Years Learning Framework (EYLF) -- Belonging, Being and Becoming -- guides every experience. Our qualified educators are passionate about early childhood and bring that passion to life in programs that are joyful, intentional, and genuinely child-centred. Families near Grantham Farm trust Flourish ELC because it is the kind of preschool that makes a real difference.",
    "A Preschool Curriculum Designed for Curious Minds",
    "At Flourish ELC, preschool means far more than structured sitting and worksheets. Language, Mathematics, and Fine Motor Skills are developed through play-based experiences that children find genuinely exciting. STEM Learning and Creative Arts cultivate the thinking and creative skills that serve children well beyond the preschool years. Music and Movement, Cultural Understanding, Social Justice, Independence, Exercise, Health, and Social and Emotional development complete a curriculum that honours childhood while building bright futures.\n\nFive nutritious meals are served each weekday -- breakfast, morning tea, lunch, afternoon tea, and a late snack. School readiness programs are embedded throughout the preschool years, ensuring children have the academic foundations, the social skills, and the emotional confidence to love kindergarten from day one.",
    "Safe, Welcoming, and Held to the Highest Standards",
    "Our purpose-built preschool near Grantham Farm meets and exceeds the National Quality Framework (NQF). Families can feel genuinely confident about the quality of care their children receive, knowing that our educators hold recognised qualifications and that our environment has been designed and assessed against Australia's national standards for early education.\n\nWe are open Monday to Friday, 7:00am to 6:00pm. Our parent communication app keeps families connected throughout the preschool day with real-time updates, photos, and milestones -- because what happens in our preschool matters, and we know you want to be part of it.",
    "Frequently Asked Questions",
    [
        ("What ages does the preschool near Grantham Farm cater for?", "Preschool programs are designed for children aged 3 to 5. Our full centre welcomes children from 6 weeks to school age."),
        ("What meals are provided?", "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. All dietary needs and allergies are accommodated."),
        ("What are the opening hours?", "Monday to Friday, 7:00am to 6:00pm."),
        ("How will I know what my child has been learning?", "Our parent communication app shares real-time updates, photos, and observations throughout the day."),
        ("How does Flourish ELC approach school readiness?", "Literacy, numeracy, independence, and social and emotional skills are built progressively through our preschool programs so children transition to kindergarten with confidence and genuine readiness."),
    ]
)

# ── PAGE 13 ── Riverstone preschool
add_page(
    doc, 13,
    "https://flourishelc.com.au/preschool-riverstone-nsw/",
    "Location Page", "preschool near riverstone", "",
    "Preschool Near Riverstone NSW | Flourish Early Learning Centre",
    "Looking for a preschool near Riverstone? Flourish ELC is purpose-built with EYLF-guided programs, five daily meals, and passionate qualified educators. Book today.",
    "Preschool Near Riverstone That Families Are Proud to Choose",
    "Flourish Early Learning Centre is a purpose-built preschool near Riverstone guided by the Early Years Learning Framework (EYLF) -- Belonging, Being and Becoming. Our qualified educators bring both professional expertise and genuine warmth to preschool, creating an environment where children feel at home, feel inspired, and feel proud of who they are. Families near Riverstone choose Flourish ELC because they can see -- and feel -- the difference a truly great preschool makes.",
    "Preschool Programs That Spark and Sustain a Love of Learning",
    "Learning at Flourish ELC is holistic and joyful. STEM Learning, Creative Arts, Language, Mathematics, Fine Motor Skills, Music and Movement, Cultural Understanding, Social Justice, Independence, Exercise, Health, and Social and Emotional development all feature in our preschool programs. Together, these learning areas build children who are curious, confident, and ready to take on whatever comes next.\n\nFive nutritious meals are prepared fresh every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. School readiness programs are part of our preschool curriculum from the outset, so children near Riverstone develop the academic, social, and emotional skills they need to love kindergarten from their very first day.",
    "A Purpose-Built Home for Preschool Learners",
    "Our facility near Riverstone has been purpose-built with preschoolers in mind and meets and exceeds the National Quality Framework (NQF). Families can be confident that everything -- the spaces, the people, the programs, and the partnerships -- is aligned with Australia's highest standards for early education. That confidence is something Flourish ELC families describe as one of the most important things we offer.\n\nWe are open Monday to Friday, 7:00am to 6:00pm. Our parent communication app keeps families genuinely connected throughout the preschool day with real-time photos, updates, and observations.",
    "Frequently Asked Questions",
    [
        ("What age does the Riverstone preschool cater for?", "Preschool programs serve children aged 3 to 5, and our full centre welcomes children from 6 weeks to school age."),
        ("Are meals included in the preschool program?", "Yes, five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. Dietary requirements and allergies are always managed carefully."),
        ("What are your preschool hours near Riverstone?", "Monday to Friday, 7:00am to 6:00pm."),
        ("How do parents receive updates about their child's day?", "Our parent communication app delivers real-time photos and observations so you always know what your preschooler is up to."),
        ("What is the EYLF and why does it matter?", "The Early Years Learning Framework guides early educators across Australia in planning and delivering high-quality learning experiences. At Flourish ELC, it underpins everything -- from the way we design our spaces to the way our educators engage with children."),
    ]
)

# ── PAGE 14 ── Mulgrave preschool
add_page(
    doc, 14,
    "https://flourishelc.com.au/preschool-mulgrave-nsw/",
    "Location Page", "preschool near mulgrave", "",
    "Preschool Near Mulgrave NSW | Flourish Early Learning Centre",
    "Looking for a preschool near Mulgrave? Flourish ELC offers EYLF-guided programs, five daily meals, qualified educators, and a warm purpose-built environment. Enquire today.",
    "Preschool Near Mulgrave: Joyful Learning, Qualified Educators, Happy Children",
    "Flourish Early Learning Centre is a warm, purpose-built preschool near Mulgrave where children are welcomed into an environment shaped by the Early Years Learning Framework (EYLF) -- Belonging, Being and Becoming. Our qualified educators are passionate about the preschool years and create daily experiences that are rich, joyful, and genuinely meaningful. Families near Mulgrave trust Flourish ELC because children love being here, and the learning shows.",
    "Preschool Programs Covering Every Dimension of Development",
    "Our preschool curriculum reaches across the full breadth of early childhood development. Language, Mathematics, and Fine Motor Skills lay the academic groundwork. STEM Learning and Creative Arts cultivate the thinking and creative skills children carry into their school years and beyond. Music and Movement, Cultural Understanding, Social Justice, Independence, Exercise, Health, and Social and Emotional development ensure that every child is seen and supported as a whole person.\n\nFive nutritious meals are served each weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. Our school readiness programs build the literacy, numeracy, independence, and social skills children need to walk into kindergarten feeling ready, capable, and excited.",
    "Purpose-Built, NQF-Aligned, and Always Welcoming",
    "Our purpose-built preschool near Mulgrave meets and exceeds the National Quality Framework (NQF), which means families can trust the quality of the educators, the spaces, and the programs their children experience every day. That standard is one Flourish ELC takes seriously -- not just as a compliance exercise, but as a reflection of our values.\n\nWe are open Monday to Friday, 7:00am to 6:00pm. Our parent communication app keeps families connected with real-time updates and photos throughout the preschool day, because we know how much those little windows into your child's world matter.",
    "Frequently Asked Questions",
    [
        ("What ages does the preschool near Mulgrave serve?", "Preschool programs are designed for children aged 3 to 5. Our full centre welcomes children from 6 weeks to school age."),
        ("What meals are included?", "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. Dietary requirements and allergies are always carefully managed."),
        ("What are the preschool opening hours?", "Monday to Friday, 7:00am to 6:00pm."),
        ("How will I know what my child is doing during the day?", "Our parent communication app delivers real-time photos and updates so you always feel close to your child's world."),
        ("How does Flourish ELC get children ready for kindergarten?", "School readiness programs develop literacy, numeracy, social and emotional skills, and independence progressively throughout the preschool years."),
    ]
)

# ── PAGE 15 ── Bull Creek childcare
add_page(
    doc, 15,
    "https://flourishelc.com.au/daycare-childcare-bull-creek-wa/",
    "Location Page", "childcare bull creek", "daycare bull creek",
    "Childcare Bull Creek WA | Flourish Early Learning Centre",
    "Warm, purpose-built childcare in Bull Creek with EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour today.",
    "Childcare in Bull Creek: Warm, Qualified, and Built Around Your Child",
    "Flourish Early Learning Centre offers families in Bull Creek a purpose-built early learning environment where children are guided by the best -- and so are their educators. Our programs are shaped by the Early Years Learning Framework (EYLF) -- Belonging, Being and Becoming -- and delivered by qualified, passionate educators who genuinely love what they do. Bull Creek families choose Flourish ELC because it is a place where children thrive, and parents feel completely at ease.",
    "Learning Programs That Cover Every Part of Childhood",
    "Our curriculum in Bull Creek brings together all the learning areas that make the early years so powerful. STEM Learning, Creative Arts, Language, Mathematics, Fine Motor Skills, Music and Movement, Cultural Understanding, Social Justice, Independence, Exercise, Health, and Social and Emotional development are all woven into daily play-based experiences. Children discover, create, connect, and grow every single day.\n\nFive nutritious meals are prepared and served each weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. Our school readiness programs ensure that when children move on to kindergarten and formal schooling, they carry the skills, the confidence, and the love of learning they need to thrive.",
    "Purpose-Built and Held to the Highest National Standards",
    "Our facility in Bull Creek is purpose-built for children aged 6 weeks to school age and meets and exceeds the National Quality Framework (NQF). Families can feel genuinely confident that the quality of care, the educators' qualifications, and the programs their children experience are assessed against Australia's highest national standards for early education.\n\nWe are open Monday to Friday, 7:00am to 6:00pm. Our parent communication app keeps families connected throughout the day with real-time updates, photos, and observations -- a genuine window into your child's world whenever you need it.",
    "Frequently Asked Questions",
    [
        ("What age groups do you cater for in Bull Creek?", "From 6 weeks to school age, with dedicated programs and spaces for infants, toddlers, and preschoolers."),
        ("What meals are included at Flourish ELC?", "Five nutritious meals daily: breakfast, morning tea, lunch, afternoon tea, and a late snack. All dietary requirements and allergies are accommodated."),
        ("What are the centre's opening hours?", "Monday to Friday, 7:00am to 6:00pm."),
        ("How do families stay connected during the day?", "Our parent communication app delivers real-time updates and photos straight to your phone throughout the day."),
        ("How does Flourish ELC prepare children for school?", "Our school readiness programs build literacy, numeracy, independence, and social and emotional skills so children move into kindergarten ready and excited."),
    ]
)

# ── PAGE 16 ── Leeming childcare
add_page(
    doc, 16,
    "https://flourishelc.com.au/daycare-childcare-leeming-wa/",
    "Location Page", "childcare near leeming", "daycare near leeming",
    "Childcare Near Leeming WA | Flourish Early Learning Centre",
    "Looking for childcare near Leeming? Flourish ELC is purpose-built with EYLF-guided programs, five daily meals, qualified educators, and a parent communication app. Enquire today.",
    "Childcare Near Leeming: Where Children Belong, Grow, and Flourish",
    "Flourish Early Learning Centre is a purpose-built early learning centre near Leeming where the Early Years Learning Framework (EYLF) -- Belonging, Being and Becoming -- is the foundation of every program, every interaction, and every space. Our qualified, passionate educators create an environment where children feel genuinely at home, genuinely seen, and genuinely free to learn at their own wonderful pace. Families near Leeming trust Flourish ELC because the warmth and quality are real.",
    "Programs That Nurture the Whole Child",
    "Learning at Flourish ELC covers all the areas that make the early years so formative and so important. STEM Learning, Creative Arts, Language, Mathematics, Fine Motor Skills, Music and Movement, Cultural Understanding, Social Justice, Independence, Exercise, Health, and Social and Emotional development are all part of our daily curriculum. Children are not just kept busy; they are genuinely growing, discovering, and connecting.\n\nFive nutritious meals are served fresh each weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. Our school readiness programs are thoughtfully integrated so that children near Leeming are academically, socially, and emotionally prepared when the time comes to start kindergarten.",
    "A Safe, Purpose-Built Environment That Meets the NQF",
    "Our purpose-built centre near Leeming meets and exceeds the National Quality Framework (NQF). That means every area of the centre -- from the learning environments and outdoor spaces to the educator qualifications and family partnerships -- has been assessed against Australia's national standards for early education. We take pride in the quality that standard represents.\n\nWe are open Monday to Friday, 7:00am to 6:00pm. Our parent communication app keeps families near Leeming genuinely connected throughout the day with real-time updates, photos, and observations.",
    "Frequently Asked Questions",
    [
        ("What ages do you cater for near Leeming?", "From 6 weeks to school age, with dedicated programs and spaces for each developmental stage."),
        ("What meals are provided daily?", "Five nutritious meals: breakfast, morning tea, lunch, afternoon tea, and a late snack. All dietary requirements and allergies are carefully managed."),
        ("What are the opening hours?", "Monday to Friday, 7:00am to 6:00pm."),
        ("How does the parent communication app work?", "It sends real-time updates, photos, and notes throughout the day so families always feel connected to their child's world."),
        ("What is school readiness at Flourish ELC?", "Our school readiness programs develop literacy, numeracy, independence, and social and emotional skills so children are genuinely ready to embrace kindergarten."),
    ]
)

# ── PAGE 17 ── Bateman childcare
add_page(
    doc, 17,
    "https://flourishelc.com.au/daycare-childcare-bateman-wa/",
    "Location Page", "childcare near bateman", "daycare near bateman",
    "Childcare Near Bateman WA | Flourish Early Learning Centre",
    "Looking for childcare near Bateman? Flourish ELC is purpose-built with EYLF-guided learning, five daily meals, qualified educators, and a real-time parent app. Book a tour.",
    "Childcare Near Bateman: Qualified Educators, Warm Community, Happy Children",
    "Flourish Early Learning Centre is a purpose-built early learning centre near Bateman where families find exactly what they have been looking for: a warm, qualified, and genuinely child-centred environment. Guided by the Early Years Learning Framework (EYLF) -- Belonging, Being and Becoming -- our programs and our people are dedicated to making each child's early years as rich and joyful as possible. Families near Bateman choose Flourish ELC because the quality is evident from the moment you walk through the door.",
    "Programs That Develop Children Across Every Area",
    "Our curriculum near Bateman is broad, deep, and full of life. STEM Learning, Creative Arts, Language, Mathematics, Fine Motor Skills, Music and Movement, Cultural Understanding, Social Justice, Independence, Exercise, Health, and Social and Emotional development are all woven into the daily rhythm of life at Flourish ELC. Children are not just supervised; they are actively growing, learning, and developing a genuine love of discovery.\n\nFive nutritious meals are prepared and served every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. Our school readiness programs build the academic, social, and emotional foundations children need to step confidently into kindergarten, ready to take on everything formal schooling brings.",
    "Purpose-Built, NQF-Aligned, and Genuinely Welcoming",
    "Our purpose-built centre near Bateman meets and exceeds the National Quality Framework (NQF). Families can be confident that the care their children receive is assessed against the highest national standards -- and that we hold ourselves to those standards every day, not just at assessment time.\n\nWe are open Monday to Friday, 7:00am to 6:00pm. Our parent communication app keeps families near Bateman genuinely connected with real-time photos, updates, and observations throughout the day.",
    "Frequently Asked Questions",
    [
        ("What age groups does Flourish ELC near Bateman cater for?", "We welcome children from 6 weeks to school age, with dedicated rooms and programs for each developmental stage."),
        ("What meals are included?", "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. Dietary requirements and allergies are always accommodated."),
        ("What are your opening hours near Bateman?", "Monday to Friday, 7:00am to 6:00pm."),
        ("How do I stay connected to my child's day?", "Our parent communication app delivers real-time updates and photos so you always know what your child is exploring and experiencing."),
        ("How does Flourish ELC prepare children for school?", "School readiness programs develop literacy, numeracy, independence, and social and emotional skills progressively so children transition to kindergarten with confidence."),
    ]
)

# ── PAGE 18 ── Rossmoyne childcare
add_page(
    doc, 18,
    "https://flourishelc.com.au/daycare-childcare-rossmoyne-wa/",
    "Location Page", "childcare near rossmoyne", "daycare near rossmoyne",
    "Childcare Near Rossmoyne WA | Flourish Early Learning Centre",
    "Looking for childcare near Rossmoyne? Flourish ELC is purpose-built with EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book today.",
    "Childcare Near Rossmoyne: Where Learning Feels Like Home",
    "Flourish Early Learning Centre is a warm, purpose-built early learning centre near Rossmoyne where children are welcomed into an environment guided by the Early Years Learning Framework (EYLF) -- Belonging, Being and Becoming. Our qualified, passionate educators are the heart of Flourish ELC, and it shows. Families near Rossmoyne choose us because their children genuinely love coming, and the development they witness speaks for itself.",
    "A Curriculum That Celebrates Every Child",
    "Our programs near Rossmoyne weave together all the learning areas that matter in a child's early years. Language, Mathematics, and Fine Motor Skills build academic foundations. STEM Learning and Creative Arts spark curiosity and inventive thinking. Music and Movement, Cultural Understanding, Social Justice, Independence, Exercise, Health, and Social and Emotional development ensure that every child -- in every dimension -- is growing and flourishing.\n\nFive nutritious meals are served each weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. Our school readiness programs are carefully integrated throughout our curriculum so that children near Rossmoyne are socially, emotionally, and academically prepared to embrace kindergarten when the time comes.",
    "Purpose-Built and NQF-Certified",
    "Our purpose-built centre near Rossmoyne meets and exceeds the National Quality Framework (NQF). Every aspect of the centre has been assessed against Australia's national quality standards -- from the learning environments and educator qualifications to the family partnerships and leadership practices. We are proud of that standard and work hard every day to uphold and exceed it.\n\nWe are open Monday to Friday, 7:00am to 6:00pm. Our parent communication app keeps families near Rossmoyne connected throughout the day with real-time updates, photos, and notes from our educators.",
    "Frequently Asked Questions",
    [
        ("What ages do you cater for near Rossmoyne?", "We welcome children from 6 weeks to school age, with dedicated programs and spaces for each developmental stage."),
        ("What meals are included daily?", "Five nutritious meals: breakfast, morning tea, lunch, afternoon tea, and a late snack prepared fresh every weekday."),
        ("What are the opening hours?", "Monday to Friday, 7:00am to 6:00pm."),
        ("How does the parent app keep me connected?", "The app delivers real-time updates, photos, and observations throughout the day so you always have a window into your child's world."),
        ("How does school readiness work at Flourish ELC?", "We build literacy, numeracy, social and emotional skills, and independence progressively through our programs so children are genuinely ready for kindergarten."),
    ]
)

# ── PAGE 19 ── Shelley childcare
add_page(
    doc, 19,
    "https://flourishelc.com.au/daycare-childcare-shelley-wa/",
    "Location Page", "childcare near shelley", "daycare near shelley",
    "Childcare Near Shelley WA | Flourish Early Learning Centre",
    "Looking for childcare near Shelley? Flourish ELC is purpose-built with EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour.",
    "Childcare Near Shelley: A Place Where Children Thrive",
    "Flourish Early Learning Centre is a purpose-built early learning centre near Shelley where families find something genuinely special: a warm community, qualified educators, and programs shaped by the Early Years Learning Framework (EYLF) -- Belonging, Being and Becoming. At Flourish ELC, children are known, celebrated, and supported in becoming the best version of themselves. Families near Shelley trust us because their children love it here, and the learning is real.",
    "Programs That Cover the Whole Child",
    "The Flourish ELC curriculum near Shelley is rich and intentional. STEM Learning, Creative Arts, Language, Mathematics, Fine Motor Skills, Music and Movement, Cultural Understanding, Social Justice, Independence, Exercise, Health, and Social and Emotional development are all part of the daily experience. Children do not just spend time here; they grow here in every meaningful sense.\n\nFive nutritious meals are prepared and served each weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. Our school readiness programs are woven through the curriculum from an early age, building the literacy, numeracy, social, and emotional skills children need to begin kindergarten with confidence, enthusiasm, and a genuine love of learning.",
    "Purpose-Built, NQF-Aligned, Open Every Weekday",
    "Our purpose-built centre near Shelley meets and exceeds the National Quality Framework (NQF), which means families can trust that every dimension of the centre -- the people, the spaces, the programs, and the partnerships -- is held to Australia's highest national standards for early education. We are proud of that commitment and live it every day.\n\nWe are open Monday to Friday, 7:00am to 6:00pm. Our parent communication app keeps families near Shelley genuinely connected throughout the day with real-time updates, photos, and observations from our educators.",
    "Frequently Asked Questions",
    [
        ("What ages do you cater for near Shelley?", "From 6 weeks to school age, with dedicated programs and spaces for infants, toddlers, and preschoolers."),
        ("What meals are provided?", "Five nutritious meals each weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. All dietary needs and allergies are carefully accommodated."),
        ("What are your opening hours?", "Monday to Friday, 7:00am to 6:00pm."),
        ("How do families stay connected during the day?", "Our parent communication app sends real-time updates, photos, and observations so you are always close to your child's world."),
        ("What does the EYLF mean for my child's day at Flourish ELC?", "The EYLF guides our educators in planning rich, intentional experiences that prioritise each child's sense of belonging, wellbeing, and engagement in learning."),
    ]
)

# ── PAGE 20 ── Willetton childcare
add_page(
    doc, 20,
    "https://flourishelc.com.au/daycare-childcare-willetton-wa/",
    "Location Page", "childcare near willetton", "daycare near willetton",
    "Childcare Near Willetton WA | Flourish Early Learning Centre",
    "Looking for childcare near Willetton? Flourish ELC is purpose-built with EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book today.",
    "Childcare Near Willetton: Warmth, Quality, and Children Who Flourish",
    "Flourish Early Learning Centre is a purpose-built childcare centre near Willetton where everything is designed around children having the very best start in life. Guided by the Early Years Learning Framework (EYLF) -- Belonging, Being and Becoming -- our programs are thoughtful, joyful, and genuinely focused on each child. Our qualified, passionate educators are what families near Willetton talk about most -- because great educators make everything else possible.",
    "A Curriculum Full of Discovery and Growth",
    "Our childcare programs near Willetton are as broad as they are rich. Language, Mathematics, and Fine Motor Skills build the academic foundations children carry into school. STEM Learning and Creative Arts nurture the inventive, curious thinkers families want their children to become. Music and Movement, Cultural Understanding, Social Justice, Independence, Exercise, Health, and Social and Emotional development complete a curriculum that honours the whole child.\n\nFive nutritious meals are served fresh each weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. Our school readiness programs are embedded throughout our curriculum, building the skills and confidence children need to begin kindergarten ready to embrace everything formal education has to offer.",
    "Purpose-Built, NQF-Aligned, and Welcoming Every Day",
    "Our purpose-built centre near Willetton meets and exceeds the National Quality Framework (NQF). That standard means every area of our centre -- the learning environments, the educator qualifications, the programs, and the partnerships with families -- has been assessed against Australia's highest national benchmarks. We take genuine pride in the quality that represents.\n\nWe are open Monday to Friday, 7:00am to 6:00pm. Our parent communication app delivers real-time updates and photos throughout the day so families near Willetton always feel connected to their child's world.",
    "Frequently Asked Questions",
    [
        ("What ages do you cater for near Willetton?", "From 6 weeks to school age, with dedicated programs and spaces for each developmental stage."),
        ("What meals are included at Flourish ELC?", "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. All dietary requirements and allergies are accommodated."),
        ("What are your opening hours near Willetton?", "Monday to Friday, 7:00am to 6:00pm."),
        ("How does the parent communication app work?", "It sends real-time updates and photos throughout the day, keeping you genuinely connected to your child's experience."),
        ("How does Flourish ELC prepare children for school?", "School readiness programs build literacy, numeracy, independence, and social and emotional skills progressively so children are ready and excited for kindergarten."),
    ]
)

print("Pages 11-20 done")

# ── PAGE 21 ── Brentwood childcare
add_page(
    doc, 21,
    "https://flourishelc.com.au/daycare-childcare-brentwood-wa/",
    "Location Page", "childcare near brentwood", "daycare near brentwood",
    "Childcare Near Brentwood WA | Flourish Early Learning Centre",
    "Looking for childcare near Brentwood? Flourish ELC is purpose-built with EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book today.",
    "Childcare Near Brentwood: Where Children Are Known, Loved, and Inspired",
    "Flourish Early Learning Centre is a purpose-built early learning centre near Brentwood where families find something they often say is hard to describe but impossible to miss: the feeling that their child truly belongs here. Our programs are guided by the Early Years Learning Framework (EYLF) -- Belonging, Being and Becoming -- and our qualified, passionate educators bring that philosophy to life with warmth and expertise. Families near Brentwood choose Flourish ELC because it just feels right -- and the learning backs that feeling up.",
    "Programs That Grow Children in Every Dimension",
    "Learning at Flourish ELC near Brentwood spans the full breadth of early childhood development. STEM Learning, Creative Arts, Language, Mathematics, Fine Motor Skills, Music and Movement, Cultural Understanding, Social Justice, Independence, Exercise, Health, and Social and Emotional development are all part of our daily curriculum. Children spend their days discovering, connecting, creating, and growing in ways they will carry with them for life.\n\nFive nutritious meals are prepared and served each weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. Our school readiness programs are thoughtfully woven through the curriculum, building the literacy, numeracy, social, and emotional skills children need to step confidently into kindergarten.",
    "Purpose-Built and Held to the Highest National Standards",
    "Our purpose-built centre near Brentwood meets and exceeds the National Quality Framework (NQF). Every element of the centre has been assessed against Australia's national quality standards -- and we hold ourselves to those standards because they matter to us, not just because they are required. That commitment is something Brentwood families feel and appreciate.\n\nWe are open Monday to Friday, 7:00am to 6:00pm. Our parent communication app keeps families near Brentwood connected throughout the day with real-time updates, photos, and observations from educators who genuinely care.",
    "Frequently Asked Questions",
    [
        ("What age groups do you cater for near Brentwood?", "From 6 weeks to school age, with dedicated programs and spaces for infants, toddlers, and preschoolers."),
        ("What meals are included each day?", "Five nutritious meals: breakfast, morning tea, lunch, afternoon tea, and a late snack. Dietary requirements and allergies are always managed carefully."),
        ("What are the opening hours?", "Monday to Friday, 7:00am to 6:00pm."),
        ("How will I know what my child is doing during the day?", "Our parent communication app delivers real-time photos and updates so you are always close to your child's world."),
        ("How does Flourish ELC prepare children for kindergarten?", "School readiness programs build literacy, numeracy, independence, and social and emotional skills progressively so children begin kindergarten with confidence and genuine readiness."),
    ]
)

# ── PAGE 22 ── Bull Creek kindergarten
add_page(
    doc, 22,
    "https://flourishelc.com.au/kindergarten-bull-creek-wa/",
    "Location Page", "kindergarten bull creek", "",
    "Kindergarten Bull Creek WA | Flourish Early Learning Centre",
    "Looking for kindergarten in Bull Creek? Flourish ELC offers EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour today.",
    "Kindergarten in Bull Creek: A Joyful, Qualified Start to School Life",
    "Flourish Early Learning Centre offers a warm, EYLF-guided kindergarten program in Bull Creek where children are prepared for school in the best possible way: through joyful, intentional, play-based learning. Our Early Years Learning Framework -- Belonging, Being and Becoming -- is not just a concept here; it is the lived experience of every child in our care. Our qualified educators know each child well and create experiences that light them up. Kindergarten in Bull Creek is wonderful when it happens at Flourish ELC.",
    "A Kindergarten Curriculum That Covers All the Foundations",
    "Our kindergarten program in Bull Creek develops the skills children need across every area. Language, Mathematics, and Fine Motor Skills lay the academic groundwork for a smooth transition into primary school. STEM Learning, Creative Arts, and Music and Movement build the curiosity and creative thinking that make children love to learn. Cultural Understanding, Social Justice, Social and Emotional development, Independence, Exercise, and Health ensure that the whole child -- not just the student -- is nurtured and celebrated.\n\nFive nutritious meals are served each weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. Our school readiness programs are the backbone of kindergarten at Flourish ELC, preparing children for primary school with confidence, resilience, and genuine excitement for the journey ahead.",
    "Qualified Educators, Purpose-Built Spaces, and Real-Time Connection",
    "Our purpose-built centre in Bull Creek meets and exceeds the National Quality Framework (NQF). Our kindergarten educators hold recognised early childhood qualifications and bring both professional expertise and genuine warmth to every session. The spaces children learn in have been thoughtfully designed to support the kind of active, engaged kindergarten experience that produces lasting results.\n\nWe are open Monday to Friday, 7:00am to 6:00pm -- hours designed to give families flexibility. Our parent communication app sends real-time updates, photos, and milestones straight to your phone so you never miss a moment of your child's kindergarten year.",
    "Frequently Asked Questions",
    [
        ("What age is kindergarten at Flourish ELC Bull Creek?", "Our kindergarten program is designed for children aged 4 to 5, in the year before they begin primary school."),
        ("What meals are provided during kindergarten?", "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack."),
        ("What are the kindergarten hours?", "Monday to Friday, 7:00am to 6:00pm."),
        ("How do parents stay connected during the kindergarten day?", "Our parent communication app delivers real-time updates and photos so families always know what their child has been discovering."),
        ("How does Flourish ELC prepare kindergarteners for primary school?", "Our school readiness programs build literacy, numeracy, independence, and social and emotional skills so children begin primary school with genuine confidence and capability."),
    ]
)

# ── PAGE 23 ── Leeming kindergarten
add_page(
    doc, 23,
    "https://flourishelc.com.au/kindergarten-leeming-wa/",
    "Location Page", "kindergarten near leeming", "",
    "Kindergarten Near Leeming WA | Flourish Early Learning Centre",
    "Looking for kindergarten near Leeming? Flourish ELC offers EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour today.",
    "Kindergarten Near Leeming: Where Children Grow Ready, Willing, and Able",
    "Flourish Early Learning Centre offers a warm, play-based kindergarten near Leeming guided by the Early Years Learning Framework (EYLF) -- Belonging, Being and Becoming. Our kindergarten year is one of the most important a child will have, and we take that responsibility joyfully and seriously. Our qualified educators design kindergarten experiences that build confident, curious, capable children who walk into primary school genuinely ready for everything that comes next. Families near Leeming trust us with that year, and we do not take it lightly.",
    "Kindergarten Programs That Develop Every Part of Your Child",
    "Our kindergarten curriculum near Leeming draws on all the learning areas that matter most in the year before school. Language, Mathematics, and Fine Motor Skills develop the academic foundations children need. STEM Learning and Creative Arts build the problem-solving and expressive skills that set children apart. Music and Movement, Cultural Understanding, Social Justice, Independence, Exercise, Health, and Social and Emotional development complete a kindergarten program that prepares children for school -- and for life.\n\nFive nutritious meals are served fresh each weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. School readiness runs through everything we do in our kindergarten program, from the way we build literacy and numeracy skills to the way we develop independence and resilience in each child.",
    "Purpose-Built, NQF-Aligned, and Open Every Weekday",
    "Our purpose-built centre near Leeming meets and exceeds the National Quality Framework (NQF). Our kindergarten educators are qualified professionals who are passionate about this stage of development and committed to delivering the best possible year for every child. The spaces they work in have been purpose-built to support active, engaged kindergarten learning.\n\nWe are open Monday to Friday, 7:00am to 6:00pm. Our parent communication app keeps families near Leeming genuinely connected with real-time updates, photos, and kindergarten milestones throughout the day.",
    "Frequently Asked Questions",
    [
        ("What age is the kindergarten program near Leeming?", "Our kindergarten program serves children aged 4 to 5 in the year before they begin primary school."),
        ("What meals are included in the kindergarten program?", "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack."),
        ("What are the kindergarten hours near Leeming?", "Monday to Friday, 7:00am to 6:00pm."),
        ("How does the parent app work during the kindergarten year?", "It sends real-time updates, photos, and milestones throughout the day so you always feel connected to your child's kindergarten experience."),
        ("How does Flourish ELC get kindergarteners ready for primary school?", "Our school readiness programs develop literacy, numeracy, independence, and social and emotional skills throughout the kindergarten year so children begin primary school with confidence and genuine capability."),
    ]
)

# ── PAGE 24 ── Bateman kindergarten
add_page(
    doc, 24,
    "https://flourishelc.com.au/kindergarten-bateman-wa/",
    "Location Page", "kindergarten near bateman", "",
    "Kindergarten Near Bateman WA | Flourish Early Learning Centre",
    "Looking for kindergarten near Bateman? Flourish ELC offers EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour today.",
    "Kindergarten Near Bateman: Play-Based, Purposeful, and Genuinely Wonderful",
    "Flourish Early Learning Centre offers a warm, EYLF-guided kindergarten program near Bateman where the year before primary school is spent in the very best way: learning through play, connecting with others, and growing into a confident, capable child. The Early Years Learning Framework -- Belonging, Being and Becoming -- shapes every experience in our kindergarten, and our qualified, passionate educators bring it to life with skill and joy. Families near Bateman choose Flourish ELC for their child's kindergarten year because they can see and feel the quality.",
    "A Kindergarten Year That Covers Every Foundation",
    "Learning at Flourish ELC in the kindergarten year is comprehensive and joyful. Language, Mathematics, and Fine Motor Skills are developed through engaging, hands-on play. STEM Learning, Creative Arts, and Music and Movement make discovery genuinely exciting. Cultural Understanding, Social Justice, Social and Emotional development, Independence, Exercise, and Health ensure that every child is growing as a whole person -- not just building academic skills, but building character, confidence, and a real love of learning.\n\nFive nutritious meals are prepared and served each weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. School readiness is the golden thread running through our kindergarten curriculum, building everything children need to begin primary school feeling ready, excited, and completely capable.",
    "Purpose-Built, Qualified, and Connected to Families",
    "Our purpose-built centre near Bateman meets and exceeds the National Quality Framework (NQF). Our kindergarten educators are qualified early childhood professionals who are passionate about this pivotal year of development. The learning spaces are purpose-built for kindergarten-age children: active, bright, and designed for the kind of engaged learning that produces lifelong results.\n\nWe are open Monday to Friday, 7:00am to 6:00pm. Our parent communication app keeps families near Bateman genuinely connected throughout the kindergarten day with real-time photos, updates, and milestones.",
    "Frequently Asked Questions",
    [
        ("What age does the kindergarten near Bateman cater for?", "Our kindergarten program is designed for children aged 4 to 5 in the year before they start primary school."),
        ("What meals are provided during kindergarten?", "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack."),
        ("What are the kindergarten hours?", "Monday to Friday, 7:00am to 6:00pm."),
        ("How will I know what my child has been doing in kindergarten?", "Our parent communication app sends real-time updates and photos so families near Bateman always feel close to their child's kindergarten year."),
        ("How does Flourish ELC prepare kindergarteners for primary school?", "Literacy, numeracy, independence, and social and emotional skills are built progressively throughout the kindergarten year so children begin primary school genuinely ready."),
    ]
)

# ── PAGE 25 ── Rossmoyne kindergarten
add_page(
    doc, 25,
    "https://flourishelc.com.au/kindergarten-rossmoyne-wa/",
    "Location Page", "kindergarten near rossmoyne", "",
    "Kindergarten Near Rossmoyne WA | Flourish Early Learning Centre",
    "Looking for kindergarten near Rossmoyne? Flourish ELC offers EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour today.",
    "Kindergarten Near Rossmoyne: A Joyful, Confident Start to School Life",
    "Flourish Early Learning Centre offers a warm, purpose-built kindergarten near Rossmoyne where children spend the year before school exactly as they should: playing, discovering, growing, and becoming. Our kindergarten programs are guided by the Early Years Learning Framework (EYLF) -- Belonging, Being and Becoming -- and delivered by qualified educators who genuinely love the kindergarten years. Families near Rossmoyne choose Flourish ELC because their children flourish here -- and because the transition to primary school that follows is confident and smooth.",
    "A Kindergarten Curriculum Built for the Whole Child",
    "Our kindergarten programs near Rossmoyne develop children across every important dimension. Language, Mathematics, and Fine Motor Skills provide the academic groundwork for primary school. Creative Arts, STEM Learning, and Music and Movement cultivate the curiosity, inventiveness, and confidence that will serve children throughout their education. Cultural Understanding, Social Justice, Independence, Exercise, Health, and Social and Emotional development complete a kindergarten curriculum that cares as much about who children are as what they know.\n\nFive nutritious meals are served fresh each weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. School readiness programs run as a thread through the entire kindergarten year, building the literacy, numeracy, social, and emotional skills children need to walk into primary school with a smile.",
    "Qualified, Warm Educators in a Purpose-Built Environment",
    "Our purpose-built kindergarten near Rossmoyne meets and exceeds the National Quality Framework (NQF). Our educators hold recognised early childhood qualifications and approach the kindergarten year with both professional rigour and genuine affection for the children in their care. The spaces children learn in are designed to inspire active, engaged kindergarten learning.\n\nWe are open Monday to Friday, 7:00am to 6:00pm. Our parent communication app keeps families near Rossmoyne genuinely connected throughout the kindergarten day with real-time updates, photos, and milestones from educators who care.",
    "Frequently Asked Questions",
    [
        ("What age is the kindergarten near Rossmoyne?", "Our kindergarten program is designed for children aged 4 to 5 -- the year before primary school begins."),
        ("What meals are included in the kindergarten program?", "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack."),
        ("What are the kindergarten opening hours?", "Monday to Friday, 7:00am to 6:00pm."),
        ("How does the parent communication app keep me connected?", "It sends real-time updates and photos throughout the kindergarten day so you always have a window into your child's world."),
        ("How does Flourish ELC prepare kindergarteners for primary school?", "Our school readiness programs develop literacy, numeracy, independence, and social and emotional skills so children begin primary school ready, confident, and genuinely excited."),
    ]
)

# ── PAGE 26 ── Shelley kindergarten
add_page(
    doc, 26,
    "https://flourishelc.com.au/kindergarten-shelley-wa/",
    "Location Page", "kindergarten near shelley", "",
    "Kindergarten Near Shelley WA | Flourish Early Learning Centre",
    "Looking for kindergarten near Shelley? Flourish ELC offers EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour today.",
    "Kindergarten Near Shelley: Where the Best Possible Year Before School Happens",
    "Flourish Early Learning Centre offers a warm, play-based kindergarten near Shelley guided by the Early Years Learning Framework (EYLF) -- Belonging, Being and Becoming. The year before primary school is one of the most important a child will have, and our qualified, passionate educators pour everything into making it wonderful. Children in our kindergarten near Shelley grow academically, socially, and emotionally -- and they do so in an environment that is safe, joyful, and genuinely child-centred. Families near Shelley choose Flourish ELC because what happens here sets children up for a lifetime.",
    "Kindergarten Programs Full of Purpose and Joy",
    "Our kindergarten curriculum near Shelley is rich and intentional. Language, Mathematics, and Fine Motor Skills lay the academic foundations for primary school. STEM Learning and Creative Arts build the problem-solving and expressive skills that make children confident, inventive learners. Music and Movement, Cultural Understanding, Social Justice, Independence, Exercise, Health, and Social and Emotional development complete a kindergarten program that develops the whole child.\n\nFive nutritious meals are prepared and served each weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. School readiness is embedded in everything we do in the kindergarten year, building the skills and confidence children need to step into primary school ready to love it.",
    "Purpose-Built, NQF-Aligned, and Genuinely Caring",
    "Our purpose-built centre near Shelley meets and exceeds the National Quality Framework (NQF). Our kindergarten educators hold recognised qualifications and bring both professional expertise and warmth to every session of every day. The learning spaces are purpose-designed for active, engaged, kindergarten-age learners.\n\nWe are open Monday to Friday, 7:00am to 6:00pm. Our parent communication app keeps families near Shelley genuinely connected with real-time photos, updates, and kindergarten milestones throughout the day.",
    "Frequently Asked Questions",
    [
        ("What age is kindergarten at Flourish ELC near Shelley?", "Our kindergarten program serves children aged 4 to 5 -- the year before they begin primary school."),
        ("What meals are provided?", "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack."),
        ("What are the opening hours?", "Monday to Friday, 7:00am to 6:00pm."),
        ("How will I know what my child has been up to during kindergarten?", "Our parent communication app delivers real-time updates and photos so you always feel connected to your child's kindergarten day."),
        ("How does Flourish ELC approach school readiness in kindergarten?", "Literacy, numeracy, independence, and social and emotional skills are built throughout the kindergarten year so children begin primary school ready, confident, and excited."),
    ]
)

# ── PAGE 27 ── Willetton kindergarten
add_page(
    doc, 27,
    "https://flourishelc.com.au/kindergarten-willetton-wa/",
    "Location Page", "kindergarten near willetton", "",
    "Kindergarten Near Willetton WA | Flourish Early Learning Centre",
    "Looking for kindergarten near Willetton? Flourish ELC offers EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour today.",
    "Kindergarten Near Willetton: Play-Based, Warm, and Built for the Best Start",
    "Flourish Early Learning Centre offers a wonderful kindergarten near Willetton where children spend the year before primary school growing in every direction. Our kindergarten programs are guided by the Early Years Learning Framework (EYLF) -- Belonging, Being and Becoming -- and delivered by qualified, passionate educators who make every day count. Families near Willetton choose our kindergarten because they see the quality in the detail: the care, the curriculum, the community, and the children who thrive here.",
    "A Kindergarten Year Across Every Learning Area",
    "Learning at Flourish ELC in the kindergarten year is comprehensive, connected, and full of life. Language, Mathematics, and Fine Motor Skills develop the academic skills children carry into primary school. STEM Learning, Creative Arts, and Music and Movement make discovery genuinely joyful. Cultural Understanding, Social Justice, Social and Emotional development, Independence, Exercise, and Health ensure children are growing as whole people -- confident, curious, and ready for whatever comes next.\n\nFive nutritious meals are served fresh each weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. School readiness is at the heart of our kindergarten curriculum, building everything children need to walk into their first day of primary school with a big smile and genuine capability.",
    "Qualified Educators, Purpose-Built Spaces, Every Weekday",
    "Our purpose-built centre near Willetton meets and exceeds the National Quality Framework (NQF). Our kindergarten educators are qualified early childhood professionals who are passionate about this pivotal developmental stage. The learning environments are designed to support the active, engaged, joyful kindergarten experience children deserve.\n\nWe are open Monday to Friday, 7:00am to 6:00pm. Our parent communication app keeps families near Willetton genuinely connected throughout the kindergarten day with real-time updates, photos, and milestones.",
    "Frequently Asked Questions",
    [
        ("What age is kindergarten at Flourish ELC near Willetton?", "Our kindergarten program serves children aged 4 to 5 in the year before they begin primary school."),
        ("What meals are included in the kindergarten program?", "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack."),
        ("What are the kindergarten hours near Willetton?", "Monday to Friday, 7:00am to 6:00pm."),
        ("How does the parent app keep me connected during the kindergarten year?", "It sends real-time updates, photos, and milestones throughout the day so you always feel close to your child's kindergarten experience."),
        ("How does Flourish ELC prepare kindergarteners for primary school?", "Literacy, numeracy, independence, and social and emotional skills are built throughout the year so children begin primary school ready, confident, and genuinely excited for what lies ahead."),
    ]
)

# ── PAGE 28 ── Brentwood kindergarten
add_page(
    doc, 28,
    "https://flourishelc.com.au/kindergarten-brentwood-wa/",
    "Location Page", "kindergarten near brentwood", "",
    "Kindergarten Near Brentwood WA | Flourish Early Learning Centre",
    "Looking for kindergarten near Brentwood? Flourish ELC offers EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app. Book your tour today.",
    "Kindergarten Near Brentwood: Your Child's Best Year Before School Starts Here",
    "Flourish Early Learning Centre offers a warm, purpose-built kindergarten near Brentwood where the year before primary school is spent exactly as it should be: playing, learning, growing, and belonging. Guided by the Early Years Learning Framework (EYLF) -- Belonging, Being and Becoming -- our kindergarten programs are shaped around each child's unique journey. Our qualified, passionate educators bring both skill and genuine heart to every kindergarten session. Families near Brentwood trust us with the kindergarten year because they see, from day one, that their child is in wonderful hands.",
    "A Kindergarten Curriculum That Develops the Whole Child",
    "Our kindergarten near Brentwood is built around the learning areas that matter most in the year before school. Language, Mathematics, and Fine Motor Skills develop the academic foundations children take into primary school. STEM Learning, Creative Arts, and Music and Movement make kindergarten genuinely exciting. Cultural Understanding, Social Justice, Independence, Exercise, Health, and Social and Emotional development ensure that every child is growing -- not just academically, but as a whole, confident, curious person.\n\nFive nutritious meals are prepared and served each weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. School readiness is woven through every part of our kindergarten curriculum, building the literacy, numeracy, social, and emotional skills children need to begin primary school feeling ready and genuinely excited.",
    "Qualified, Caring Educators in a Purpose-Built Setting",
    "Our purpose-built centre near Brentwood meets and exceeds the National Quality Framework (NQF). Our kindergarten educators hold recognised early childhood qualifications and are committed to making the kindergarten year as meaningful as it is joyful. The learning spaces are purpose-designed for active, engaged kindergarten-age learners.\n\nWe are open Monday to Friday, 7:00am to 6:00pm. Our parent communication app keeps families near Brentwood genuinely connected throughout the kindergarten day with real-time updates, photos, and milestones from educators who care deeply.",
    "Frequently Asked Questions",
    [
        ("What age is the kindergarten near Brentwood?", "Our kindergarten program is designed for children aged 4 to 5 -- the year before they begin primary school."),
        ("What meals are provided during the kindergarten year?", "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack."),
        ("What are the kindergarten opening hours near Brentwood?", "Monday to Friday, 7:00am to 6:00pm."),
        ("How do families stay connected during the kindergarten day?", "Our parent communication app delivers real-time updates and photos so families near Brentwood always feel close to their child's kindergarten world."),
        ("How does Flourish ELC prepare kindergarteners for primary school?", "School readiness programs build literacy, numeracy, independence, and social and emotional skills throughout the kindergarten year so children begin primary school with confidence and capability."),
    ]
)

# ── PAGE 29 ── Areas We Serve (childcare near me)
add_page(
    doc, 29,
    "https://flourishelc.com.au/areas-we-serve/",
    "Location Page", "childcare near me", "daycare near me, preschool near me, kindergarten near me",
    "Find Flourish Early Learning Centres Near You",
    "Looking for childcare near me? Flourish ELC has warm, purpose-built centres with EYLF-guided programs, five daily meals, qualified educators, and a real-time parent app.",
    "Childcare Near Me: Find a Flourish Early Learning Centre Close to Your Family",
    "Whatever brought you to searching for childcare near me -- a new baby, a house move, or simply the feeling that it is time -- you have come to the right place. Flourish Early Learning Centre is a family of purpose-built early learning centres guided by the Early Years Learning Framework (EYLF): Belonging, Being and Becoming. Our qualified, passionate educators create environments where children feel at home, families feel supported, and learning feels like the joyful adventure it is meant to be.",
    "What Makes Flourish ELC Special Wherever You Find Us",
    "Every Flourish ELC centre -- whether you are searching for childcare, daycare, preschool, or kindergarten near you -- is built around the same beautiful principles. Our programs cover Creative Arts, STEM Learning, Language, Mathematics, Fine Motor Skills, Music and Movement, Cultural Understanding, Social Justice, Independence, Exercise, Health, and Social and Emotional development. These are not just learning areas; they are the threads of a rich, whole-child experience that children and families love.\n\nFive nutritious meals are served fresh every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. Our centres are open 7:00am to 6:00pm Monday to Friday. Our parent communication app keeps every family connected with real-time updates and photos throughout the day. School readiness programs are woven through our curriculum so that children transition to kindergarten and primary school with confidence, capability, and enthusiasm.",
    "Qualified Educators, Purpose-Built Spaces, Real-Time Connection",
    "Every Flourish ELC centre is staffed by qualified educators who meet and exceed the National Quality Framework (NQF). Our facilities are purpose-built for children from 6 weeks to school age -- not repurposed spaces, but environments designed from the ground up around how young children learn, rest, play, and grow. Families searching for preschool or kindergarten near me will find that same consistent standard across every Flourish ELC location.\n\nOur parent communication app is one of the things families love most. It delivers real-time photos, updates, and observations throughout the day so even when you are at work, you are never far from your child's world. It is a small but meaningful reflection of how much we value the partnership between families and educators.",
    "Frequently Asked Questions",
    [
        ("What age groups do Flourish ELC centres cater for?", "All our centres welcome children from 6 weeks through to school age, with dedicated programs and spaces for each developmental stage."),
        ("What meals does Flourish ELC provide?", "Five nutritious meals every weekday: breakfast, morning tea, lunch, afternoon tea, and a late snack. All dietary requirements and allergies are carefully managed."),
        ("What are the opening hours across Flourish ELC centres?", "Our centres are open Monday to Friday, 7:00am to 6:00pm."),
        ("How does the parent communication app work?", "The app sends real-time updates, photos, and milestones throughout the day so families always feel genuinely connected to their child's world."),
        ("What does the EYLF mean for my child at Flourish ELC?", "The Early Years Learning Framework guides how our educators plan, observe, and respond to each child -- always with belonging, wellbeing, and engagement at the heart of practice. It is the foundation of everything we do."),
    ]
)

print("Pages 21-29 done")

# Save
output_path = "/home/invoi/fahad_projects/clients/flourishelc.com.au/flourishelc.com.au_v2.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
