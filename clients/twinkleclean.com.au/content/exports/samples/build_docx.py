"""House-style .docx builder for Twinkle Clean content deliverable.

House style (per client spec):
- "Page N" at the very beginning of each page's content.
- Headings labelled H1:/H2:/H3: (with colon), bold, SAME font size as body, GREY.
- FAQ questions are H3.
- All body content BLACK.
Reads clean markdown (## / ### / - / **bold**); the H1:/H2:/H3: labels, Page N and
colours are applied only here at render time, so the markdown source stays gate-valid.
"""
import re, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_BREAK

BODY_PT = 11
GREY = RGBColor(0x59, 0x59, 0x59)   # heading grey
BLACK = RGBColor(0x00, 0x00, 0x00)  # body black

def add_inline(paragraph, text, *, bold_all=False, color=BLACK, size=BODY_PT):
    """Add text to a paragraph, honouring **bold** spans."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        b = bold_all
        t = part
        if part.startswith("**") and part.endswith("**"):
            t = part[2:-2]
            b = True
        run = paragraph.add_run(t)
        run.font.size = Pt(size)
        run.font.bold = b
        run.font.color.rgb = color

def add_heading(doc, level, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"H{level}: {text}")
    run.font.size = Pt(BODY_PT)   # same size as body
    run.font.bold = True
    run.font.color.rgb = GREY

def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_inline(p, text)

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    add_inline(p, text)

def add_page_label(doc, n, first):
    p = doc.add_paragraph()
    if not first:
        p.add_run().add_break(WD_BREAK.PAGE)
    run = p.add_run(f"Page {n}")
    run.font.size = Pt(BODY_PT)
    run.font.bold = True
    run.font.color.rgb = BLACK

def render_markdown(doc, md, page_no, first):
    add_page_label(doc, page_no, first)
    for raw in md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            add_heading(doc, 3, line[4:].strip())
        elif line.startswith("## "):
            add_heading(doc, 2, line[3:].strip())
        elif line.startswith("# "):
            add_heading(doc, 1, line[2:].strip())
        elif line.startswith("- "):
            add_bullet(doc, line[2:].strip())
        else:
            add_body(doc, line)

def build(entries, out_path):
    """entries: list of (page_no, markdown_text). Renders all into one docx."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(BODY_PT)
    style.font.color.rgb = BLACK
    for i, (page_no, md) in enumerate(entries):
        render_markdown(doc, md, page_no, first=(i == 0))
    doc.save(out_path)
    return out_path

if __name__ == "__main__":
    # sample: single page
    md_path = Path(sys.argv[1])
    out = Path(sys.argv[2])
    build([(1, md_path.read_text())], out)
    print("wrote", out)
