"""House-style .docx builder for the Twinkle Clean content deliverable.

House style (per client spec):
- "Page N" at the very beginning of each page's content. No table of contents.
- Heading TAG ("H1:"/"H2:"/"H3:") is GREY; the heading TEXT itself is BLACK.
  Both bold, same font size as body.
- FAQ questions are H3.
- All body content BLACK. Links render as real hyperlinks in standard blue,
  underlined, NOT bold.

Reads clean markdown (## / ### / - / **bold** / [text](url)). The H1:/H2:/H3:
labels, Page N and colours are applied only here at render time, so the markdown
source stays valid for the content audit gate.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BODY_PT = 11
BODY_FONT = "Calibri"
TITLE_PT = 18          # "Page N" title — larger than everything else
TITLE_FONT = "Georgia"  # the only element in a different font, so it stands out
GREY = RGBColor(0x59, 0x59, 0x59)    # heading TAG only
BLACK = RGBColor(0x00, 0x00, 0x00)   # body + heading text
LINK_BLUE = RGBColor(0x05, 0x63, 0xC1)  # standard Word hyperlink blue

# [text](url)
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
BOLD_RE = re.compile(r"(\*\*[^*]+\*\*)")


def _style_run(run, *, bold=False, color=BLACK, size=BODY_PT):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    return run


def _add_hyperlink(paragraph, text, url):
    """Insert a real clickable hyperlink: standard blue, underlined, not bold."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(BODY_PT * 2))  # half-points
    rPr.append(sz)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_text_with_bold(paragraph, text, *, color=BLACK, bold_all=False):
    """Add text honouring **bold** spans."""
    for part in BOLD_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            _style_run(paragraph.add_run(part[2:-2]), bold=True, color=color)
        else:
            _style_run(paragraph.add_run(part), bold=bold_all, color=color)


def _add_inline(paragraph, text):
    """Add text, rendering markdown links as real blue hyperlinks."""
    pos = 0
    for m in LINK_RE.finditer(text):
        if m.start() > pos:
            _add_text_with_bold(paragraph, text[pos:m.start()])
        _add_hyperlink(paragraph, m.group(1), m.group(2))
        pos = m.end()
    if pos < len(text):
        _add_text_with_bold(paragraph, text[pos:])


def add_heading(doc, level, text):
    """Grey 'HN:' tag + black heading text, both bold, body size."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    _style_run(p.add_run(f"H{level}: "), bold=True, color=GREY)
    _style_run(p.add_run(text), bold=True, color=BLACK)


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    _add_inline(p, text)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    _add_inline(p, text)


def add_page_label(doc, n, first):
    """Page N: a real Word Heading 1.

    Uses the built-in "Heading 1" STYLE (not just big text) so the document has a
    navigable outline: Word's Navigation pane and Google Docs' outline / document
    tabs both read heading styles. The Georgia look is layered on top of the style.

    Every page starts on a new page.
    """
    p = doc.add_paragraph(style="Heading 1")
    if not first:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"Page {n}")
    run.font.name = TITLE_FONT
    run.font.size = Pt(TITLE_PT)
    run.font.bold = True
    run.font.color.rgb = BLACK
    # Heading styles carry a theme font; pin the East Asian/complex-script slots
    # too so Word/Google Docs do not substitute the theme face back in.
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), TITLE_FONT)


def add_blank(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _style_run(p.add_run(""), color=BLACK)


def add_meta_line(doc, label, value, *, as_link=False):
    """Grey bold label + black value. URLs render as blue hyperlinks."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _style_run(p.add_run(f"{label}: "), bold=True, color=GREY)
    if as_link:
        _add_hyperlink(p, value, value)
    else:
        _style_run(p.add_run(value), color=BLACK)


def add_meta_block(doc, meta):
    """Header block, in order:
    URL -> [blank] -> primary keyword -> secondary keywords -> [blank]
        -> meta title -> meta description -> [blank]
    Supporting keywords are deliberately not shown.
    """
    add_meta_line(doc, "URL", meta["url"], as_link=True)
    add_blank(doc)

    kws = meta["keywords"]
    add_meta_line(doc, "Primary Keyword", kws[0] if kws else "")
    if len(kws) > 1:
        add_meta_line(doc, "Secondary Keywords", ", ".join(kws[1:]))
    add_blank(doc)

    add_meta_line(doc, "Meta Title", meta["meta_title"])
    add_meta_line(doc, "Meta Description", meta["meta_description"])
    add_blank(doc)


def render_markdown(doc, md, page_no, first, meta=None):
    add_page_label(doc, page_no, first)
    if meta:
        add_meta_block(doc, meta)
    for raw in md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            add_heading(doc, 3, line[4:].strip())
        elif line.startswith("## "):
            add_heading(doc, 2, line[3:].strip())
        elif line.startswith("# "):
            add_heading(doc, 1, line[2:].strip())
        elif line.startswith("- "):
            add_bullet(doc, line[2:].strip())
        else:
            add_body(doc, line)


def build(entries, out_path):
    """entries: list of (page_no, markdown_text) or (page_no, markdown_text, meta).

    Renders one combined .docx. No table of contents.
    """
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_PT)
    normal.font.color.rgb = BLACK
    for i, item in enumerate(entries):
        page_no, md = item[0], item[1]
        meta = item[2] if len(item) > 2 else None
        render_markdown(doc, md, page_no, first=(i == 0), meta=meta)
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    md_path, out = Path(sys.argv[1]), Path(sys.argv[2])
    build([(1, md_path.read_text())], out)
    print("wrote", out)
